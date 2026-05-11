from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"e:\face")
SOURCE_MD = BASE_DIR / "换脸系统_初稿_学术化版_精修后.md"
SOURCE_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_精修后_图文整改版.docx"
OUTPUT_MD = BASE_DIR / "换脸系统_初稿_学术化版_终稿统一整改版.md"
OUTPUT_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_终稿统一整改版.docx"

INSERTIONS = [
    (
        "图 2.1 人脸替换典型处理流程图",
        "图 2.1 将本课题中的人脸替换任务分解为素材进入系统、身份信息提取、目标区域处理、结果融合以及输出保存等关键阶段。该流程说明，本文讨论的对象并非单一算法算子，而是一条由前端交互、线程调度、底层引擎和结果管理共同构成的系统链路。只有这些环节能够连续衔接，系统才具备从输入到输出的完整可用性。",
    ),
    (
        "表 2.1 传统方法与深度学习方法对比",
        "表 2.1 通过实现依赖、适用场景、结果自然度与工程稳定性等维度，概括了两类技术路线在系统中的不同角色。对本科毕业设计而言，这种比较的意义并不在于简单判断孰优孰劣，而在于解释为什么当前项目需要同时保留传统方法与 InsightFace 两条处理路径，并据此在效果质量与环境兼容性之间取得平衡。",
    ),
    (
        "表 2.2 关键技术与系统作用映射表",
        "表 2.2 将算法、媒体处理、界面开发和数据管理等技术与其在系统中的职责逐一对应，使读者能够从工程分工角度理解每项技术存在的必要性。相较于孤立罗列技术名词，这种映射方式更有助于说明系统的实现逻辑，即不同组件并非并列堆叠，而是围绕同一处理流程形成彼此依赖的协作关系。",
    ),
    (
        "图 3.1 主要用户业务场景图",
        "图 3.1 直观展示了视频模式与摄像头模式两类核心业务场景的差异。前者强调稳定输出、结果保存与回放验证，后者则更关注实时显示、状态切换与交互连续性。正是由于两类场景在时延要求、控制重点和输出形态上存在明显区别，系统需求分析才不能停留在抽象功能罗列层面，而必须结合具体使用过程进行拆解。",
    ),
    (
        "表 3.1 功能需求汇总表",
        "表 3.1 对素材加载、处理启动、结果保存、历史回放与状态提示等核心功能进行了结构化整理。该表的价值在于把用户视角下的操作诉求转换为可实现的系统能力，并进一步明确这些功能之间并非彼此孤立，而是共同构成从输入素材到输出结果的连续业务链路，为后续详细设计提供约束依据。",
    ),
    (
        "表 3.2 非功能需求汇总表",
        "表 3.2 说明了界面可用性、处理稳定性、结构可维护性与扩展性等非功能目标。对于桌面多媒体原型而言，非功能需求虽然不直接产出换脸结果，却决定了系统能否在长时间运行、素材反复切换及多次演示过程中保持可控状态。因此，这部分内容实际上构成了工程实现质量的重要评价基础。",
    ),
    (
        "图 3.2 系统用例关系图",
        "图 3.2 从用户操作顺序的角度呈现了系统用例之间的依赖关系。通过该图可以看出，视频模式和摄像头模式虽然共享素材选择与结果查看等基础动作，但在前置条件和执行路径上存在清晰区分。借助这种关系梳理，后续界面状态控制、按钮启用逻辑以及线程清理策略便具备了更明确的设计依据。",
    ),
    (
        "表 3.3 系统可行性分析表",
        "表 3.3 将技术基础、工程条件与实现风险集中到同一框架下进行评估。通过该表可以看出，当前项目既具备成熟的技术栈支撑，也已经积累了数据库记录、媒体文件与运行样例等客观材料，因此总体方案具有较强的实施可行性。与此同时，路径治理和结果补登记等问题也被明确暴露出来，为后续设计与测试提供了现实约束。",
    ),
    (
        "图 4.1 系统总体架构图",
        "图 4.1 从分层角度揭示了用户交互层、业务控制层、媒体处理层和数据持久化层之间的关系。该图强调，系统的可用性并不是由某一个模块单独决定的，而是取决于界面逻辑、处理线程、底层引擎和数据记录是否能够稳定协同。对软件工程论文而言，这种架构表达有助于把分散代码归纳为边界清晰的整体系统。",
    ),
    (
        "图 4.2 数据流与接口流示意图",
        "图 4.2 进一步说明了素材、参数、处理结果与元数据在不同模块之间的流动路径。通过这一数据视角可以看到，前端并非直接承担全部逻辑，而是依靠后端接口和本地文件共同完成资源组织。这种设计既有助于减少界面层与数据层之间的耦合，也为后续增加查询、筛选与任务管理能力预留了扩展空间。",
    ),
    (
        "表 4.2 主要接口设计表",
        "表 4.2 对图片、视频、输出结果与任务等核心资源接口进行了汇总。其意义不仅在于说明接口地址或字段名称，更在于体现系统如何以统一的数据结构连接前端交互与后端存储。只有当接口输入输出保持稳定，素材列表展示、处理结果登记以及历史记录回放等功能才能在不同模块之间形成可重复、可验证的调用链路。",
    ),
    (
        "图 4.3 数据库关系示意图",
        "图 4.3 将 FaceImage、InputVideo、OutputVideo 与 ProcessingTask 之间的关联关系以图形方式呈现出来。相较于单独阅读字段表，关系示意图更能揭示结果记录如何回溯到源素材，以及后续任务管理如何在现有模型基础上继续扩展。对于需要强调结果可追踪性的系统而言，这种实体关系表达具有直接的工程解释价值。",
    ),
    (
        "图 4.4 系统部署结构图",
        "图 4.4 说明了当前系统采用的本机部署形态，即桌面前端、本地 Django 服务、SQLite 数据库与媒体目录共同运行在同一台设备上。这种部署方式虽然不追求生产级分布式能力，但非常契合毕业设计阶段对调试便利性、演示稳定性与环境可控性的要求，也使论文中的测试结果能够建立在明确而可复现的运行边界之上。",
    ),
    (
        "图 5.1 系统主界面截图",
        "图 5.1 展示了系统主界面在单窗口条件下整合素材选择、预览显示、参数控制与状态反馈的总体布局。该界面并未将功能拆散到多个独立窗口，而是通过区域化组织缩短用户的操作路径，使换脸处理、结果查看与参数调整能够在一次连续演示中完成。这种界面设计直接服务于答辩场景下的可用性与操作连贯性要求。",
    ),
    (
        "图 5.2 主界面局部功能区域截图",
        "图 5.2 进一步放大了主界面中的局部功能区域，便于观察模式切换、素材选择、状态提示与控制按钮之间的相互配合。通过该图可以更直观地理解，系统如何在同一前端中兼容视频模式与摄像头模式两类工作流，并通过控件显隐和状态更新降低用户误操作概率，这也是界面逻辑设计的重要体现。",
    ),
    (
        "表 5.1 前端素材加载与数据管理关键方法表",
        "表 5.1 从方法层面对素材加载与数据管理功能进行了拆解，使前端如何与后端接口、本地目录扫描及异步加载逻辑协同工作变得更加清晰。该表的作用在于把零散的函数名称转化为具有职责边界的实现单元，从而说明系统并非简单调用若干接口，而是围绕稳定加载、异常容忍与结果保存形成了相对完整的组织结构。",
    ),
    (
        "图 5.4 视频模式输入与输出样例",
        "图 5.4 将视频模式中的输入素材与输出结果并置展示，有助于把前述流程图中的抽象环节对应到真实样例上。对于系统实现型论文而言，这类样例的意义不只在于展示视觉效果，更在于证明前端选择、后台处理、结果写出与回放验证之间已经构成可执行的完整闭环，因此能够支撑后续测试与分析章节的展开。",
    ),
    (
        "图 5.5 摄像头模式处理流程图",
        "图 5.5 突出了摄像头模式在持续采集、实时显示与动态切换方面的特殊要求。与视频模式相比，摄像头模式并不以单次输出文件为目标，而更强调线程持续运行期间的界面响应、状态反馈与人脸切换控制。因此，该流程图不仅描述了处理步骤本身，也从侧面说明了为什么系统必须采用线程化设计来支撑实时交互。",
    ),
    (
        "表 5.2 输出结果归档与回放机制表",
        "表 5.2 说明了输出结果如何从处理完成状态进入元数据登记、列表展示与后续回放流程。相较于只在目录中生成视频文件，这种归档机制显著提升了系统的可追踪性与可复查性，使输出结果能够被重新加载、筛选和验证。对毕业设计而言，这一机制也是系统区别于临时算法 demo 的重要标志之一。",
    ),
    (
        "表 6.2 项目数据基础统计表",
        "表 6.2 将数据库记录数、本地可用路径数、输出目录文件数以及视频时长、帧率等信息进行了集中汇总。该表的作用在于把图 6.1 中的总体观察进一步量化，使测试章节不再停留于感性描述，而能够基于明确数据讨论系统当前的样本规模、资源一致性现状及后续分析的适用边界，从而提升论证的客观性。",
    ),
    (
        "表 6.3 测试方案与验收指标表",
        "表 6.3 对功能完整性、界面稳定性、数据一致性、结果可视化与异常处理等测试项目进行了条目化整理。通过这一表格，可以看出本文的测试并非随意操作若干功能点，而是围绕系统实际调用链建立了相对明确的验收标准。这样的组织方式有助于保证后续测试结果描述与前文需求、设计章节之间保持一致。",
    ),
    (
        "表 6.4 功能验证结果表",
        "表 6.4 归纳了各项核心功能在当前版本中的验证情况，使视频处理、摄像头演示、结果保存与历史回放等能力的完成度获得了更清晰的呈现。该表同样保留了异常容忍与同步不足等现实问题，从而避免测试章节只强调成功案例而忽略局限性。这种记录方式更符合工程验证的真实特点，也增强了论文结论的可信度。",
    ),
    (
        "表 6.5 稳定性与性能观察记录表",
        "表 6.5 并不追求算法 benchmark 式的高精度性能评测，而是围绕输入样本属性、处理耗时、回放表现与资源一致性等工程现象进行归纳。通过该表可以进一步看出，当前系统的主要风险集中在历史路径治理与结果补登记方面，而非基本流程完全失效。因此，这一表格既支撑了稳定性分析，也为后续改进建议提供了证据基础。",
    ),
    (
        "表 6.6 现存问题与改进建议表",
        "表 6.6 将路径失效、结果未补登记、任务调度尚未充分发挥作用等问题与对应改进方向进行对应整理。与正文分析相结合后，该表能够帮助读者快速把握系统当前最需要优先完善的工程环节，并理解这些问题并非孤立缺陷，而是与资源治理、任务管理和结果留痕机制密切相关，因而具有较强的后续实施价值。",
    ),
]

TAIL_INSERTIONS = [
    (
        "6 系统测试与结果分析",
        "表 5.3 从静态结构角度总结了界面控制类、处理线程、数据管理组件与底层引擎各自承担的任务。结合前面的时序图可以进一步看出，系统之所以能够在长耗时处理过程中保持基本响应，依赖的并不只是线程本身，而是线程职责划分、信号传递与状态清理机制共同作用的结果。这种静态与动态结合的说明，有助于增强全文论证的一致性。",
    ),
]


def strip_review_wrapper(text: str) -> str:
    lines = text.splitlines()
    if lines and lines[0].strip() == "修改后的正文":
        lines = lines[2:]
    if "修改简报" in "\n".join(lines):
        joined = "\n".join(lines)
        joined = joined[: joined.index("修改简报")].rstrip()
        lines = joined.splitlines()
    return "\n".join(lines).strip() + "\n"


def remove_bold_markers(text: str) -> str:
    return re.sub(r"^\*\*(.+?)\*\*$", r"\1", text, flags=re.M)


def normalize_spacing(text: str) -> str:
    text = text.replace("\u3000", "")
    text = re.sub(r"[ \t]+$", "", text, flags=re.M)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"([A-Za-z0-9\]])\s+([,;:.])", r"\1\2", text)
    text = re.sub(r"([,;:])([A-Za-z])", r"\1 \2", text)
    text = re.sub(r"(\bKeywords:)\s+", r"\1 ", text)
    return text


def insert_after_line(lines: list[str], anchor: str, paragraph: str) -> None:
    for index, line in enumerate(lines):
        if line.strip() == anchor:
            next_nonempty = ""
            for candidate in lines[index + 1 :]:
                if candidate.strip():
                    next_nonempty = candidate.strip()
                    break
            if next_nonempty == paragraph:
                return
            insert_at = index + 1
            while insert_at < len(lines) and lines[insert_at].strip() == "":
                insert_at += 1
            lines[insert_at:insert_at] = ["", paragraph, ""]
            return
    raise ValueError(f"Anchor not found in markdown: {anchor}")


def update_markdown() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")
    text = strip_review_wrapper(text)
    text = remove_bold_markers(text)
    lines = text.splitlines()
    for anchor, paragraph in INSERTIONS:
        insert_after_line(lines, anchor, paragraph)
    text = "\n".join(lines)
    text = normalize_spacing(text)
    OUTPUT_MD.write_text(text, encoding="utf-8")


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_style(paragraph: Paragraph, style_id: str = "Normal") -> None:
    paragraph.style = style_id
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.append(p_style)
    p_style.set(qn("w:val"), style_id)


def clear_bold(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.bold = False


def find_body_paragraph(doc: Document, text: str, start: int = 80) -> Paragraph:
    for paragraph in doc.paragraphs[start:]:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found in body: {text}")


def next_nonempty_paragraph(paragraph: Paragraph) -> Paragraph | None:
    current = paragraph._p.getnext()
    while current is not None:
        if current.tag != qn("w:p"):
            current = current.getnext()
            continue
        para = Paragraph(current, paragraph._parent)
        if (para.text or "").strip():
            return para
        current = current.getnext()
    return None


def insert_after_paragraph(anchor: Paragraph, text: str) -> Paragraph:
    next_para = next_nonempty_paragraph(anchor)
    if next_para is not None and next_para.text.strip() == text:
        set_style(next_para)
        clear_bold(next_para)
        return next_para

    new_para = paragraph_after(anchor)
    set_style(new_para)
    run = new_para.add_run(text)
    run.bold = False
    return new_para


def previous_nonempty_paragraph(paragraph: Paragraph) -> Paragraph | None:
    current = paragraph._p.getprevious()
    while current is not None:
        if current.tag != qn("w:p"):
            current = current.getprevious()
            continue
        para = Paragraph(current, paragraph._parent)
        if (para.text or "").strip():
            return para
        current = current.getprevious()
    return None


def insert_before_paragraph(anchor: Paragraph, text: str) -> Paragraph:
    prev_para = previous_nonempty_paragraph(anchor)
    if prev_para is not None and prev_para.text.strip() == text:
        set_style(prev_para)
        clear_bold(prev_para)
        return prev_para

    new_para = paragraph_before(anchor)
    set_style(new_para)
    run = new_para.add_run(text)
    run.bold = False
    return new_para


def normalize_docx(doc: Document) -> None:
    for paragraph in doc.paragraphs[80:]:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name == "Normal" and len(text) > 70:
            clear_bold(paragraph)


def update_docx() -> None:
    doc = Document(SOURCE_DOCX)
    normalize_docx(doc)
    for anchor_text, paragraph_text in INSERTIONS:
        anchor_para = find_body_paragraph(doc, anchor_text)
        insert_after_paragraph(anchor_para, paragraph_text)
    for anchor_text, paragraph_text in TAIL_INSERTIONS:
        anchor_para = find_body_paragraph(doc, anchor_text)
        insert_before_paragraph(anchor_para, paragraph_text)
    doc.save(OUTPUT_DOCX)


def main() -> None:
    update_markdown()
    update_docx()
    print(f"output_md={OUTPUT_MD}")
    print(f"output_docx={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
