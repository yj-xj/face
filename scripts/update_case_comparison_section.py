from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches


BASE_DIR = Path(r"e:\face")
DOCX_PATH = BASE_DIR / "换脸系统_初稿_学术化版_精修后_已加入方法对比.docx"
MD_PATH = BASE_DIR / "换脸系统_初稿_学术化版_精修后.md"
OUTPUT_DOCX_PATH = BASE_DIR / "换脸系统_初稿_学术化版_精修后_图文整改版.docx"
ASSET_DIR = BASE_DIR / "thesis_assets_generated"

SECTION_START = "6.5 典型案例分析"
SECTION_END = "6.6 现存问题与改进建议"

REAL_FIGURE = ASSET_DIR / "comparison_real_source.png"
PAINTED_FIGURE = ASSET_DIR / "comparison_painted_source.png"

CAMERA_FIGURE = BASE_DIR / "image" / "摄像头模式.png"

FIGURE_REAL_CAPTION = "图 6.3 真实照片源人脸条件下的代表性结果对比图"
FIGURE_REAL_ANALYSIS = (
    "图 6.3 选取真实照片作为源人脸，分别展示传统方法与 InsightFace 的代表性输出。"
    "从图中可以看出，传统方法能够完成基本的人脸替换，但在人脸边缘衔接、亮度过渡与整体融合感方面仍保留较明显的拼接痕迹；"
    "InsightFace 输出的人脸轮廓更完整，五官比例与肤色过渡也更接近自然照片。"
    "需要说明的是，当前图像材料来源于不同测试环节，因此该图更适合用于比较两类方案的视觉特征，而不构成严格控制变量的算法 benchmark。"
)

FIGURE_PAINTED_CAPTION = "图 6.4 绘制源人脸条件下的代表性结果对比图"
FIGURE_PAINTED_ANALYSIS = (
    "图 6.4 采用绘制人物图像作为源人脸。"
    "在此情形下，传统方法虽然能够完成几何对齐，但由于绘制纹理与真实照片之间存在显著域差异，输出结果容易出现局部覆盖感较强、纹理衔接不连续等问题；"
    "InsightFace 在该场景下仍能保持一定的人脸结构连续性，但受绘制风格与真实照片统计分布差异影响，其结果自然度仍弱于真实照片之间的换脸效果。"
    "这一现象说明，绘制人脸向真实照片的跨域迁移比真实照片之间的换脸任务更具挑战性。"
)

FIGURE_CAMERA_CAPTION = "图 6.5 摄像头模式运行界面"
FIGURE_CAMERA_ANALYSIS = (
    "除静态图像结果外，图 6.5 进一步展示了系统在摄像头模式下的实际运行状态。"
    "界面能够同时完成待替换人脸选择、摄像头启停、快照保存与实时显示，说明本系统并非仅能离线生成结果图，而是已经具备较完整的交互式演示能力。"
    "结合图 6.3 与图 6.4 的结果可以认为，系统已经能够在不同素材类型与不同运行模式下形成可观察、可分析的输出结果。"
)

SECTION_PARAGRAPHS = [
    "为了使测试分析不局限于统计表与文字说明，本文结合 image 目录中新增的样例图像，对传统方法与 InsightFace 的输出结果进行了重新整理，并生成两组对比图用于说明系统在不同素材条件下的表现差异。需要指出的是，当前图像材料分别来自静态图像处理、历史输出样例与摄像头运行界面，因此本节重点关注工程可用性与视觉表现特征，而不将其视为严格控制变量的数值对照实验。",
]


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        (Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc")),
        (Path(r"C:\Windows\Fonts\simhei.ttf")),
        (Path(r"C:\Windows\Fonts\simsun.ttc")),
    ]
    for path in candidates:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def crop_center_square(image: Image.Image) -> Image.Image:
    image = image.convert("RGB")
    w, h = image.size
    side = min(w, h)
    left = (w - side) // 2
    top = (h - side) // 2
    return image.crop((left, top, left + side, top + side))


def build_comparison_figure(output_path: Path, title: str, items: list[tuple[Path, str]]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    canvas_w = 1800
    margin = 70
    title_h = 110
    cell_gap = 36
    cell_w = (canvas_w - margin * 2 - cell_gap * (len(items) - 1)) // len(items)
    cell_h = 620
    label_h = 64
    image_box_h = cell_h - label_h - 26
    canvas_h = margin * 2 + title_h + cell_h

    canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = get_font(34, bold=True)
    label_font = get_font(24, bold=True)

    draw.text((margin, 28), title, fill=(30, 41, 59), font=title_font)
    y = margin + title_h - 20

    for index, (image_path, label) in enumerate(items):
        x = margin + index * (cell_w + cell_gap)
        draw.rounded_rectangle((x, y, x + cell_w, y + cell_h), radius=20, outline=(203, 213, 225), width=3, fill=(248, 250, 252))
        draw.rounded_rectangle((x + 12, y + 12, x + cell_w - 12, y + 12 + label_h), radius=14, fill=(226, 232, 240))
        bbox = draw.textbbox((0, 0), label, font=label_font)
        text_x = x + (cell_w - (bbox[2] - bbox[0])) / 2
        text_y = y + 12 + (label_h - (bbox[3] - bbox[1])) / 2 - 2
        draw.text((text_x, text_y), label, fill=(15, 23, 42), font=label_font)

        with Image.open(image_path) as im:
            prepared = crop_center_square(im)
            prepared = ImageOps.contain(prepared, (cell_w - 34, image_box_h - 30))

        panel = Image.new("RGB", (cell_w - 24, image_box_h), "white")
        panel_draw = ImageDraw.Draw(panel)
        panel_draw.rounded_rectangle((0, 0, panel.width - 1, panel.height - 1), radius=18, outline=(226, 232, 240), width=2, fill="white")
        px = (panel.width - prepared.width) // 2
        py = (panel.height - prepared.height) // 2
        panel.paste(prepared, (px, py))

        canvas.paste(panel, (x + 12, y + 12 + label_h + 14))

    canvas.save(output_path)


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    return new_para


def set_paragraph_style(paragraph: Paragraph, style_id: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.append(p_style)
    p_style.set(qn("w:val"), style_id)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        paragraph._p = paragraph._element = None


def clear_between(doc: Document, start_text: str, end_text: str) -> tuple[Paragraph, Paragraph]:
    start_para = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == start_text)
    end_para = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == end_text)

    current = start_para._p.getnext()
    while current is not None and current is not end_para._p:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    return start_para, end_para


def add_text_paragraph(anchor: Paragraph, text: str, *, bold: bool = False, center: bool = False) -> Paragraph:
    new_para = paragraph_after(anchor)
    new_para.style = "Normal"
    set_paragraph_style(new_para, "Normal")
    run = new_para.add_run(text)
    run.bold = bold
    if center:
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_para.paragraph_format.line_spacing = 1.0
    return new_para


def add_picture_paragraph(anchor: Paragraph, image_path: Path, width_inches: float = 6.1) -> Paragraph:
    new_para = paragraph_after(anchor)
    new_para.style = "Normal"
    set_paragraph_style(new_para, "Normal")
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return new_para


def build_markdown_section() -> str:
    parts = [
        "6.5 典型案例分析",
        "",
    ]
    for paragraph in SECTION_PARAGRAPHS:
        parts.extend([paragraph, ""])
    parts.extend(
        [
            FIGURE_REAL_CAPTION,
            "",
            f"**{FIGURE_REAL_ANALYSIS}**",
            "",
            FIGURE_PAINTED_CAPTION,
            "",
            f"**{FIGURE_PAINTED_ANALYSIS}**",
            "",
            FIGURE_CAMERA_CAPTION,
            "",
            f"**{FIGURE_CAMERA_ANALYSIS}**",
            "",
        ]
    )
    return "\n".join(parts).strip()


def update_markdown() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    start = text.index("6.5 典型案例分析")
    end = text.index("6.6 现存问题与改进建议")
    replacement = build_markdown_section() + "\n\n"
    new_text = text[:start] + replacement + text[end:]
    MD_PATH.write_text(new_text, encoding="utf-8")


def update_docx() -> None:
    doc = Document(DOCX_PATH)
    start_para, _ = clear_between(doc, SECTION_START, SECTION_END)

    current = start_para
    for paragraph in SECTION_PARAGRAPHS:
        current = add_text_paragraph(current, paragraph, bold=False, center=False)

    current = add_picture_paragraph(current, REAL_FIGURE)
    current = add_text_paragraph(current, FIGURE_REAL_CAPTION, bold=True, center=True)
    current = add_text_paragraph(current, FIGURE_REAL_ANALYSIS, bold=True, center=False)

    current = add_picture_paragraph(current, PAINTED_FIGURE)
    current = add_text_paragraph(current, FIGURE_PAINTED_CAPTION, bold=True, center=True)
    current = add_text_paragraph(current, FIGURE_PAINTED_ANALYSIS, bold=True, center=False)

    current = add_picture_paragraph(current, CAMERA_FIGURE, width_inches=6.3)
    current = add_text_paragraph(current, FIGURE_CAMERA_CAPTION, bold=True, center=True)
    current = add_text_paragraph(current, FIGURE_CAMERA_ANALYSIS, bold=True, center=False)

    doc.save(OUTPUT_DOCX_PATH)


def main() -> None:
    build_comparison_figure(
        REAL_FIGURE,
        "真实照片源人脸条件下的代表性结果对比",
        [
            (BASE_DIR / "image" / "face_origin1.jpg", "源人脸"),
            (BASE_DIR / "image" / "传统_face_swap1.png", "传统方法结果"),
            (BASE_DIR / "image" / "face_swap1.png", "InsightFace 结果"),
        ],
    )
    build_comparison_figure(
        PAINTED_FIGURE,
        "绘制源人脸条件下的代表性结果对比",
        [
            (BASE_DIR / "image" / "face_origin2.jpg", "源人脸"),
            (BASE_DIR / "image" / "传统_face_swap2.png", "传统方法结果"),
            (BASE_DIR / "image" / "face_swap2.png", "InsightFace 结果"),
        ],
    )
    update_markdown()
    update_docx()
    print(f"real_figure={REAL_FIGURE}")
    print(f"painted_figure={PAINTED_FIGURE}")
    print(f"output_docx={OUTPUT_DOCX_PATH}")
    print(f"output_md={MD_PATH}")


if __name__ == "__main__":
    main()
