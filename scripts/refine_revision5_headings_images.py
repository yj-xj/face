from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
BLACK = RGBColor(0, 0, 0)
SPECIAL_TITLES = {"ABSTRACT", "参考文献", "致谢"}


def find_source() -> Path:
    matches = [p for p in ROOT.glob("*修订版5.docx") if not p.name.startswith("~$")]
    if not matches:
        raise FileNotFoundError("Could not find the non-temp revision 5 docx.")
    return matches[0]


def build_output_path(source: Path) -> Path:
    match = re.search(r"修订版(\d+)$", source.stem)
    if match:
        next_rev = int(match.group(1)) + 1
        base_stem = source.stem[: match.start()] + f"修订版{next_rev}"
        candidate = source.with_name(base_stem + source.suffix)
        while candidate.exists():
            next_rev += 1
            candidate = source.with_name(source.stem[: match.start()] + f"修订版{next_rev}" + source.suffix)
        return candidate
    candidate = source.with_name(source.stem + "_版式微调版" + source.suffix)
    counter = 1
    while candidate.exists():
        candidate = source.with_name(source.stem + f"_版式微调版{counter}" + source.suffix)
        counter += 1
    return candidate


def normalized_title(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def is_special_title(text: str) -> bool:
    compact = normalized_title(text)
    return compact in {"摘要", "目录"} or text.strip() in SPECIAL_TITLES


def set_paragraph_runs_black(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = BLACK


def apply_heading_layout(doc: DocumentObject) -> None:
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style_name = paragraph.style.name if paragraph.style else ""

        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_paragraph_runs_black(paragraph)
            continue

        if style_name == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_paragraph_runs_black(paragraph)
            continue

        if style_name == "Heading 3":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_paragraph_runs_black(paragraph)
            continue

        if is_special_title(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_paragraph_runs_black(paragraph)


def find_paragraph_by_prefix(doc: DocumentObject, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text.startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found for prefix: {prefix}")


def insert_heading3_once(anchor: Paragraph, heading_text: str) -> None:
    prev = anchor._p.getprevious()
    if prev is not None:
        prev_para = Paragraph(prev, anchor._parent)
        if (prev_para.text or "").strip() == heading_text:
            return

    inserted = anchor.insert_paragraph_before(heading_text, style="Heading 3")
    inserted.alignment = WD_ALIGN_PARAGRAPH.LEFT
    inserted.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_runs_black(inserted)


def add_third_level_headings(doc: DocumentObject) -> None:
    insertions = [
        (
            "在本系统的传统处理链中，人脸替换的核心并不是神经网络推理",
            "2.2.1 传统处理链的实现入口",
        ),
        (
            "从具体算法机制看，三角剖分的意义在于避免对整张人脸执行单一刚性变换",
            "2.2.2 三角剖分与局部仿射映射",
        ),
        (
            "在逐三角形映射完成之后，系统还需要对变形后的人脸纹理进行区域回填与边界融合",
            "2.2.3 边界融合与效果补偿",
        ),
        (
            "从实际代码看，系统中的深度学习换脸并不是一个泛泛的概念",
            "2.3.1 深度学习处理链的组成",
        ),
        (
            "基于深度学习的人脸替换可抽象为身份特征提取与目标属性保持的联合映射过程",
            "2.3.2 身份表征与生成机制",
        ),
        (
            "从模型优化目标看，此类方法通常将总损失写为",
            "2.3.3 优化目标与工程含义",
        ),
        (
            "从代码结构上看，系统的并行协作主要围绕三类对象展开",
            "5.7.1 核心对象分工",
        ),
        (
            "这种分工方式的优势在于，当算法处理或媒体读写出现延迟时",
            "5.7.2 协作机制与运行价值",
        ),
        (
            "本章测试的定位是工程验证，而非算法竞赛式 benchmark。",
            "6.1.1 测试定位与验证重点",
        ),
        (
            "测试环境延续了项目开发阶段的单机模式。",
            "6.1.2 环境组成与运行条件",
        ),
        (
            "截至本文撰写时，系统数据库中已经记录 19 条人脸图片",
            "6.1.3 样本基础与统计范围",
        ),
    ]

    for prefix, heading_text in insertions:
        anchor = find_paragraph_by_prefix(doc, prefix)
        insert_heading3_once(anchor, heading_text)


def clear_fixed_line_spacing(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.attrib.pop(qn("w:line"), None)
        spacing.attrib.pop(qn("w:lineRule"), None)


def normalize_image_paragraphs(doc: DocumentObject) -> int:
    image_paragraphs = 0
    for paragraph in doc.paragraphs:
        if not paragraph._p.xpath(".//w:drawing"):
            continue
        image_paragraphs += 1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        clear_fixed_line_spacing(paragraph)
    return image_paragraphs


def main() -> None:
    source = find_source()
    output = build_output_path(source)
    shutil.copy2(source, output)

    doc = Document(str(output))
    apply_heading_layout(doc)
    add_third_level_headings(doc)
    image_count = normalize_image_paragraphs(doc)
    doc.save(str(output))

    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    print(f"IMAGES_FIXED={image_count}")


if __name__ == "__main__":
    main()
