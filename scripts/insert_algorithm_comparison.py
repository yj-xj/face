from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"e:\face")
DOCX_PATH = BASE_DIR / "换脸系统_初稿_学术化版_精修后.docx"
FALLBACK_DOCX_PATH = BASE_DIR / "换脸系统_初稿_学术化版_精修后_已加入方法对比.docx"
MD_PATH = BASE_DIR / "换脸系统_初稿_学术化版_精修后.md"

SECTION_START = "6.5 典型案例分析"
SECTION_END = "6.6 现存问题与改进建议"

COMPARISON_PARAGRAPHS = [
    "在已有静态样例与视频结果展示的基础上，本文进一步补充传统方法与深度学习方法在不同素材条件下的对比分析。考虑到当前系统同时保留了传统几何配准方案与 InsightFace 方案，且测试材料既包含绘制类人脸，也包含真实照片类人脸，因此将比较场景划分为三类：绘制人脸与真实照片人脸之间的跨域换脸、绘制人脸之间的同域换脸，以及真实照片人脸之间的同域换脸。比较重点并不在于给出统一的算法评分值，而在于考察两类方法在人脸身份保持、边界自然度、纹理协调性与结果稳定性等方面的表现差异。",
    "在绘制人脸与真实照片人脸的跨域场景中，两类方法均面临较为明显的域差异问题。传统方法主要依赖关键点对齐与局部区域融合，当绘制素材的轮廓特征较为清晰时，可以完成基本的人脸映射，但在颜色过渡、纹理衔接与光照一致性方面较易保留拼接痕迹。相比之下，InsightFace 在真实照片目标上的身份迁移能力更强，整体结果通常更接近自然视觉效果；但由于绘制素材与模型训练分布之间存在偏移，部分样例中仍可能出现特征提取不稳定或风格协调性不足的现象。由此可见，跨域换脸是三类情形中难度最高的一类，深度学习方法虽整体优于传统方法，但并不能完全消除风格差异带来的影响。",
    "在绘制人脸之间的换脸场景中，传统方法的处理链路相对直接。当两幅素材在朝向、比例与五官布局上较为接近时，该方法能够较稳定地完成区域替换，并在一定程度上保留原有插画风格。然而，由于其缺少对高层语义结构的建模能力，当素材之间在线条风格、表情幅度或局部细节上存在明显差异时，输出结果仍容易出现边缘衔接不自然或局部结构失真的问题。InsightFace 在该类场景中的表现则更依赖素材是否保留较强的真实人脸结构特征。若绘制风格偏向写实，其结果通常较为自然；若素材风格高度抽象化，深度特征匹配的稳定性会明显下降，其优势未必始终显著。因此，在纯绘制素材条件下，传统方法并非完全失去意义，反而可作为具有较强可解释性的补充方案。",
    "在真实照片人脸之间的换脸场景中，深度学习方法的优势最为明显。InsightFace 在身份特征保持、人脸轮廓过渡、表情适配以及整体真实感方面普遍优于传统方法，所得结果更适合直接用于视频展示与系统演示。传统方法虽然能够完成基本对齐与替换，但在姿态变化、光照差异或局部遮挡较为明显时，更容易产生边界僵硬、颜色不一致以及局部结构拉伸等问题。从当前系统的工程目标来看，若需要获得质量更高、可直接回放的换脸结果，真实照片之间的换脸任务更适合作为深度学习方法的主要应用场景。",
    "综合上述三类测试情形可以看出，传统方法与深度学习方法之间并非简单的替代关系，而更接近主辅并存的工程组合。传统方法实现链路清晰、运行依赖较少、便于定位问题，适合作为兼容性方案或特定素材条件下的补充路径；深度学习方法则在真实人脸换脸和复杂场景适应方面表现更优，应作为当前系统的主要输出方案。基于这一认识，本文在系统实现中保留了两类方法并行存在的设计，这不仅有助于提升系统的适用范围，也更符合原型系统在演示、验证与工程分析阶段的实际需求。",
]


def insert_paragraph_after(paragraph: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    if template._p.pPr is not None:
        new_p.append(deepcopy(template._p.pPr))
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = template.style

    run = new_para.add_run(text)
    template_run = next((candidate for candidate in template.runs if candidate._r.rPr is not None), None)
    if template_run is not None:
        run._r.insert(0, deepcopy(template_run._r.rPr))
    run.bold = False
    return new_para


def find_paragraph_index(doc: Document, text: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"Paragraph not found: {text}")


def update_docx() -> tuple[int, Path]:
    doc = Document(DOCX_PATH)
    start_index = find_paragraph_index(doc, SECTION_START)
    end_index = find_paragraph_index(doc, SECTION_END)
    section_paragraphs = doc.paragraphs[start_index + 1 : end_index]
    non_empty_texts = [paragraph.text.strip() for paragraph in section_paragraphs if paragraph.text.strip()]

    if COMPARISON_PARAGRAPHS[0] in non_empty_texts:
        return 0, DOCX_PATH

    template = next(
        paragraph
        for paragraph in section_paragraphs
        if paragraph.text.strip() and paragraph.style and paragraph.style.name == "Normal" and not any(bool(run.bold) for run in paragraph.runs if run.text.strip())
    )
    anchor = next(paragraph for paragraph in reversed(section_paragraphs) if paragraph.text.strip())

    current = anchor
    for paragraph_text in COMPARISON_PARAGRAPHS:
        current = insert_paragraph_after(current, paragraph_text, template)

    try:
        doc.save(DOCX_PATH)
        return len(COMPARISON_PARAGRAPHS), DOCX_PATH
    except PermissionError:
        doc.save(FALLBACK_DOCX_PATH)
        return len(COMPARISON_PARAGRAPHS), FALLBACK_DOCX_PATH


def update_markdown() -> int:
    text = MD_PATH.read_text(encoding="utf-8")
    if COMPARISON_PARAGRAPHS[0] in text:
        return 0

    marker = "\n\n6.6 现存问题与改进建议"
    if marker not in text:
        raise ValueError("Markdown section marker not found.")

    block = "\n\n".join(COMPARISON_PARAGRAPHS)
    text = text.replace(marker, f"\n\n{block}{marker}", 1)
    MD_PATH.write_text(text, encoding="utf-8")
    return len(COMPARISON_PARAGRAPHS)


def main() -> None:
    docx_count, docx_path = update_docx()
    md_count = update_markdown()
    print(f"docx_inserted={docx_count}")
    print(f"md_inserted={md_count}")
    print(f"docx_path={docx_path}")
    print(f"md_path={MD_PATH}")


if __name__ == "__main__":
    main()
