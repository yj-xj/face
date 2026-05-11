from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
SOURCE_PATTERN = "*修订版8.docx"

NEW_STRUCTURE_PARAGRAPH = (
    "全文共由六章构成。第一章介绍研究背景、研究意义、国内外研究现状及本文主要工作；"
    "第二章说明与系统实现密切相关的关键技术和理论基础；第三章面向业务场景展开系统需求分析；"
    "第四章围绕系统设计与实现展开论述，先说明系统架构、模块划分、数据流、数据库和部署方式，"
    "再介绍界面组织、素材管理、处理流程与线程协作等实现内容；第五章从工程测试视角对系统进行验证与分析；"
    "第六章总结全文并提出后续展望。"
)

NEW_CHAPTER4_INTRO = (
    "本章将系统总体设计与详细实现合并展开。内容先从架构分层、模块职责、接口组织、数据库关系和部署方式说明整体设计，"
    "再结合前端界面、双模式交互、素材管理、处理流程与线程协作分析关键实现，以呈现系统从方案设计到落地运行的完整链路。"
)

NEW_SECTION_4_1_INTRO = (
    "本节在需求分析的基础上，对当前项目的整体结构做工程化整理。后文将依次说明设计目标、分层架构、模块职责、接口组织、"
    "数据库关系和本机部署形态，把前端、后端与脚本中的实现内容归纳为边界清楚、职责明确的系统方案。"
)

NEW_SECTION_4_2_INTRO = (
    "本节聚焦系统已经落地的关键实现，而不是停留在设计层描述。内容将从前端主界面、双模式交互、素材加载、视频处理、"
    "摄像头处理、结果归档以及线程协作几个方面展开，说明现有代码如何把换脸能力组织成一套能够连续运行、能够回放结果、"
    "也能够保留处理记录的桌面原型。"
)

TEXT_REPLACEMENTS = {
    "全文共由七章构成。第一章介绍研究背景、研究意义、国内外研究现状及本文主要工作；第二章说明与系统实现密切相关的关键技术和理论基础；第三章面向业务场景展开系统需求分析；第四章从系统架构、模块划分、数据流、数据库和部署方式等角度给出总体设计；第五章重点说明系统的详细实现过程；第六章从工程测试视角对系统进行验证与分析；第七章总结全文并提出后续展望。": NEW_STRUCTURE_PARAGRAPH,
    "图5-1": "图4-5",
    "图5-2": "图4-6",
    "图5-3": "图4-7",
    "图5-4": "图4-8",
    "图5-5": "图4-9",
    "图5-6": "图4-10",
    "图5-7": "图4-11",
    "表5-1": "表4-4",
    "表5-2": "表4-5",
    "表5-3": "表4-6",
    "图6-1": "图5-1",
    "图6-3": "图5-2",
    "图6-4": "图5-3",
    "表6-1": "表5-1",
    "表6-2": "表5-2",
    "表6-3": "表5-3",
    "表6-4": "表5-4",
    "表6-5": "表5-5",
    "图 5.1": "图 4.5",
    "图 5.2": "图 4.6",
    "图 5.3": "图 4.7",
    "图 5.4": "图 4.8",
    "图 5.5": "图 4.9",
    "图 5.6": "图 4.10",
    "图 5.7": "图 4.11",
    "表 5.1": "表 4.4",
    "表 5.2": "表 4.5",
    "表 5.3": "表 4.6",
    "图 6.1": "图 5.1",
    "图 6.3": "图 5.2",
    "图 6.4": "图 5.3",
    "表 6.1": "表 5.1",
    "表 6.2": "表 5.2",
    "表 6.3": "表 5.3",
    "表 6.4": "表 5.4",
    "表 6.5": "表 5.5",
}

FINAL_FIGURE_EXPLANATIONS = {
    "图3-1 主要用户业务场景图": "该图对应系统的两类核心使用场景，即离线视频处理与摄像头实时演示。",
    "图3-2 系统用例关系图": "该图展示了素材选择、处理执行、结果查看等用例之间的主要依赖关系。",
    "图4-1 系统总体架构图": "该图概括了界面层、控制层、处理层与数据层之间的总体分层关系。",
    "图4-2 数据流与接口流示意图": "该图展示了素材、参数与结果在前端、后端和本地文件之间的流转路径。",
    "图4-3 数据库关系示意图": "该图说明了人脸图片、输入视频、输出结果与处理任务之间的关联关系。",
    "图4-4 系统部署结构图": "该图展示了桌面前端、本地 Django 服务、SQLite 数据库与媒体目录的本机部署方式。",
    "图4-5 系统主界面截图": "该图展示了主界面的整体布局，以及素材区、预览区和控制区的组织方式。",
    "图4-6 主界面局部功能区域截图": "该图展示了模式切换、素材选择和状态提示等局部功能区域的配合关系。",
    "图4-7 视频模式处理流程图": "该图对应视频模式从参数准备到结果保存与回放的主要处理流程。",
    "图4-8 视频模式输入与输出样例": "该图并列展示了视频模式的输入素材与输出结果。",
    "图4-9 摄像头模式处理流程图": "该图展示了摄像头模式下实时采集、换脸处理与界面反馈的基本流程。",
    "图4-10 摄像头模式运行界面截图": "该图展示了摄像头模式运行时的实时画面和控制区域布局。",
    "图4-11 关键线程与模块协作时序图": "该图展示了界面控制类、处理线程、底层引擎与数据管理组件之间的协作时序。",
    "图5-1 系统样本数据统计图": "该图统计了当前系统中人脸图片、输入视频和输出视频的数量分布。",
    "图5-2 真实照片源人脸条件下的代表性结果对比图": "该图给出了真实照片源人脸条件下的代表性处理结果。",
    "图5-3 绘制源人脸条件下的代表性结果对比图": "该图给出了绘制源人脸条件下的代表性处理结果。",
}

HEADING_REWRITES = {
    "第四章 系统总体设计": ("第四章 系统设计与实现", "Heading 1"),
    "4.1 设计目标与原则": ("4.1.1 设计目标与原则", "Heading 3"),
    "4.2 总体架构设计": ("4.1.2 总体架构设计", "Heading 3"),
    "4.3 模块划分与职责": ("4.1.3 模块划分与职责", "Heading 3"),
    "4.4 接口设计": ("4.1.4 接口设计", "Heading 3"),
    "4.5 数据库设计": ("4.1.5 数据库设计", "Heading 3"),
    "4.6 部署结构设计": ("4.1.6 部署结构设计", "Heading 3"),
    "第五章 系统详细设计与实现": ("4.2 系统详细设计与实现", "Heading 2"),
    "5.1 前端主界面实现": ("4.2.1 前端主界面实现", "Heading 3"),
    "5.2 双模式界面组织与交互逻辑": ("4.2.2 双模式界面组织与交互逻辑", "Heading 3"),
    "5.3 素材加载与数据库管理实现": ("4.2.3 素材加载与数据库管理实现", "Heading 3"),
    "5.4 视频模式处理流程实现": ("4.2.4 视频模式处理流程实现", "Heading 3"),
    "5.5 摄像头模式处理流程实现": ("4.2.5 摄像头模式处理流程实现", "Heading 3"),
    "5.6 输出归档与结果回放实现": ("4.2.6 输出归档与结果回放实现", "Heading 3"),
    "5.7 关键线程与类协作实现": ("4.2.7 关键线程与类协作实现", "Heading 3"),
    "第六章 系统测试与结果分析": ("第五章 系统测试与结果分析", "Heading 1"),
    "6.1 测试环境与测试目标": ("5.1 测试环境与测试目标", "Heading 2"),
    "6.1.1 测试定位与验证重点": ("5.1.1 测试定位与验证重点", "Heading 3"),
    "6.1.2 环境组成与运行条件": ("5.1.2 环境组成与运行条件", "Heading 3"),
    "6.1.3 样本基础与统计范围": ("5.1.3 样本基础与统计范围", "Heading 3"),
    "6.2 测试方案设计": ("5.2 测试方案设计", "Heading 2"),
    "6.3 功能测试与验证结果": ("5.3 功能测试与验证结果", "Heading 2"),
    "6.4 稳定性与性能观察": ("5.4 稳定性与性能观察", "Heading 2"),
    "6.5 典型案例分析": ("5.5 典型案例分析", "Heading 2"),
    "第七章 总结与展望": ("第六章 总结与展望", "Heading 1"),
    "7.1 工作总结": ("6.1 工作总结", "Heading 2"),
    "7.2 后续展望": ("6.2 后续展望", "Heading 2"),
}


def find_source() -> Path:
    matches = sorted(p for p in ROOT.glob(SOURCE_PATTERN) if not p.name.startswith("~$"))
    if not matches:
        raise FileNotFoundError("未找到修订版8文档。")
    return matches[0]


def build_output_path(source: Path) -> Path:
    match = re.search(r"修订版(\d+)$", source.stem)
    if match:
        next_rev = int(match.group(1)) + 1
        prefix = source.stem[: match.start()]
        candidate = source.with_name(f"{prefix}修订版{next_rev}{source.suffix}")
        while candidate.exists():
            next_rev += 1
            candidate = source.with_name(f"{prefix}修订版{next_rev}{source.suffix}")
        return candidate
    counter = 9
    candidate = source.with_name(f"{source.stem}_修订版{counter}{source.suffix}")
    while candidate.exists():
        counter += 1
        candidate = source.with_name(f"{source.stem}_修订版{counter}{source.suffix}")
    return candidate


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def paragraph_has_math(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))


def find_paragraph(doc: DocumentObject, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if (paragraph.text or "").strip() == exact_text:
            return paragraph
    raise ValueError(f"未找到段落: {exact_text}")


def paragraph_index(doc: DocumentObject, target: Paragraph) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError("未找到目标段落索引。")


def copy_run_format(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    dst_run.font.color.rgb = src_run.font.color.rgb

    src_rpr = src_run._element.rPr
    if src_rpr is None:
        return
    src_fonts = src_rpr.find(qn("w:rFonts"))
    if src_fonts is None:
        return

    dst_rpr = dst_run._element.get_or_add_rPr()
    dst_fonts = dst_rpr.find(qn("w:rFonts"))
    if dst_fonts is None:
        dst_fonts = OxmlElement("w:rFonts")
        dst_rpr.append(dst_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        value = src_fonts.get(qn(key))
        if value:
            dst_fonts.set(qn(key), value)


def copy_paragraph_format(src: Paragraph, dst: Paragraph) -> None:
    dst.alignment = src.alignment
    src_pf = src.paragraph_format
    dst_pf = dst.paragraph_format

    attrs = (
        "left_indent",
        "right_indent",
        "first_line_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacing_rule",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
    )
    for attr in attrs:
        try:
            setattr(dst_pf, attr, getattr(src_pf, attr))
        except Exception:
            continue


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    template_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.text = new_text
    if template_run and paragraph.runs:
        copy_run_format(template_run, paragraph.runs[0])


def insert_paragraph_after(anchor: Paragraph, text: str, style_name: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.style = style_name
    if text:
        paragraph.add_run(text)
    return paragraph


def apply_heading_layout(paragraph: Paragraph, style_name: str) -> None:
    paragraph.style = style_name
    paragraph.paragraph_format.first_line_indent = Pt(0)
    if style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.page_break_before = True
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.page_break_before = False


def rewrite_headings(doc: DocumentObject) -> None:
    for old_text, (new_text, style_name) in HEADING_REWRITES.items():
        paragraph = find_paragraph(doc, old_text)
        replace_paragraph_text(paragraph, new_text)
        apply_heading_layout(paragraph, style_name)


def add_merged_chapter_structure(doc: DocumentObject) -> None:
    chapter4 = find_paragraph(doc, "第四章 系统总体设计")
    intro4 = doc.paragraphs[paragraph_index(doc, chapter4) + 1]
    first_subheading = find_paragraph(doc, "4.1 设计目标与原则")
    chapter5 = find_paragraph(doc, "第五章 系统详细设计与实现")
    intro5 = doc.paragraphs[paragraph_index(doc, chapter5) + 1]

    chapter_intro = insert_paragraph_after(chapter4, NEW_CHAPTER4_INTRO, intro4.style.name)
    copy_paragraph_format(intro4, chapter_intro)
    if intro4.runs and chapter_intro.runs:
        copy_run_format(intro4.runs[0], chapter_intro.runs[0])

    section_4_1 = insert_paragraph_after(chapter_intro, "4.1 系统总体设计", first_subheading.style.name)
    copy_paragraph_format(first_subheading, section_4_1)
    if first_subheading.runs and section_4_1.runs:
        copy_run_format(first_subheading.runs[0], section_4_1.runs[0])
    apply_heading_layout(section_4_1, "Heading 2")

    replace_paragraph_text(intro4, NEW_SECTION_4_1_INTRO)
    replace_paragraph_text(intro5, NEW_SECTION_4_2_INTRO)


def apply_text_replacements(doc: DocumentObject) -> int:
    if not TEXT_REPLACEMENTS:
        return 0
    pattern = re.compile("|".join(re.escape(key) for key in sorted(TEXT_REPLACEMENTS, key=len, reverse=True)))
    changed = 0

    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph) or paragraph_has_math(paragraph):
            continue
        text = paragraph.text or ""
        if not text:
            continue
        new_text = pattern.sub(lambda match: TEXT_REPLACEMENTS[match.group(0)], text)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)
            changed += 1
    return changed


def shorten_figure_explanations(doc: DocumentObject) -> int:
    changed = 0
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs[:-1]):
        caption = (paragraph.text or "").strip()
        if caption not in FINAL_FIGURE_EXPLANATIONS:
            continue
        explanation = paragraphs[index + 1]
        if paragraph_has_drawing(explanation) or (explanation.style and explanation.style.name.startswith("Heading")):
            continue
        new_text = FINAL_FIGURE_EXPLANATIONS[caption]
        if (explanation.text or "").strip() != new_text:
            replace_paragraph_text(explanation, new_text)
            changed += 1
    return changed


def collect_residuals(doc: DocumentObject) -> list[str]:
    residual_patterns = [
        "第五章 系统详细设计与实现",
        "第七章 总结与展望",
        "全文共由七章构成",
        "图6-",
        "表6-",
        "图5-4",
        "图5-5",
        "图5-6",
        "图5-7",
        "表5-6",
    ]
    hits: list[str] = []
    for pattern in residual_patterns:
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if pattern in text:
                hits.append(f"{pattern} -> {text}")
                break
    return hits


def verify_figure_order(doc: DocumentObject) -> list[str]:
    issues: list[str] = []
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        caption = (paragraph.text or "").strip()
        if caption not in FINAL_FIGURE_EXPLANATIONS:
            continue
        prev_has_image = index > 0 and paragraph_has_drawing(paragraphs[index - 1])
        next_text = (paragraphs[index + 1].text or "").strip() if index + 1 < len(paragraphs) else ""
        if not prev_has_image:
            issues.append(f"{caption}: 图题前未检测到图片段落")
        if next_text != FINAL_FIGURE_EXPLANATIONS[caption]:
            issues.append(f"{caption}: 图后说明未按预期更新")
    return issues


def main() -> None:
    source = find_source()
    output = build_output_path(source)
    shutil.copy2(source, output)

    doc = Document(str(output))

    add_merged_chapter_structure(doc)
    rewrite_headings(doc)
    text_changed = apply_text_replacements(doc)
    explanation_changed = shorten_figure_explanations(doc)

    doc.save(str(output))

    verified = Document(str(output))
    residuals = collect_residuals(verified)
    figure_issues = verify_figure_order(verified)

    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    print(f"TEXT_PARAGRAPHS_UPDATED={text_changed}")
    print(f"FIGURE_EXPLANATIONS_UPDATED={explanation_changed}")
    print(f"RESIDUAL_COUNT={len(residuals)}")
    for item in residuals:
        print(f"RESIDUAL={item}")
    print(f"FIGURE_ORDER_ISSUE_COUNT={len(figure_issues)}")
    for item in figure_issues:
        print(f"FIGURE_ORDER_ISSUE={item}")


if __name__ == "__main__":
    main()
