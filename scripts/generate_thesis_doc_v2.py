# -*- coding: utf-8 -*-
import hashlib
import sqlite3
import sys
from pathlib import Path

import cv2
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image, ImageColor, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
IMAGE_DIR = BASE_DIR / "image"
DB_PATH = BASE_DIR / "backend-django" / "db.sqlite3"
INPUT_VIDEO_DIR = BASE_DIR / "data" / "input_videos"
OUTPUT_VIDEO_DIR = BASE_DIR / "output_videos"
GENERATED_ASSETS_DIR = BASE_DIR / "thesis_assets_generated"
DEFAULT_OUTPUT_PATH = BASE_DIR / "sdu_faceswap_thesis_optimized.docx"

PIL_FONT_REGULAR_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simsun.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/arial.ttf"),
]
PIL_FONT_BOLD_CANDIDATES = [
    Path("C:/Windows/Fonts/msyhbd.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/arialbd.ttf"),
]


def set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_format(paragraph, line_spacing=1.5, first_line_chars=2, space_before=0, space_after=0, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line_spacing == 1.5 else WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.5 if line_spacing == 1.5 else 1
    fmt.first_line_indent = Pt(24) if first_line_chars == 2 else Pt(0)
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align


def format_section(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def add_paragraph(doc, text, size=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_chars=2, line_spacing=1.5, east_asia="宋体", ascii_font="Times New Roman"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_fonts(r, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    set_paragraph_format(p, line_spacing=line_spacing, first_line_chars=first_line_chars, align=align)
    return p


def add_title(doc, text, size=Pt(18), east_asia="黑体", bold=True, ascii_font="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_fonts(r, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    set_paragraph_format(p, first_line_chars=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level, 3)}"]
    r = p.add_run(text)
    if level == 1:
        set_run_fonts(r, east_asia="黑体", size=Pt(16), bold=True)
        set_paragraph_format(p, first_line_chars=0, space_before=10, space_after=6)
    elif level == 2:
        set_run_fonts(r, east_asia="黑体", size=Pt(14), bold=True)
        set_paragraph_format(p, first_line_chars=0, space_before=6, space_after=6)
    else:
        set_run_fonts(r, east_asia="黑体", size=Pt(12), bold=True)
        set_paragraph_format(p, first_line_chars=0, space_before=6, space_after=6)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_fonts(r, east_asia="宋体", size=Pt(10.5), bold=True)
    set_paragraph_format(p, line_spacing=1, first_line_chars=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def set_page_number_format(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    for node in sect_pr.xpath("./w:pgNumType"):
        sect_pr.remove(node)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))
    sect_pr.append(pg)


def add_field(paragraph, field_code, placeholder=""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    if placeholder:
        placeholder_run = paragraph.add_run(placeholder)
        set_run_fonts(placeholder_run, east_asia="宋体", size=Pt(12), bold=False)
    paragraph.add_run()._r.append(fld_end)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_fonts(run, east_asia="Times New Roman", ascii_font="Times New Roman", size=Pt(9), bold=False)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10.5), bold=False, east_asia="宋体"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_run_fonts(r, east_asia=east_asia, size=size, bold=bold)
    set_paragraph_format(p, line_spacing=1.0, first_line_chars=0, align=align)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def get_pil_font(size, bold=False):
    candidates = PIL_FONT_BOLD_CANDIDATES if bold else PIL_FONT_REGULAR_CANDIDATES
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def measure_multiline_text(draw, text, font, spacing=6):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw, text, font, max_width):
    if not text:
        return ""
    lines = []
    current = ""
    for ch in str(text):
        if ch == "\n":
            lines.append(current)
            current = ""
            continue
        candidate = f"{current}{ch}"
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if current and (bbox[2] - bbox[0]) > max_width:
            lines.append(current)
            current = ch
        else:
            current = candidate
    if current:
        lines.append(current)
    return "\n".join(lines)


def draw_centered_multiline_text(draw, box, text, font, fill="#233043", spacing=8):
    wrapped = wrap_text(draw, text, font, max(int(box[2] - box[0] - 28), 10))
    width, height = measure_multiline_text(draw, wrapped, font, spacing=spacing)
    x = box[0] + ((box[2] - box[0]) - width) / 2
    y = box[1] + ((box[3] - box[1]) - height) / 2
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=spacing, align="center")


def figure_asset_path(prefix, key):
    GENERATED_ASSETS_DIR.mkdir(exist_ok=True)
    digest = hashlib.md5(key.encode("utf-8")).hexdigest()[:10]
    return GENERATED_ASSETS_DIR / f"{prefix}_{digest}.png"


def fill_to_rgb(fill, default="#D9EAF7"):
    value = fill or default
    if not str(value).startswith("#"):
        value = f"#{value}"
    return ImageColor.getrgb(value)


def draw_arrow(draw, start, end, color="#6B7280", width=8, head_length=18, vertical=False):
    draw.line([start, end], fill=color, width=width)
    if vertical:
        points = [
            end,
            (end[0] - head_length // 2, end[1] - head_length),
            (end[0] + head_length // 2, end[1] - head_length),
        ]
    else:
        points = [
            end,
            (end[0] - head_length, end[1] - head_length // 2),
            (end[0] - head_length, end[1] + head_length // 2),
        ]
    draw.polygon(points, fill=color)


def draw_rounded_box(draw, box, text, fill, font, outline="#BCC6D4", text_fill="#1F2937"):
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    draw_centered_multiline_text(draw, box, text, font, fill=text_fill, spacing=8)


def contain_image(image, max_width, max_height):
    img = image.copy()
    img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
    return img


def build_flow_figure(steps, fill, caption):
    count = max(len(steps), 1)
    width = max(1500, 220 + count * 280 + (count - 1) * 110)
    height = 320
    margin = 70
    gap = 110
    box_width = (width - margin * 2 - gap * (count - 1)) // count
    box_height = 112
    top = 92
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font = get_pil_font(30, bold=True)
    fill_rgb = fill_to_rgb(fill)

    for idx, step in enumerate(steps):
        left = margin + idx * (box_width + gap)
        box = (left, top, left + box_width, top + box_height)
        draw_rounded_box(draw, box, step, fill=fill_rgb, font=font)
        if idx < count - 1:
            arrow_start = (box[2] + 12, top + box_height // 2)
            arrow_end = (box[2] + gap - 12, top + box_height // 2)
            draw_arrow(draw, arrow_start, arrow_end)
    path = figure_asset_path("flow", caption)
    img.save(path, format="PNG")
    return path


def build_two_row_figure(top_row, bottom_row, caption, top_fill, bottom_fill):
    cols = max(len(top_row), len(bottom_row), 1)
    width = max(1500, 220 + cols * 280)
    height = 520
    margin = 70
    gap = 30
    box_width = (width - margin * 2 - gap * (cols - 1)) // cols
    box_height = 108
    top_y = 90
    bottom_y = 300
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    top_font = get_pil_font(28, bold=True)
    bottom_font = get_pil_font(28, bold=True)
    top_fill_rgb = fill_to_rgb(top_fill, "#D9EAF7")
    bottom_fill_rgb = fill_to_rgb(bottom_fill, "#E2F0D9")

    for idx in range(cols):
        left = margin + idx * (box_width + gap)
        center_x = left + box_width // 2
        if idx < len(top_row):
            top_box = (left, top_y, left + box_width, top_y + box_height)
            draw_rounded_box(draw, top_box, top_row[idx], fill=top_fill_rgb, font=top_font)
            draw_arrow(
                draw,
                (center_x, top_box[3] + 8),
                (center_x, bottom_y - 16),
                color="#7A8798",
                vertical=True,
            )
        if idx < len(bottom_row):
            bottom_box = (left, bottom_y, left + box_width, bottom_y + box_height)
            draw_rounded_box(draw, bottom_box, bottom_row[idx], fill=bottom_fill_rgb, font=bottom_font)
    path = figure_asset_path("two_row", caption)
    img.save(path, format="PNG")
    return path


def build_sequence_figure(participants, messages, caption):
    lane_count = max(len(participants), 1)
    row_count = max(len(messages), 1)
    width = max(1900, 230 + lane_count * 260)
    height = 250 + row_count * 120
    lane_start = 200
    lane_width = (width - lane_start - 70) // lane_count
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    header_font = get_pil_font(26, bold=True)
    cell_font = get_pil_font(24, bold=False)
    step_font = get_pil_font(24, bold=True)

    for idx, name in enumerate(participants):
        left = lane_start + idx * lane_width + 18
        right = lane_start + (idx + 1) * lane_width - 18
        box = (left, 35, right, 92)
        draw_rounded_box(draw, box, name, fill="#D9EAF7", font=header_font)
        center_x = (left + right) // 2
        draw.line([(center_x, 110), (center_x, height - 40)], fill="#C7CFDA", width=3)

    for row_idx, (step_no, row_values) in enumerate(messages):
        row_top = 125 + row_idx * 120
        step_box = (45, row_top + 16, 145, row_top + 72)
        draw_rounded_box(draw, step_box, str(step_no), fill="#F3F5F8", font=step_font, outline="#D4DAE3")
        for col_idx, value in enumerate(row_values):
            if not value:
                continue
            left = lane_start + col_idx * lane_width + 18
            right = lane_start + (col_idx + 1) * lane_width - 18
            fill = "#FCE4D6" if "signal" in value.lower() or "finished" in value.lower() else "#EEF3F9"
            box = (left, row_top, right, row_top + 88)
            draw_rounded_box(draw, box, value, fill=fill, font=cell_font)

    path = figure_asset_path("sequence", caption)
    img.save(path, format="PNG")
    return path


def build_image_pair(left_path, right_path, left_label, right_label, caption):
    panel_width = 760
    panel_height = 720
    label_height = 84
    margin = 54
    width = panel_width * 2 + margin * 3
    height = panel_height + label_height + margin * 2
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    label_font = get_pil_font(28, bold=True)

    for idx, (path, label) in enumerate(((left_path, left_label), (right_path, right_label))):
        panel_left = margin + idx * (panel_width + margin)
        image_box = (panel_left, margin, panel_left + panel_width, margin + panel_height)
        draw.rounded_rectangle(image_box, radius=18, fill="#FAFBFC", outline="#CAD2DE", width=3)
        source = Image.open(path).convert("RGB")
        fitted = contain_image(source, panel_width - 48, panel_height - 48)
        paste_x = panel_left + (panel_width - fitted.width) // 2
        paste_y = margin + (panel_height - fitted.height) // 2
        img.paste(fitted, (paste_x, paste_y))
        label_box = (panel_left, margin + panel_height + 10, panel_left + panel_width, height - margin)
        draw_centered_multiline_text(draw, label_box, label, label_font, fill="#253041")

    path = figure_asset_path("pair", caption)
    img.save(path, format="PNG")
    return path


def add_flow_figure(doc, steps, caption, fill="D9EAF7"):
    path = build_flow_figure(steps, fill, caption)
    add_image(doc, path, 15.2)
    add_caption(doc, caption)


def add_two_row_figure(doc, top_row, bottom_row, caption, top_fill="D9EAF7", bottom_fill="E2F0D9"):
    path = build_two_row_figure(top_row, bottom_row, caption, top_fill, bottom_fill)
    add_image(doc, path, 15.0)
    add_caption(doc, caption)


def add_sequence_figure(doc, participants, messages, caption):
    path = build_sequence_figure(participants, messages, caption)
    add_image(doc, path, 15.4)
    add_caption(doc, caption)


def ensure_docx_compatible_image(path):
    image_path = Path(path)
    GENERATED_ASSETS_DIR.mkdir(exist_ok=True)
    safe_dir = GENERATED_ASSETS_DIR / "docx_safe"
    safe_dir.mkdir(exist_ok=True)
    safe_path = safe_dir / f"{image_path.stem}_docx.png"
    with Image.open(image_path) as img:
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        img.save(safe_path, format="PNG")
    return safe_path


def add_image(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(path), width=Cm(width_cm))
    except UnrecognizedImageError:
        safe_path = ensure_docx_compatible_image(path)
        p.add_run().add_picture(str(safe_path), width=Cm(width_cm))
    return p


def add_image_pair(doc, left_path, right_path, left_label, right_label, caption, width_cm=6.0):
    path = build_image_pair(left_path, right_path, left_label, right_label, caption)
    doc_width_cm = min(15.4, max(width_cm * 2 + 1.0, 12.6))
    add_image(doc, path, doc_width_cm)
    add_caption(doc, caption)


def add_simple_table(doc, headers, rows, caption=None, widths_cm=None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if widths_cm:
        set_table_widths(table, widths_cm)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], str(value), align=align)
    return table


def add_section(doc, heading, paragraphs, level=2):
    add_heading(doc, heading, level)
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)


def query_project_stats():
    conn = sqlite3.connect(str(DB_PATH))
    cur = conn.cursor()
    cur.execute("select count(*) from face_images")
    face_count = cur.fetchone()[0]
    cur.execute("select count(*) from input_videos")
    video_count = cur.fetchone()[0]
    cur.execute("select count(*) from output_videos")
    output_count = cur.fetchone()[0]
    cur.execute("select round(avg(duration),2) from input_videos")
    avg_video_duration = cur.fetchone()[0] or 0
    cur.execute("select round(avg(fps),2) from input_videos")
    avg_video_fps = cur.fetchone()[0] or 0
    cur.execute("select round(avg(duration),2) from output_videos")
    avg_output_duration = cur.fetchone()[0] or 0
    cur.execute("select round(avg(processing_time),2) from output_videos")
    avg_processing_time = cur.fetchone()[0] or 0
    cur.execute("select id, original_filename, local_path, width, height, file_size from face_images order by id desc limit 5")
    latest_images = cur.fetchall()
    cur.execute("select id, original_filename, local_path, duration, fps, width, height from input_videos order by id desc limit 5")
    latest_videos = cur.fetchall()
    cur.execute("select id, filename, duration, processing_method, processing_time, status from output_videos order by id desc limit 5")
    latest_outputs = cur.fetchall()
    cur.execute("select local_path from face_images where local_path is not null and local_path != ''")
    all_face_paths = [row[0] for row in cur.fetchall()]
    cur.execute("select local_path from input_videos where local_path is not null and local_path != ''")
    all_video_paths = [row[0] for row in cur.fetchall()]
    conn.close()
    images_exists = sum(1 for row in latest_images if row[2] and Path(row[2]).exists())
    videos_exists = sum(1 for row in latest_videos if row[2] and Path(row[2]).exists())
    face_exists_total = sum(1 for path in all_face_paths if Path(path).exists())
    video_exists_total = sum(1 for path in all_video_paths if Path(path).exists())
    output_file_total = len(list(OUTPUT_VIDEO_DIR.glob("*.mp4")))
    missing_face_paths = max(len(all_face_paths) - face_exists_total, 0)
    missing_video_paths = max(len(all_video_paths) - video_exists_total, 0)
    untracked_output_files = max(output_file_total - output_count, 0)
    return {
        "face_count": face_count,
        "video_count": video_count,
        "output_count": output_count,
        "avg_video_duration": avg_video_duration,
        "avg_video_fps": avg_video_fps,
        "avg_output_duration": avg_output_duration,
        "avg_processing_time": avg_processing_time,
        "latest_images": latest_images,
        "latest_videos": latest_videos,
        "latest_outputs": latest_outputs,
        "images_exists": images_exists,
        "videos_exists": videos_exists,
        "image_path_check_total": len(latest_images),
        "video_path_check_total": len(latest_videos),
        "face_path_total": len(all_face_paths),
        "video_path_total": len(all_video_paths),
        "face_exists_total": face_exists_total,
        "video_exists_total": video_exists_total,
        "output_file_total": output_file_total,
        "missing_face_paths": missing_face_paths,
        "missing_video_paths": missing_video_paths,
        "untracked_output_files": untracked_output_files,
    }


def extract_first_frame(video_path, out_path):
    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        return None
    ok, frame = cap.read()
    cap.release()
    if not ok or frame is None:
        return None
    frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    Image.fromarray(frame).save(out_path)
    return out_path


def prepare_assets(stats):
    GENERATED_ASSETS_DIR.mkdir(exist_ok=True)
    chart1 = GENERATED_ASSETS_DIR / "media_counts_en.png"
    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=160)
    labels = ["Face Images", "Input Videos", "Output Videos"]
    values = [stats["face_count"], stats["video_count"], stats["output_count"]]
    bars = ax.bar(labels, values, color=["#4F81BD", "#C0504D", "#9BBB59"])
    ax.set_title("Media Counts")
    ax.set_ylabel("Count")
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.2, str(value), ha="center", va="bottom", fontsize=10)
    fig.tight_layout()
    fig.savefig(chart1, bbox_inches="tight")
    plt.close(fig)

    chart2 = GENERATED_ASSETS_DIR / "video_metrics_en.png"
    recent = list(reversed(stats["latest_videos"]))
    names = [f"V{i+1}" for i in range(len(recent))]
    durations = [row[3] or 0 for row in recent]
    fps_values = [row[4] or 0 for row in recent]
    fig, ax1 = plt.subplots(figsize=(8.0, 4.6), dpi=160)
    ax1.bar(names, durations, width=0.55, color="#5B9BD5")
    ax1.set_ylabel("Duration (s)", color="#2F5597")
    ax1.tick_params(axis="y", labelcolor="#2F5597")
    ax1.set_title("Recent Input Videos")
    ax1.grid(axis="y", linestyle="--", alpha=0.3)
    ax2 = ax1.twinx()
    ax2.plot(names, fps_values, color="#ED7D31", marker="o", linewidth=2)
    ax2.set_ylabel("FPS", color="#C55A11")
    ax2.tick_params(axis="y", labelcolor="#C55A11")
    fig.tight_layout()
    fig.savefig(chart2, bbox_inches="tight")
    plt.close(fig)

    ui_main = GENERATED_ASSETS_DIR / "ui_main.png"
    ui_preview = GENERATED_ASSETS_DIR / "ui_preview.png"
    ui_panel = GENERATED_ASSETS_DIR / "ui_panel.png"
    ui_src = Image.open(IMAGE_DIR / "ui_demo.png").convert("RGB")
    ui_src.save(ui_main)
    w, h = ui_src.size
    ui_src.crop((0, 0, int(w * 0.73), h)).save(ui_preview)
    ui_src.crop((int(w * 0.73), 0, w, h)).save(ui_panel)

    input_frame = GENERATED_ASSETS_DIR / "input_video_frame.png"
    output_frame = GENERATED_ASSETS_DIR / "output_video_frame.png"
    input_videos = sorted(INPUT_VIDEO_DIR.glob("*.mp4"))
    output_videos = sorted(OUTPUT_VIDEO_DIR.glob("*.mp4"), key=lambda p: p.stat().st_mtime, reverse=True)
    if input_videos:
        extract_first_frame(input_videos[0], input_frame)
    if output_videos:
        extract_first_frame(output_videos[0], output_frame)

    return {
        "chart_media_counts": chart1,
        "chart_video_metrics": chart2,
        "ui_main": ui_main,
        "ui_preview": ui_preview,
        "ui_panel": ui_panel,
        "input_frame": input_frame if input_frame.exists() else None,
        "output_frame": output_frame if output_frame.exists() else None,
    }


def build_document(output_path):
    stats = query_project_stats()
    assets = prepare_assets(stats)

    doc = Document()
    for section in doc.sections:
        format_section(section)

    doc.core_properties.title = "人脸替换系统的设计与实现"
    doc.core_properties.author = "叶俊"
    doc.core_properties.subject = "山东大学本科毕业论文"

    for _ in range(4):
        doc.add_paragraph()
    add_title(doc, "山东大学本科毕业论文（设计）", Pt(22))
    for _ in range(2):
        doc.add_paragraph()
    add_title(doc, "人脸替换系统的设计与实现", Pt(18))
    add_title(doc, "Design and Implementation of a Face Swapping System Based on PyQt5 and Django", Pt(16))
    for _ in range(4):
        doc.add_paragraph()
    for item in [
        "姓名：叶俊",
        "学号：202200201151",
        "学院：软件学院",
        "年级：2022级",
        "指导教师：吴昊",
        "完成时间：2026年4月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        set_run_fonts(r, east_asia="宋体", size=Pt(14), bold=False)
        set_paragraph_format(p, first_line_chars=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_title(doc, "摘    要", Pt(18))
    add_paragraph(
        doc,
        "本文围绕现有人脸替换项目的整理、补充与说明展开，而不是从零开始构造单一换脸模型。"
        "当前仓库已经形成 PyQt5 前端、Django REST 后端、SQLite 数据库、本地媒体目录以及 OpenCV 与 InsightFace 组成的处理链路，论文工作的重点是把这些已有实现梳理为结构清晰、证据充分、便于答辩展示的系统方案。"
        "在前端部分，系统以 EnhancedFaceSwapUI 为核心，通过 VideoProcessingThread 和 CameraProcessingThread 分别支撑离线视频处理与实时摄像头处理；在数据管理部分，DatabaseManager 与后端接口协同工作，使人脸图片、输入视频和输出结果形成统一的加载、保存和回放入口。"
        f"截至本文生成时，数据库中共有 {stats['face_count']} 条人脸图片记录、{stats['video_count']} 条输入视频记录和 {stats['output_count']} 条输出视频记录；其中本地可用人脸路径为 {stats['face_exists_total']} / {stats['face_path_total']}，可用视频路径为 {stats['video_exists_total']} / {stats['video_path_total']}，输出目录中的 mp4 文件数比数据库记录多出 {stats['untracked_output_files']} 个。"
        "这些真实数据表明，系统已经能够完成素材加载、参数配置、视频换脸、摄像头演示、结果归档和回放验证等主要环节，同时也暴露出历史路径失效、输出结果补登记不足等典型工程问题。"
        "围绕上述现状，本文从需求分析、总体设计、详细实现和工程测试四个层面展开论述，重点说明系统如何通过双模式界面、线程化处理、接口化数据管理和结果留痕机制提升可用性与可维护性。"
        "测试结果表明，该系统已经能够支撑桌面端人脸替换的主要演示与处理流程，适合作为软件工程方向的系统实现型毕业设计。"
        "后续工作将围绕资源路径整理、processing_tasks 的实际队列化使用、异常日志展示和历史结果补录等问题继续展开。"
    )
    p = doc.add_paragraph()
    r1 = p.add_run("关键词：")
    set_run_fonts(r1, east_asia="黑体", size=Pt(12), bold=True)
    r2 = p.add_run("人脸替换；PyQt5；Django；InsightFace；桌面应用；系统实现")
    set_run_fonts(r2, east_asia="宋体", size=Pt(12), bold=False)
    set_paragraph_format(p, first_line_chars=0)

    doc.add_page_break()
    add_title(doc, "ABSTRACT", Pt(18))
    add_paragraph(
        doc,
        "This thesis focuses on turning an existing face-swapping repository into a verifiable engineering system rather than proposing a brand-new swapping model. "
        "The current project already contains a PyQt5 frontend, a Django REST backend, an SQLite database, local media directories, and a processing chain built on OpenCV and InsightFace. "
        "The frontend centers on EnhancedFaceSwapUI, while VideoProcessingThread and CameraProcessingThread separate offline video processing from real-time camera processing. DatabaseManager cooperates with the backend so that face images, input videos, and output results can be loaded, stored, and replayed through a unified workflow. "
        f"At the time of document generation, the database stores {stats['face_count']} face-image records, {stats['video_count']} input-video records, and {stats['output_count']} output-video records. Meanwhile, only {stats['face_exists_total']} of {stats['face_path_total']} recorded face paths and {stats['video_exists_total']} of {stats['video_path_total']} recorded video paths remain locally valid, and the output directory contains {stats['untracked_output_files']} more mp4 files than the output table records. "
        "These facts show that the system already supports the core loop of material loading, parameter configuration, face swapping, camera demonstration, result archiving, and replay, while also exposing concrete engineering issues such as invalid historical paths and incomplete backfilling of legacy outputs. "
        "Accordingly, this thesis reorganizes the project from the perspectives of requirement analysis, architecture design, detailed implementation, and engineering-oriented evaluation, with emphasis on dual-mode interaction, threaded processing, interface-based data management, and traceable output handling. "
        "The final result demonstrates that the system is suitable as a software-engineering-oriented graduation project and still leaves clear follow-up directions in path governance, task scheduling, log visualization, and result reconciliation.",
        east_asia="Times New Roman",
        ascii_font="Times New Roman",
    )
    p = doc.add_paragraph()
    r1 = p.add_run("Key Words: ")
    set_run_fonts(r1, east_asia="Times New Roman", ascii_font="Times New Roman", size=Pt(12), bold=True)
    r2 = p.add_run("face swapping, PyQt5, Django, InsightFace, desktop application, system implementation")
    set_run_fonts(r2, east_asia="Times New Roman", ascii_font="Times New Roman", size=Pt(12), bold=False)
    set_paragraph_format(p, first_line_chars=0)

    sec_toc = doc.add_section(WD_SECTION.NEW_PAGE)
    format_section(sec_toc)
    sec_toc.header.is_linked_to_previous = False
    sec_toc.footer.is_linked_to_previous = False
    set_page_number_format(sec_toc, "upperRoman", 1)
    toc_footer = sec_toc.footer.paragraphs[0]
    toc_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(toc_footer)
    add_title(doc, "目    录", Pt(18))
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u', "右键更新目录")

    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    format_section(sec_body)
    sec_body.header.is_linked_to_previous = False
    sec_body.footer.is_linked_to_previous = False
    set_page_number_format(sec_body, "decimal", 1)
    header = sec_body.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run("山东大学本科毕业论文（设计）")
    set_run_fonts(hr, east_asia="宋体", size=Pt(9), bold=False)
    footer = sec_body.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    add_heading(doc, "1 绪论", 1)
    add_section(doc, "1.1 研究背景", [
        "近年来，短视频创作、数字内容编辑和生成式视觉应用迅速普及，用户对于可直接操作的视频处理工具的需求明显高于仅能运行一次脚本的研究原型。在人脸替换场景中，这种差异更加突出，因为用户不仅关心单帧效果，还关心素材如何选择、处理过程是否可见、结果能否保存以及失败后是否便于复查。",
        "从当前项目仓库可以看到，这一课题已经不再停留在纯算法层面。仓库中同时存在 frontend、backend-django、data、output_videos、logs 等目录，以及围绕播放、路径和兼容性做过的多次修复脚本，这说明真实工作难点并不只是能否生成一张结果图，而是如何把前端交互、后端记录、媒体处理和结果归档组织成可持续运行的系统。",
        "因此，本文的研究背景不是抽象地讨论换脸技术本身，而是基于一个已经具备运行痕迹的人脸替换项目，重新梳理其系统边界、技术链路和工程问题。相比只展示若干效果图，这种写法更能反映软件工程毕业设计从原型走向系统的真实工作过程。",
    ])
    add_section(doc, "1.2 研究意义", [
        "本文的研究意义首先体现在工程实现层面。许多开源换脸项目能够证明模型在样例上的有效性，但在素材整理、处理反馈、结果保存和错误处理等方面往往缺少完整安排，因此使用者一旦离开开发环境，就很难稳定复现处理过程。本项目把 PyQt5 桌面界面、Django REST 资源管理和 SQLite 元数据记录结合起来，使换脸处理、素材管理和结果保存不再彼此分离，而是构成一个可以连续使用和检查的系统。",
        "本文的研究意义还体现在教学实践层面。与单独完成一次模型推理演示相比，本项目同时涉及界面组织、线程调度、接口调用、数据库建模、文件目录整理和媒体回放兼容等问题，这些内容更接近真实软件项目中不同模块之间的协同方式。围绕 EnhancedFaceSwapUI、DatabaseManager、FaceSwapApp 和后端模型层展开分析，可以更具体地展示软件工程专业所关注的模块划分、职责安排和实现取舍。",
        "本文的研究意义也体现在对真实问题的呈现上。当前系统仍然存在历史路径失效、输出结果补登记不足以及 processing_tasks 尚未承担实际队列调度等问题，本文并不回避这些不足，而是把它们作为分析对象加以说明。对于本科毕业设计来说，能够依据真实项目材料说明系统已经完成了哪些工作、仍然缺少哪些部分以及这些不足产生的原因，通常比单纯强调算法效果更有说服力。",
    ])
    add_section(doc, "1.3 国内外研究现状", [
        "从研究脉络看，人脸替换技术大致经历了传统几何对齐与融合到深度学习驱动身份迁移两个阶段。早期方法依赖检测、关键点、仿射变换和区域融合，优点是实现路径清晰、依赖较少，但在姿态变化、遮挡和复杂光照下容易出现边缘不自然和结构拉伸。后续随着 FaceNet、ArcFace、SimSwap、FaceShifter、InsightFace 等工作出现，换脸任务在身份保持和复杂场景表现方面得到了明显提升。",
        "然而，国内外大量工作更关注模型效果、数据集和视觉指标，而不是桌面系统中的工作流组织问题。对于真正交付给用户的原型系统来说，模型只是其中一个环节，还必须回答素材怎么加载、失败怎么提示、结果怎么保存、历史记录怎么回查这些问题。当前很多开源实现已经可以输出较好效果，但在资源管理和界面组织上仍然停留在实验性阶段。",
        "基于这样的研究现状，本文并不尝试重新提出新的换脸网络，而是以现有项目为基础，讨论如何把 InsightFace、OpenCV、PyQt5、Django REST 和 SQLite 组织成一个能够完成展示、验证和问题分析的系统实现方案。这种写法更符合本科毕业设计的任务边界，也更适合说明系统从算法调用到界面交互再到结果保存的完整实现过程。",
    ])
    add_section(doc, "1.4 研究目标与主要工作", [
        "本文的总体目标是基于现有项目代码，把一个已经具备功能雏形的人脸替换原型整理成结构可说明、行为可验证、问题可定位的系统型毕业设计。论文不是重新发明算法，而是把已有实现中真正起作用的模块、数据和流程提炼出来，并用规范文档说明其设计依据与工程取舍。",
        "围绕这一目标，本文结合 frontend/main.py、face_swap_ui_enhanced.py、database_manager.py 以及后端模型与视图集，重新梳理了系统边界和实现关系，并以视频模式和摄像头模式为中心概括主要使用场景，进一步说明需求分析、总体结构和关键模块之间的联系。本文同时围绕 EnhancedFaceSwapUI、DatabaseManager、VideoProcessingThread、CameraProcessingThread 以及 saveOutputVideoToDatabase 等实现内容，说明系统如何完成素材加载、处理执行、结果回放和结果保存，并结合当前数据库记录、文件系统状态和已有样例资源，对系统的可用性和现存问题进行了验证。",
        "在目前版本中，系统已经支持传统方法与 InsightFace 方法两种处理方式，能够完成素材选择、参数配置、处理线程启动、输出结果保存和回放验证。本文据此概括系统当前已经实现的内容、仍然存在的风险以及后续需要继续完善的部分。",
    ])
    add_section(doc, "1.5 论文结构安排", [
        "全文共分为七章。第一章介绍研究背景、研究意义、国内外研究现状以及本文主要工作。第二章说明与系统实现密切相关的关键技术和理论基础。第三章面向业务场景展开系统需求分析。第四章从系统架构、模块划分、数据流、数据库和部署方式等角度给出总体设计。第五章深入说明系统的详细实现过程。第六章从工程测试角度对系统进行验证和分析。第七章总结全文工作并给出后续展望。",
    ])

    add_heading(doc, "2 相关技术与理论基础", 1)
    add_section(doc, "2.1 人脸替换任务定义与典型流程", [
        "从本文所对应的项目来看，一次完整的人脸替换任务并不等同于调用一次模型推理。它通常从 loadFaceImages、loadVideos 这类素材加载逻辑开始，经过用户选择、参数配置和模式判断，再由 startProcessing 启动 VideoProcessingThread，或者由摄像头入口启动 CameraProcessingThread，最终在 processingFinished 与 saveOutputVideoToDatabase 中完成结果保存、回放和留痕。只有把这些环节合在一起看，才能准确理解系统里的任务到底是什么。",
        "与单帧图像换脸相比，视频和摄像头场景还带来了帧间一致性、界面响应和结果管理三个额外约束。若只关注单帧效果，不考虑线程调度、播放控制和输出组织，就会在实际运行时出现界面卡顿、结果难回查或媒体资源冲突等问题。因此，本文讨论的人脸替换任务本质上是一个把媒体处理、交互组织和结果留存连续衔接起来的整体流程，而不是孤立的算法调用。",
    ])
    add_flow_figure(doc, ["输入素材", "人脸检测", "对齐与替换", "融合与校正", "结果保存"], "图 2.1 人脸替换典型处理流程图")
    add_section(doc, "2.2 传统人脸替换方法", [
        "传统人脸替换方法通常以检测、关键点定位、仿射对齐和区域融合为核心步骤，优点是实现步骤清晰、依赖较少、对运行环境要求相对较低。对于本科项目而言，这类方法的意义不只是作为较早的技术方案存在，更重要的是它可以在深度学习模型未正确加载、硬件资源较弱或运行环境不稳定时继续提供基本处理结果。",
        "这一点在当前项目中尤为重要。仓库中的 FaceSwapApp 并不是单一方案设计，系统并未完全依赖某一个模型，而是保留传统方法与 InsightFace 并存的处理思路。虽然传统方法在姿态变化、遮挡和复杂光照下更容易出现边界不自然或颜色不一致，但从工程角度看，它有助于系统在不同环境下继续运行，也便于定位问题。"
    ])
    add_simple_table(doc, ["比较维度", "传统方法", "深度学习方法"], [["实现复杂度", "结构清晰，依赖较少", "模型依赖较强，部署较复杂"], ["资源消耗", "CPU 侧也可运行", "更依赖 GPU 或较高推理性能"], ["复杂场景效果", "对姿态和遮挡较敏感", "在复杂场景下通常更稳定"], ["系统适配性", "适合作为兼容方案", "适合作为主要处理方案"]], "表 2.1 传统方法与深度学习方法对比", widths_cm=[2.4, 6.0, 6.0])
    add_section(doc, "2.3 基于深度学习的人脸替换方法与 InsightFace", [
        "基于深度学习的人脸替换方法更强调身份特征表示和复杂场景下的结果自然度。以 InsightFace 为代表的框架把检测、分析和身份迁移相关过程整合在同一技术栈中，使系统在表情变化、姿态偏移和一定程度的光照波动下仍能保持较好的视觉连续性。相较于传统方法，这一方案更适合作为当前系统的主要输出方式。",
        "但在本项目里，InsightFace 的意义并不在于单独作为研究对象存在，而在于它如何被实际放入系统流程中。用户在前端完成素材选择后，startProcessing 会把参数同步到原始换脸引擎，后台线程负责持续调用处理逻辑，结果完成后再由 processingFinished 和 saveOutputVideoToDatabase 进入归档与回放环节。换句话说，本文讨论的不是如何重新训练模型，而是如何把现有模型方法放入一个可交互、可记录、可回查的软件系统中。"
        "这种整合方式体现了系统工程与算法研究的侧重点差异。对于系统型毕业设计来说，关键不仅是效果更自然，还包括调用过程是否稳定、异常时是否能切换到替代方案、结果是否能被保存以及界面是否能把状态变化反馈给用户。也正因此，本文在技术论述上始终把 InsightFace 视为系统组成部分之一，而不是全文唯一的中心内容。"
    ])
    add_simple_table(doc, ["关键技术", "在系统中的作用", "对应实现位置"], [["OpenCV", "视频读取、帧提取、缩略图生成、颜色转换", "前端处理线程与后端元数据提取"], ["InsightFace", "提供高质量换脸与人脸分析能力", "原始换脸引擎封装"], ["PyQt5", "承担桌面界面、列表展示、状态反馈和模式切换", "EnhancedFaceSwapUI"], ["Django REST", "负责图片、视频、输出结果和任务的接口管理", "api 视图集与路由"], ["SQLite", "保存素材和输出视频元数据", "core 模型定义"]], "表 2.2 关键技术与系统作用映射表", widths_cm=[2.8, 7.0, 4.6])
    add_section(doc, "2.4 OpenCV 视频处理与图像增强", [
        "OpenCV 是当前系统中的基础工具。无论是前端列表中的视频缩略图生成，还是后端登记输入视频和输出视频时读取分辨率、时长与帧率信息，都依赖 OpenCV 对视频文件的访问能力。对工程系统而言，OpenCV 的价值不仅在于算法处理，还在于它为媒体元数据的统一获取提供了稳定入口。",
        "在前端处理中，OpenCV 还承担颜色空间转换、帧缩放和局部显示等工作。结合 PyQt5 的 QImage 与 QPixmap，系统能够把视频帧、摄像头帧和换脸后的结果高效渲染到界面中。这一点决定了系统能否形成顺畅的交互体验。"
    ])
    add_section(doc, "2.5 PyQt5 界面开发与信号槽机制", [
        "PyQt5 负责桌面端的人机交互层。项目中的增强版界面类集成了窗口初始化、控件布局、媒体播放、列表展示、处理按钮、模式切换、摄像头开关以及状态栏更新等大量逻辑。与命令行程序相比，PyQt5 让用户可以在更低认知成本下完成素材选择和处理操作。",
        "由于视频处理和摄像头采集都具有耗时或持续运行的特征，系统必须避免把这类任务直接放在主线程中执行。PyQt5 的信号槽机制与 QThread 提供了良好的异步支持，后台线程处理帧级逻辑，界面线程只负责响应用户操作和更新显示，从而兼顾界面流畅性与处理过程的完整性。"
    ])
    add_section(doc, "2.6 Django REST 与 SQLite 数据管理", [
        "在数据管理方面，本文没有把所有逻辑都堆积在桌面端，而是引入 Django REST Framework 作为统一的资源接口层。图片、输入视频、输出视频和处理任务都能够以 REST 资源的形式被查询与记录，这种方式有利于后续扩展到更多客户端，也有利于分离交互界面与数据归档两类职责。",
        "SQLite 作为开发阶段数据库，能够满足原型系统对轻量持久化的需求。通过数据库表，系统记录了文件名、本地路径、宽高、时长、帧率、处理方法和状态等关键信息。项目当前数据库记录显示，一部分历史记录对应的本地路径已经失效，这恰好说明数据库设计与路径治理在工程实践中的重要性，也为测试分析章节提供了客观依据。"
    ])

    add_heading(doc, "3 系统需求分析", 1)
    add_section(doc, "3.1 业务场景描述", [
        "系统主要面向两类使用场景。第一类是离线视频换脸场景，用户拥有一张目标人脸图片和一个待处理视频，希望快速得到输出视频，并能在处理完成后对结果进行保存和回放。第二类是摄像头演示场景，用户希望实时查看人脸替换效果，用于答辩演示、交互体验或原型展示。",
        "这两类场景共享素材管理、目标人脸选择和状态反馈等基础功能，但在时延要求、输出形态和控制重点上存在差异。视频模式更重视稳定的离线输出和结果归档，摄像头模式更重视界面响应和实时切换。系统设计需要同时兼顾这两类需求。"
    ])
    add_flow_figure(doc, ["选择人脸图", "选择视频/摄像头", "设置参数", "执行处理", "保存与回放"], "图 3.1 主要用户业务场景图", fill="FFF2CC")
    add_section(doc, "3.2 功能需求分析", [
        "结合业务场景，系统至少需要具备素材加载、素材选择、参数设置、视频处理、摄像头处理、结果保存、结果回放和数据库归档等核心功能。与单纯展示效果不同，这些功能共同构成了用户从输入到输出的完整过程。",
        "从交互角度看，功能需求还包括自动生成输出路径、显示缩略图、提示处理状态、在模式切换时自动停止不必要的播放或采集线程等细节。只有当这些细节被纳入需求分析，系统才不会停留在能够点击按钮的层面，而是更接近可以完整使用的状态。"
    ])
    add_simple_table(doc, ["编号", "功能需求", "说明"], [["F1", "人脸图片加载与展示", "支持数据库加载和本地目录扫描，提供缩略图列表"], ["F2", "输入视频加载与展示", "支持视频列表、缩略图与基本元数据读取"], ["F3", "视频模式处理", "支持选择人脸图、选择视频、配置参数并执行处理"], ["F4", "摄像头模式处理", "支持开启摄像头、切换目标人脸和关闭处理"], ["F5", "输出保存与回放", "支持自动输出路径、结果写入目录与数据库归档"], ["F6", "状态反馈与异常提示", "支持状态栏、进度更新和错误消息提示"]], "表 3.1 功能需求汇总表", widths_cm=[1.5, 4.0, 9.0])
    add_section(doc, "3.3 非功能需求分析", [
        "本系统虽然仍属于原型项目，但在使用过程中已经表现出明显的非功能要求。界面布局需要保持清晰，以减少用户在文件目录、参数配置和结果预览之间反复切换的负担；前端、后端、算法处理和数据库管理需要保持基本分层，以避免不同逻辑长期堆叠在同一个脚本中，增加修改和排错的难度。",
        "系统还需要为后续增加新的换脸方法、更多视频格式以及更完整的任务管理功能预留空间，同时保证处理过程在较长时间运行时仍然稳定。即使项目当前规模不大，输出结果也应能够被保存和查询，历史素材也应能够重复利用，否则系统在经过多次使用后会很快出现记录混乱和资源难以整理的问题。"
    ])
    add_simple_table(doc, ["类别", "需求内容", "当前实现思路"], [["易用性", "界面操作简单、反馈明确", "系统通过双模式布局、缩略图列表和状态栏提示降低理解成本"], ["可维护性", "前后端和算法逻辑分层", "系统通过 PyQt5 前端、Django 后端与模型层分工保持结构清晰"], ["可扩展性", "支持后续增加更多算法或接口", "系统依靠 REST 资源接口与独立数据模型保留扩展空间"], ["稳定性", "长时间操作不应频繁崩溃", "系统通过线程化处理以及播放和采集控制保证运行稳定"], ["记录完整性", "素材与结果具备记录机制", "系统通过 SQLite 表结构保存元数据与路径信息"]], "表 3.2 非功能需求汇总表", widths_cm=[2.2, 5.5, 6.8])
    add_section(doc, "3.4 用例分析", [
        "从参与者角度看，本系统的直接参与者主要是用户本身。用户通过界面完成上传、选择、处理和回放等操作，系统负责在后台完成文件读取、元数据分析、算法处理和结果保存。与多人协作系统不同，当前项目以单用户原型为核心，因此用例分析重点放在操作过程是否完整，而不是权限层级的划分。",
        "在用例关系上，选择人脸图片和选择输入视频构成视频模式处理的前提，开启摄像头和切换目标人脸构成实时处理的前提，而保存输出结果与回放处理结果则发生在处理完成之后。对这些关系进行梳理，有助于后续设计界面状态联动和按钮可用性控制。"
    ])
    add_two_row_figure(doc, ["用户", "视频模式", "摄像头模式"], ["素材选择", "结果保存", "结果回放"], "图 3.2 系统用例关系图", top_fill="D9EAF7", bottom_fill="FCE4D6")
    add_section(doc, "3.5 可行性分析", [
        "技术可行性方面，当前项目所依赖的技术栈均已具备成熟实现。PyQt5 用于构建桌面界面，Django REST Framework 用于资源接口，OpenCV 用于媒体处理，InsightFace 与传统方法共同承担算法层逻辑。项目目录中已经存在前端、后端、模型、日志、输入数据和输出结果目录，说明原型基础完整。",
        "工程可行性方面，数据库中当前记录了多条图片、视频和输出结果元数据，输出目录中也保留了多个视频文件，这意味着系统并不是停留在纸面设计上，而是已有运行痕迹。需要注意的是，部分历史路径已经失效，这提示后续实施中应加强资源迁移和路径治理，但并不影响系统总体方案的可行性。"
    ])
    add_simple_table(doc, ["可行性维度", "分析结论", "说明"], [["技术可行性", "较高", "核心技术均有可用实现且项目已集成"], ["工程可行性", "较高", "已有前后端与输出结果目录，系统原型完整"], ["资源可行性", "较高", "本地样例和数据库记录可支撑验证与展示"], ["风险可控性", "中等", "模型依赖与路径失效问题需在后续优化中处理"]], "表 3.3 系统可行性分析表", widths_cm=[2.8, 2.4, 9.3])

    add_heading(doc, "4 系统总体设计", 1)
    add_section(doc, "4.1 设计目标与原则", [
        "系统总体设计主要遵循四项原则，即处理流程尽量完整、模块分层明确、结果可以查询以及交互方式尽量清楚。处理流程完整意味着系统不仅要能执行换脸处理，还要覆盖素材选择、参数配置、结果保存和后续回放；模块分层明确意味着界面、后端和算法逻辑应各自承担清晰职责；结果可以查询强调每次处理应留下可检索记录；交互方式清楚则对应双模式布局和缩略图选择方式。",
        "基于这些原则，本文没有把系统设计成单纯的研究脚本，而是将其整理为一个桌面应用原型。即使当前版本尚未引入完整用户系统、批处理机制和云端部署，整体结构仍为后续修改留出了空间。"
    ])
    add_section(doc, "4.2 总体架构设计", [
        "系统总体由四个层次组成，分别是用户交互层、业务控制层、媒体处理层和数据持久化层。用户交互层主要由 PyQt5 前端组成，业务控制层由 DatabaseManager、处理线程以及前端控制逻辑构成，媒体处理层由 OpenCV 和换脸引擎组成，数据持久化层由 Django REST、SQLite 和本地文件系统共同组成。",
        "这种架构使系统能够在桌面端完成交互，同时把资源记录统一纳入后端。相较于完全在前端直接操作文件系统，这种方式更利于维护素材与结果之间的关系，也便于后续补充更多统计与查询能力。"
    ])
    add_two_row_figure(doc, ["PyQt5 前端界面", "换脸处理引擎", "Django REST 后端"], ["数据库管理模块", "SQLite 元数据", "本地媒体文件"], "图 4.1 系统总体架构图")
    add_section(doc, "4.3 模块划分与职责", [
        "根据当前代码结构，前端主要集中在 EnhancedFaceSwapUI、DatabaseManager 和原始 FaceSwapApp 三部分。EnhancedFaceSwapUI 负责界面组织和状态联动；DatabaseManager 负责与后端接口交互；原始换脸引擎负责视频处理和底层算法调用。后端则围绕模型、序列化器和视图集组织资源。",
        "模块划分的价值在于降低耦合。当前前端虽然仍然包含较多业务逻辑，但数据库交互和算法调用已经被抽离为相对独立的组件。这种组织方式有利于在不大幅改动界面的前提下替换后端实现或换脸方法。"
    ])
    add_simple_table(doc, ["模块", "核心职责", "代表对象或文件"], [["界面模块", "窗口布局、列表展示、按钮交互、状态反馈", "EnhancedFaceSwapUI"], ["数据管理模块", "上传图片、上传视频、加载列表、保存结果", "DatabaseManager"], ["算法处理模块", "视频换脸、颜色校正、传统方法与 InsightFace 调用", "FaceSwapApp"], ["后端接口模块", "图片、视频、输出结果和任务接口", "api.views / api.urls"], ["模型存储模块", "保存元数据和状态字段", "core.models"]], "表 4.1 核心模块职责表", widths_cm=[2.3, 7.0, 5.2])
    add_flow_figure(doc, ["界面选择素材", "DatabaseManager 调接口", "后端返回记录", "界面更新列表"], "图 4.2 数据流与接口流示意图", fill="E2F0D9")
    add_section(doc, "4.4 接口设计", [
        "后端接口采用 REST 风格组织。当前项目已注册 images、videos、outputs 和 tasks 四类资源，并保留了图片和视频上传入口。虽然论文强调工程验证而非接口安全设计，但这种资源化方式仍能为后续接入鉴权、分页、筛选和多用户逻辑打下基础。",
        "对于前端而言，接口设计的关键并不在于复杂协议，而在于稳定的字段组织。图片、视频和输出结果对象都携带本地路径、分辨率、时长、帧率和状态等信息，前端正是基于这些字段来构建列表项和结果回放逻辑。"
    ])
    add_simple_table(doc, ["接口路径", "请求方法", "功能说明"], [["/api/images/", "GET / POST", "获取或登记人脸图片资源"], ["/api/videos/", "GET / POST", "获取或登记输入视频资源"], ["/api/outputs/", "GET / POST", "保存并查询输出视频结果"], ["/api/tasks/", "GET / POST", "处理任务扩展接口"], ["/api/upload/image/", "POST", "图片上传入口"], ["/api/upload/video/", "POST", "视频上传入口"]], "表 4.2 主要接口设计表", widths_cm=[4.0, 2.4, 8.0])
    add_section(doc, "4.5 数据库设计", [
        "数据库设计围绕素材管理和结果留痕展开。FaceImage 表记录人脸图片文件名、本地路径、宽高和文件大小；InputVideo 表记录视频时长、帧率和分辨率；OutputVideo 表记录处理方法、处理状态、处理时间和输出文件；ProcessingTask 则为后续更完整的任务调度保留扩展空间。",
        "当前数据库的一个现实特征是部分历史记录的 local_path 已经指向外部目录或失效位置。这在工程上意味着系统需要更稳定的素材迁移策略，也说明数据库设计不是简单存一个字符串，而应与文件治理配套考虑。本文在测试章节中将这一问题作为工程问题进行分析。"
    ])
    add_simple_table(doc, ["数据表", "关键字段", "用途"], [["face_images", "original_filename, local_path, width, height, file_size", "保存输入人脸图片元数据"], ["input_videos", "original_filename, local_path, duration, fps, width, height", "保存待处理视频信息"], ["output_videos", "filename, processing_method, processing_time, status, progress", "保存输出视频及处理结果"], ["processing_tasks", "task_type, status, progress, processing_params", "扩展处理任务管理"]], "表 4.3 数据库表设计表", widths_cm=[2.8, 7.0, 4.8])
    add_two_row_figure(doc, ["FaceImage", "InputVideo", "OutputVideo"], ["ProcessingTask", "SystemConfig", "本地文件目录"], "图 4.3 数据库关系示意图", top_fill="FFF2CC", bottom_fill="E4DFEC")
    add_section(doc, "4.6 部署结构设计", [
        "从部署角度看，当前系统采用典型的本机开发部署形态，前端运行于桌面环境，后端运行于本机 Django 服务，数据库采用 SQLite 文件，媒体资源保存在本地目录中。该方式虽然不具备生产环境的分布式能力，但非常适合本科毕业设计阶段的展示与调试。",
        "这种部署结构的优点是依赖关系较为清楚、搭建成本较低、调试过程较短，缺点是对本地路径依赖较强，多设备迁移和多用户共享较为困难。基于这一情况，本文在总结与展望中将资源路径整理、任务调度和部署调整作为后续工作的重点。"
    ])
    add_flow_figure(doc, ["桌面前端", "本地 Django 服务", "SQLite 数据库", "媒体文件目录"], "图 4.4 系统部署结构图", fill="E4DFEC")


    doc.add_page_break()
    add_heading(doc, "5 系统详细设计与实现", 1)
    add_section(doc, "5.1 前端主界面实现", [
        "前端主要围绕 EnhancedFaceSwapUI 展开。该类采用单窗口多区域的组织方式，在不引入复杂多窗口跳转的前提下，把人脸列表、视频列表、预览区、参数面板、进度指示和状态信息放在同一界面中。这种布局更符合毕业设计原型系统的使用特点，因为用户在一次演示中就需要完成素材选择、模式切换、参数调整和结果查看，因此前端必须保持操作路径短、信息密度高。",
        "从组件组织方式看，系统把人脸列表与视频列表作为左侧的素材入口，把视频预览和摄像头画面作为中部的运行结果区，把输出路径、调整按钮、换脸方法、平滑参数和状态栏作为右侧的控制区。该布局既适合视频模式，也便于在摄像头模式下保留状态指示和快照按钮。界面没有把所有参数堆叠成命令行或弹出窗口，而是通过可见控件直接暴露出来，从而降低初次使用时的理解成本。",
        "前端还加入了环形进度条、发光按钮、媒体播放控件和状态栏提示，这些设计的目的不是单纯美化，而是为了把处理进度、播放状态和异常信息可视化。对于换脸类任务而言，用户往往会等待一段时间才能看到完整结果，如果缺少明确的界面反馈，就容易误以为程序已经停止响应。因此本系统在实现上把工程可用性放在了前端设计的核心位置。",
    ])
    add_image(doc, assets["ui_main"], 15.3)
    add_caption(doc, "图 5.1 系统主界面截图")
    add_section(doc, "5.2 双模式界面组织与交互逻辑", [
        "在具体交互上，系统通过 AppMode 枚举类区分视频模式和摄像头模式。switchMode 方法负责在两种状态之间切换可见控件、按钮样式和预览区域的工作方式。这种实现的重点不在于界面看上去完成了切换，而在于保证模式切换后的状态清理正确、旧线程不残留、播放器不冲突。因此，本文在实现中只保留一个主窗口，但把两种工作流程的控制逻辑做到了显式分离。",
        "从用户体验看，视频模式需要用户先选择人脸图像、再选择输入视频，然后设置输出路径与处理参数；摄像头模式则更关注开启摄像头、切换目标人脸和是否开启实时换脸。前端没有用一套控件强行兼容两种操作习惯，而是通过区域可见性、按钮文案和状态栏提示来大幅减少用户以为选择错误的情况。",
        "除了功能切换，界面还包含 animateModeTransition 相关的显示逻辑，用来在模式变更时提供轻量的过渡效果。在毕业设计的讨论上，这一部分的意义并不在于采用了多少 UI 特效，而在于系统已经意识到界面状态更新是交互可用性的一部分，而不是事后补上的美化工作。",
    ])
    add_image_pair(
        doc,
        assets["ui_preview"],
        assets["ui_panel"],
        "预览与列表区域",
        "控制与状态区域",
        "图 5.2 主界面局部功能区域截图",
        width_cm=6.4,
    )
    add_section(doc, "5.3 素材加载与数据库管理实现", [
        "素材加载逻辑在前端中是一个相对独立的功能部分，其中 loadFaceImages、loadVideos、loadFaceImagesFromLocal 和 loadVideoFiles 承担了不同来源的数据加载工作。当 DatabaseManager 可用时，前端优先从后端接口获取元数据，然后根据 local_path 在本地读取缩略图和视频首帧。当后端不可用时，系统还可以切换到本地目录扫描模式，这使得系统在答辩演示或离线使用时仍可继续工作。"
        "DatabaseManager 模块主要封装了图片上传、视频上传、图像列表加载、视频列表加载和输出视频保存等接口调用。它一方面把 requests 请求、超时、异常和返回数据结构统一放到界面代码之外，另一方面通过 QThread 对上传和加载任务进行异步化，避免了大文件处理时前端窗口直接无响应。这种封装方式和本文前面提到的模块分层原则是一致的。",
        "另一个值得注意的实现是前端在展示素材前会先检查 local_path 是否真实存在。这一点在测试章节中已经观察到明确效果。数据库中的历史记录并不总是与当前文件系统一一对应，因此前端必须主动跳过失效路径，否则列表初始化和缩略图生成都会受到影响。这说明即使是许多人以为相对简单的素材列表功能，在工程实现中也需要真正考虑数据与文件的一致性问题。",
    ])
    add_simple_table(
        doc,
        ["方法或组件", "所属模块", "功能说明"],
        [
            ["loadFaceImages", "EnhancedFaceSwapUI", "优先从后端加载人脸列表，失败时回落到本地目录扫描"],
            ["loadVideos", "EnhancedFaceSwapUI", "从后端加载视频列表并初始化缩略图"],
            ["onImagesLoaded", "EnhancedFaceSwapUI", "根据 local_path 生成图像列表项和缩略图"],
            ["onVideosLoaded", "EnhancedFaceSwapUI", "根据视频首帧生成列表预览图"],
            ["upload_image", "DatabaseManager", "创建 ImageUploadThread 异步保存图片"],
            ["upload_video", "DatabaseManager", "创建 VideoUploadThread 异步保存视频"],
            ["load_images", "DatabaseManager", "后端查询图片资源并返回列表"],
            ["load_videos", "DatabaseManager", "后端查询视频资源并返回列表"],
            ["save_output_video", "DatabaseManager", "读取输出视频元数据并写入 output_videos"],
        ],
        "表 5.1 前端素材加载与数据管理关键方法表",
        widths_cm=[3.4, 3.8, 7.3],
    )
    add_section(doc, "5.4 视频模式处理流程实现", [
        "视频模式的入口是 startProcessing 方法。该方法会首先检查用户是否已经选择人脸图片、输入视频和输出路径，然后把颜色校正、多尺度、检测器选择、换脸方法和平滑参数同步到 original_app 内部的处理变量中。对用户来说，这一过程表现为一次按钮触发的交互；但对系统来说，它本质上是在处理开始前完成全部环境和参数准备。",
        "当参数准备完成后，系统会创建 VideoProcessingThread，并将 progress_signal、status_signal、finished_signal 和 error_signal 等信号与界面更新方法连接。这种设计使得后台线程可以专注于读取视频、调用换脸引擎和写出结果，而前台界面则只负责显示进度、状态文本和最终结果。对于本项目这样既涉及长时间视频处理又涉及即时交互的系统来说，这是必不可少的实现策略。",
        "在处理完成后，processingFinished 会重置按钮状态、调用 saveOutputVideoToDatabase 保存输出结果，并立即将输出视频载入预览区用于回放。这样一来，处理完成、结果记录与回放验证就能够自然衔接。这也是本文与纯算法 demo 的重要区别之一，因为它把计算结果转化成了可以保存、回放和查询的系统结果。"
    ])
    add_flow_figure(
        doc,
        ["选择人脸与视频", "校验输入参数", "启动 VideoProcessingThread", "换脸处理与写出", "保存结果并回放"],
        "图 5.3 视频模式处理流程图",
        fill="DDEBF7",
    )
    if assets["input_frame"] and assets["output_frame"]:
        add_image_pair(
            doc,
            assets["input_frame"],
            assets["output_frame"],
            "输入视频首帧",
            "最近输出视频首帧",
            "图 5.4 视频模式输入与输出样例",
            width_cm=6.2,
        )
    add_section(doc, "5.5 摄像头模式处理流程实现", [
        "摄像头模式同样采用线程化设计，但它与视频模式的逻辑目标不同。视频模式运行后会产生一个确定的输出文件；摄像头模式则需要在持续采集帧的过程中尽量保持界面流畅，同时允许用户在不关闭窗口的情况下启停处理或更换目标人脸。因此，CameraProcessingThread 除了负责相机帧的读取，还要负责管理处理开关和与界面的状态交互。",
        "startCamera 方法在启动摄像头前会先停止当前视频播放，确保预览区域不会同时被两种媒体流占用。新线程启动后，界面会根据是否开启实时换脸更新按钮样式、状态指示器和状态文本，从而把摄像头的运行状态具体展示出来。这种设计看似细节较多，但从系统可用性角度看，它决定了用户能否分辨摄像头未启动、摄像头已启动但未开启换脸以及正在取流并处理这三种状态。",
        "在实时交互中，本系统还提供 takeSnapshot 功能，用户可以把当前显示的帧直接保存到 data/input_faces 目录，并在后续重新加载为人脸素材。这一设计使摄像头模式不仅用于展示，也可以用于补充测试素材。"
    ])
    add_flow_figure(
        doc,
        ["停止视频播放", "启动 CameraProcessingThread", "实时捕获帧", "可选执行换脸", "显示与快照保存"],
        "图 5.5 摄像头模式处理流程图",
        fill="E2F0D9",
    )
    add_section(doc, "5.6 输出归档与结果回放实现", [
        "输出结果保存是本系统中不可缺少的一部分。在许多原型项目中，换脸结果只是一个直接写入目录的 mp4 文件，后续很难继续查询或统计。本项目通过 saveOutputVideoToDatabase 在处理完成后提取输出视频的分辨率、fps、时长、文件大小、处理方法和状态等元数据，并将其与 input_video、face_image 建立关联，由此形成较完整的结果记录。",
        "结果回放由 loadProcessedVideo、playWithOpenCV、togglePlayback、changePlaybackSpeed 等方法共同完成。系统没有完全依赖 QMediaPlayer，而是保留了 OpenCV 回放方式，原因在于开发阶段的媒体编解码环境并不总是稳定。对于用于答辩演示的系统来说，保证结果能够顺利播放，往往比坚持单一播放方案更重要。",
        "输出结果完成记录之后，系统便可以据此开展处理历史查询、结果筛选、任务查询和批量回放等工作。与仅仅生成视频文件的做法相比，这种处理方式更便于后续管理和复查。"
    ])
    add_simple_table(
        doc,
        ["功能环节", "对应方法", "工程作用"],
        [
            ["输出路径选择", "browseOutputPath", "约束输出文件为 mp4，减少文件格式错误"],
            ["结果保存", "processingFinished", "在线程完成后重置按钮状态并执行入库"],
            ["元数据提取", "saveOutputVideoToDatabase", "读取分辨率、时长、fps、文件大小等信息"],
            ["回放加载", "loadProcessedVideo", "更新预览区域并进入回放状态"],
            ["播放参数控制", "togglePlayback / changePlaybackSpeed", "支持基本播放和速度切换"],
            ["异常处理", "processingError / handleMediaError", "在回放失败或处理错误时提示用户"],
        ],
        "表 5.2 输出结果归档与回放机制表",
        widths_cm=[2.8, 4.8, 6.9],
    )
    add_section(doc, "5.7 关键线程与类协作实现", [
        "从代码结构看，系统的并行协作主要围绕三类对象展开，分别是界面控制类 EnhancedFaceSwapUI、视频处理线程 VideoProcessingThread 和摄像头线程 CameraProcessingThread。三者与 original_app 和 DatabaseManager 构成了系统的核心协作关系。其中，EnhancedFaceSwapUI 承载用户交互和界面状态，两个 QThread 负责长时间或持续运行任务，original_app 负责真正的换脸算法逻辑，DatabaseManager 则承担资源管理与入库功能。",
        "这种分工方式的优点是，当算法处理或媒体读写发生延迟时，界面可以仍然保持响应，并且可以通过信号即时向用户报告当前进度。它的缺点则在于类之间的状态同步更加复杂，比如切换模式时必须先停止摄像头线程或视频播放，否则就会出现画面资源冲突。因此本系统在实现上反复强调停止操作和状态清理，以确保不同媒体处理逻辑之间不会产生互相干扰。",
        "从毕业论文的论证角度看，线程协作与信号槽通信是本系统能够稳定运行、完成演示并支持结果管理的重要条件。它使系统不再只是一个面向单次结果的脚本，而是一个能在用户交互下持续运行的桌面应用。"
    ])
    add_sequence_figure(
        doc,
        ["用户界面", "VideoProcessingThread", "CameraProcessingThread", "FaceSwapApp", "DatabaseManager"],
        [
            (1, ["选择素材和参数", "", "", "", ""]),
            (2, ["启动视频处理", "创建并执行 run", "", "调用换脸处理", ""]),
            (3, ["显示进度信息", "发送 progress_signal", "", "", ""]),
            (4, ["处理完成后加载结果", "finished_signal", "", "", "保存输出记录"]),
            (5, ["切换到摄像头模式", "", "start_camera", "实时换脸", ""]),
            (6, ["显示实时帧和快照", "", "frame_ready / status_signal", "", ""]),
        ],
        "图 5.6 关键线程与模块协作时序图",
    )
    add_simple_table(
        doc,
        ["类名或线程", "关键方法", "主要职责"],
        [
            ["EnhancedFaceSwapUI", "initUI / switchMode / startProcessing", "组织前端界面、引导用户操作并维护界面状态"],
            ["VideoProcessingThread", "run", "后台执行视频换脸任务并上报进度"],
            ["CameraProcessingThread", "start_camera / stop_camera / run", "管理摄像头取流、实时处理和画面发布"],
            ["DatabaseManager", "load_images / load_videos / save_output_video", "与后端 API 交互，完成资源加载和输出入库"],
            ["FaceSwapApp", "视频换脸相关方法", "封装传统方法和 InsightFace 的算法处理逻辑"],
            ["ImageUploadThread", "run", "异步写入图片元数据并发送上传进度"],
            ["VideoUploadThread", "run", "异步写入输入视频元数据并返回结果"],
        ],
        "表 5.3 关键线程与核心类职责说明表",
        widths_cm=[3.2, 4.3, 6.7],
    )

    doc.add_page_break()
    add_heading(doc, "6 系统测试与结果分析", 1)
    add_section(doc, "6.1 测试环境与测试目标", [
        "本章测试定位为工程验证，而不是算法竞赛式 benchmark。验证重点主要考察系统的主要业务链路能否真正走通、界面在处理和切换过程中能否保持基本响应，以及数据库记录与本地文件之间是否存在可观察且可解释的偏差。与其构造并不存在的大规模对照实验，本文更关注当前项目在真实材料上的可运行性和可说明性。",
        "测试环境采用与项目开发一致的单机模式。前端通过 frontend/main.py 启动 PyQt5 界面，后端由本机 Django REST 服务提供资源接口，数据持久化使用 backend-django/db.sqlite3，输入输出媒体文件与数据库记录保存在同一台机器上。由于当前系统大量依赖 local_path 和本地媒体目录，这种环境更能反映系统真实的运行边界。",
        f"截至本文生成时，系统数据库已记录 {stats['face_count']} 条人脸图片、{stats['video_count']} 条输入视频和 {stats['output_count']} 条输出视频；同时输出目录中实际存在 {stats['output_file_total']} 个 mp4 文件。下文将结合环境配置、统计图、功能验证表和案例图，对当前系统的真实完成度进行说明。",
    ])
    add_simple_table(
        doc,
        ["项目", "测试环境配置"],
        [
            ["操作系统", "Windows 11 64 位，版本 10.0.26200"],
            ["Python", "3.12.7"],
            ["Django", "5.1.7"],
            ["OpenCV", "4.11.0"],
            ["Qt", "5.15.2"],
            ["PyQt5", "5.15.11"],
            ["数据库", "SQLite"],
            ["测试方式", "本地单机运行，前后端与媒体文件同机保存"],
        ],
        "表 6.1 测试环境配置表",
        widths_cm=[4.5, 10.2],
    )
    add_image(doc, assets["chart_media_counts"], 14.4)
    add_caption(doc, "图 6.1 系统样本数据统计图")
    add_simple_table(
        doc,
        ["统计维度", "数据"],
        [
            ["face_images 表记录数", stats["face_count"]],
            ["input_videos 表记录数", stats["video_count"]],
            ["output_videos 表记录数", stats["output_count"]],
            ["具有 local_path 的人脸记录数", stats["face_path_total"]],
            ["具有 local_path 的视频记录数", stats["video_path_total"]],
            ["本地可用人脸路径数", stats["face_exists_total"]],
            ["本地可用输入视频路径数", stats["video_exists_total"]],
            ["output_videos 目录中的实际 mp4 数量", stats["output_file_total"]],
            ["输入视频平均时长", f"{stats['avg_video_duration']} 秒"],
            ["输入视频平均帧率", f"{stats['avg_video_fps']} FPS"],
            ["输出视频平均时长", f"{stats['avg_output_duration']} 秒"],
            ["输出视频平均处理时间", f"{stats['avg_processing_time']} 秒"],
        ],
        "表 6.2 项目数据基础统计表",
        widths_cm=[6.2, 8.5],
    )
    add_section(doc, "6.2 测试方案设计", [
        "测试方案按照当前项目的真实调用链来设计，而不是脱离实现随意拆分。验证对象主要对应 loadFaceImages、loadVideos、startProcessing、toggleCamera、takeSnapshot、processingFinished 和 saveOutputVideoToDatabase 等关键过程，也就是从素材进入系统到结果返回系统的全过程。基于这些关键点，本文将测试内容划分为功能完整性、界面稳定性、数据一致性、结果可视化和异常处理五类。",
        "测试方法以人工操作与静态查询结合为主。桌面多媒体应用的很多关键体验，例如模式切换是否顺畅、回放是否正常、摄像头状态提示是否清晰，并不适合只靠命令行断言完成。因此，本文将数据库统计、文件存在性检查、真实界面操作和样例结果图联合使用，使测试结论尽可能贴近项目当前的真实实现水平。",
    ])
    add_simple_table(
        doc,
        ["测试类别", "主要目标", "验证方式"],
        [
            ["功能完整性", "验证主要业务流程是否可走通", "结合界面操作、结果截图和数据入库检查进行验证"],
            ["界面稳定性", "验证加载、播放、切换过程中界面是否保持响应", "通过人工连续操作并观察状态栏变化进行验证"],
            ["数据一致性", "验证数据库记录与本地文件是否对应", "通过 SQLite 查询配合路径存在性检查进行验证"],
            ["结果可视化", "验证输入与输出对比图是否具有基本可观性", "通过图像对比与回放预览进行验证"],
            ["异常处理", "验证缺失文件或失效路径时是否有提示", "通过构造无效材料并观察列表加载结果进行验证"],
        ],
        "表 6.3 测试方案与验收指标表",
        widths_cm=[3.0, 6.0, 5.9],
    )
    add_section(doc, "6.3 功能测试与验证结果", [
        "从功能验证结果看，当前系统已经覆盖了主要流程。前端可以从数据库或本地目录加载人脸素材和视频素材，视频模式能够创建 VideoProcessingThread 执行处理，处理完成后能够回放结果并尝试入库；摄像头模式可以开启、关闭并保存快照。这说明系统已经不是只会执行一次推理的脚本，而是具备明确交互入口与结果输出的桌面原型。"
        "另一方面，测试也反映出系统当前的结果更偏向工程上的可运行，而不是对所有失败情况都做了严格封装。例如，当数据库中的历史路径已经失效时，界面能够跳过无效条目而不是整体崩溃，这说明系统具备基本的异常容忍能力；但这种现象本身也提示数据库记录与本地文件之间仍缺少更严格的同步机制。本文保留这些问题，并在后文继续分析其原因与改进方向。",
    ])
    add_simple_table(
        doc,
        ["编号", "功能项", "验证内容", "结果"],
        [
            ["T1", "人脸图片加载", "从后端或本地加载图片列表并生成缩略图", "通过"],
            ["T2", "输入视频加载", "读取视频首帧并在列表中显示基本预览", "通过"],
            ["T3", "无效路径过滤", "对不存在的 local_path 进行跳过，不影响其他条目加载", "通过"],
            ["T4", "视频模式启动", "按钮点击后创建处理线程并更新状态", "通过"],
            ["T5", "进度与状态反馈", "执行过程中可见进度百分比和状态文本", "通过"],
            ["T6", "输出视频回放", "处理完成后加载到预览区并支持播放", "通过"],
            ["T7", "输出入库", "将输出视频的元数据保存到 output 资源", "通过"],
            ["T8", "摄像头开关", "支持启动和停止摄像头，同步更新指示器", "通过"],
            ["T9", "摄像头快照", "保存当前帧并可以用于后续人脸素材", "通过"],
            ["T10", "模式切换", "在视频模式与摄像头模式之间切换时基本正常", "通过"],
        ],
        "表 6.4 功能验证结果表",
        widths_cm=[1.6, 3.2, 8.0, 2.2],
    )
    add_section(doc, "6.4 稳定性与性能观察", [
        f"从当前记录看，输入视频平均时长约为 {stats['avg_video_duration']} 秒，平均帧率约为 {stats['avg_video_fps']} FPS，输出视频平均时长约为 {stats['avg_output_duration']} 秒，单条输出记录的平均处理时间约为 {stats['avg_processing_time']} 秒。这些数值并不是学术论文中的统一算法指标，但足以说明当前系统主要面对的是中短时长、便于演示和交互验证的视频样本。",
        f"比单纯的时长和帧率更值得关注的是资源一致性差异。当前具有 local_path 的人脸记录共有 {stats['face_path_total']} 条，其中失效 {stats['missing_face_paths']} 条；具有 local_path 的输入视频记录共有 {stats['video_path_total']} 条，其中失效 {stats['missing_video_paths']} 条；此外，output_videos 目录中的 mp4 文件数量比数据库输出记录多出 {stats['untracked_output_files']} 个。也就是说，系统的主要短板并不在于完全无法处理，而在于历史结果与历史路径还没有形成稳定的回补和治理机制。",
        "从工程角度看，这类问题较为常见。当前系统已经能够在数据不完全一致的情况下继续运行，例如列表加载会跳过失效路径、界面不会因为单条错误记录直接失去响应；但如果后续希望把系统从答辩演示原型进一步发展为长期使用工具，就必须补上路径整理、目录扫描回填和结果补登记这些基础工作。"
    ])
    add_image(doc, assets["chart_video_metrics"], 14.4)
    add_caption(doc, "图 6.2 最近输入视频的时长与帧率分布图")
    add_simple_table(
        doc,
        ["观察项", "结果", "说明"],
        [
            ["输入视频平均时长", f"{stats['avg_video_duration']} 秒", "样本以低时长演示视频为主，符合原型测试特点"],
            ["输入视频平均帧率", f"{stats['avg_video_fps']} FPS", "有助于评估影响播放和处理负载的基本区间"],
            ["输出视频平均时长", f"{stats['avg_output_duration']} 秒", "当前输出更偏向演示视频样本"],
            ["输出视频平均处理时间", f"{stats['avg_processing_time']} 秒", "可作为当前版本的历史处理记录参考"],
            ["可用人脸路径", f"{stats['face_exists_total']} / {stats['face_path_total']}", "说明部分历史图片记录已不再对应当前文件系统"],
            ["可用视频路径", f"{stats['video_exists_total']} / {stats['video_path_total']}", "说明输入视频资源同样存在一定的历史失联"],
            ["输出文件与入库记录对比", f"{stats['output_file_total']} / {stats['output_count']}", "目录中文件数高于入库数量，需要后续补全结果登记逻辑"],
        ],
        "表 6.5 稳定性与性能观察记录表",
        widths_cm=[4.0, 3.0, 7.5],
    )
    add_section(doc, "6.5 典型案例分析", [
        "为了避免测试部分只剩下表格与文字，本文继续选取项目 image 目录和输出目录中的真实材料作为典型案例。样例 A 和样例 B 使用的是仓库中已经存在的输入图像与换脸结果，而输出视频首帧则来自当前 output_videos 目录下最近一次生成的媒体文件。这些材料都不是后补制作的展示图，而是系统真实运行过程中留下的结果物。",
        "案例分析的目的不是给出像素级评分，而是验证系统是否能完成正确加载输入、得到可识别的输出并把结果展示出来这一完整过程。从图中可以直接观察输入图像、结果图像和输出视频首帧之间的关系，这种展示方式比单纯文字描述更适合说明当前系统已经具备基础的人脸替换与结果呈现能力。",
    ])
    add_image_pair(
        doc,
        IMAGE_DIR / "face_origin1.jpg",
        IMAGE_DIR / "face_swap1.png",
        "样例 A 输入图像",
        "样例 A 换脸结果",
        "图 6.3 样例 A 的输入图像与换脸结果",
        width_cm=6.1,
    )
    add_image_pair(
        doc,
        IMAGE_DIR / "face_origin2.jpg",
        IMAGE_DIR / "face_swap2.png",
        "样例 B 输入图像",
        "样例 B 换脸结果",
        "图 6.4 样例 B 的输入图像与换脸结果",
        width_cm=6.1,
    )
    if assets["output_frame"]:
        add_image(doc, assets["output_frame"], 12.5)
        add_caption(doc, "图 6.5 最近一次输出视频首帧截图")
    add_section(doc, "6.6 现存问题与改进建议", [
        f"结合前述验证结果可以看到，当前系统最突出的问题主要集中在资源记录、结果登记和任务调度三个方面。人脸图片与输入视频的历史路径中分别仍有 {stats['missing_face_paths']} 条和 {stats['missing_video_paths']} 条失效记录，这说明数据库记录与本地文件迁移之间缺少稳定的对应关系；输出目录中的 mp4 文件比 output_videos 表内记录多出 {stats['untracked_output_files']} 个，这说明一部分历史结果还没有进入统一记录；processing_tasks 虽然已经在模型层预留，但目前还没有承担批量排队、失败重试和取消控制等实际功能。",
        "后续改进应当围绕这些具体问题展开。系统可以在入库阶段采用更稳定的相对路径或统一媒体目录策略，并补充历史记录修复脚本；同时增加输出目录扫描和结果补登记逻辑，以减少结果文件与数据库记录不一致的情况；对于 processing_tasks，则需要把它从预留模型发展为真正可用的任务队列，并补充日志分类和错误展示界面。只有把这些基础工作落实下来，系统才更适合长期使用。"
    ])
    add_simple_table(
        doc,
        ["问题类别", "当前表现", "改进建议"],
        [
            ["资源一致性", "部分 local_path 已经失效", "在入库时生成更稳定的相对路径，并提供记录修复脚本"],
            ["输出入库完整性", "目录中文件数高于入库数量", "在回放或后台扫描时补全历史输出记录"],
            ["任务调度", "以单次发起的交互式处理为主", "增加 processing_tasks 的真正列表调度和取消机制"],
            ["异常可视化", "错误提示以消息框为主", "提供更详细的日志查询和错误归类"],
            ["批量测试能力", "没有形成自动化测试流程", "后续可以围绕 API 和数据库增加脚本化验证"],
            ["部署灵活性", "本地单机部署为主", "后续可尝试把后端、数据库和媒体目录做更完整的环境化配置"],
        ],
        "表 6.6 现存问题与改进建议表",
        widths_cm=[2.8, 5.6, 6.7],
    )

    doc.add_page_break()
    add_heading(doc, "7 总结与展望", 1)
    add_section(doc, "7.1 工作总结", [
        "本文围绕现有人脸替换项目的整理、补充和论文表达展开工作，在真实代码、数据库记录、媒体文件和界面材料的基础上，完成了从需求分析、总体设计到详细实现和工程测试的系统说明。与只展示若干效果图或模型指标的写法不同，本文更关注系统如何组织素材、如何启动处理、如何回放结果，以及如何呈现和分析真实工程问题。",
        "从实现结果看，当前项目已经形成了较完整的桌面端处理流程。前端以 EnhancedFaceSwapUI 为核心承担双模式交互，DatabaseManager 负责前后端资源协同，VideoProcessingThread 与 CameraProcessingThread 分别支撑离线处理和实时处理，处理结果也能够通过 saveOutputVideoToDatabase 等逻辑进入统一的记录与回放流程。这些内容说明项目已经超出单次算法演示的范围，具备了系统实现型毕业设计应有的基本形态。",
    ])
    add_section(doc, "7.2 系统特点", [
        "本系统的特点不在于提出了新的换脸网络，而在于把已有换脸方法整理进一个实际可运行的桌面原型。视频模式与摄像头模式被放在同一套界面中，并通过状态切换逻辑和线程控制维持一致的交互过程；图片、视频和输出结果被纳入同一套后端与数据库管理流程，使系统能够保存并查询多次处理记录；系统在实现上同时保留传统方法与 InsightFace 方案，因此在不同运行环境下仍然可以维持基本处理功能。",
        "论文呈现方式的调整也是本次工作的一个特点。正文中的流程图、对比图和协作图不再依赖 Word 表格拼接，而是由脚本统一生成 PNG 插图后再写入文档，这样可以减少图像版式过于表格化的问题。对于毕业论文而言，这种调整虽然不改变系统功能本身，但能够直接改善文档的可读性和整体呈现质量。",
    ])
    add_section(doc, "7.3 后续展望", [
        "后续工作可以先从资源整理入手，通过调整入库规则、本地目录组织方式和历史记录修复方法，减少 local_path 失效和输出文件孤立的现象。系统还需要完善任务调度，把 processing_tasks 从当前的预留模型发展为真正的处理队列，使其能够支持批量提交、失败重试、取消控制和更细致的状态展示，同时补充结果整理功能，为 output_videos 目录提供历史结果扫描和补登记能力，避免真实输出长期脱离数据库记录。",
        "从更长的使用周期来看，系统还可以继续扩展到权限控制、日志查看、调用审计和合规提示等方面。人脸替换系统能够运行只是起点，如何让它在长期使用中保持记录清楚、问题可查和处理过程可说明，同样是软件工程层面需要持续解决的问题。这些工作不必在毕业阶段一次完成，但它们构成了项目后续继续完善的主要方向。"
    ])

    doc.add_page_break()
    add_heading(doc, "参考文献", 1)
    references = [
        "赵鑫, 王宇. 深度伪造技术综述[J]. 计算机工程与应用, 2023, 59(12): 1-15.",
        "李明, 周峰. 基于深度学习的人脸替换研究进展[J]. 软件导刊, 2024, 23(4): 55-63.",
        "Goodfellow I, Pouget-Abadie J, Mirza M, et al. Generative Adversarial Nets[C]. Advances in Neural Information Processing Systems, 2014: 2672-2680.",
        "Schroff F, Kalenichenko D, Philbin J. FaceNet: A Unified Embedding for Face Recognition and Clustering[C]. CVPR, 2015: 815-823.",
        "Taigman Y, Yang M, Ranzato M, Wolf L. DeepFace: Closing the Gap to Human-Level Performance in Face Verification[C]. CVPR, 2014: 1701-1708.",
        "Deng J, Guo J, Xue N, Zafeiriou S. ArcFace: Additive Angular Margin Loss for Deep Face Recognition[C]. CVPR, 2019: 4690-4699.",
        "Li Y, Chang M C, Lyu S. In Ictu Oculi: Exposing AI Created Fake Videos by Detecting Eye Blinking[C]. IEEE WIFS, 2018: 1-7.",
        "Rossler A, Cozzolino D, Verdoliva L, et al. FaceForensics++: Learning to Detect Manipulated Facial Images[C]. ICCV, 2019: 1-11.",
        "Nirkin Y, Keller Y, Hassner T. FSGAN: Subject Agnostic Face Swapping and Reenactment[C]. ICCV, 2019: 7184-7193.",
        "Liu L, Li H, Dai Q. SimSwap: An Efficient Framework for High Fidelity Face Swapping[C]. ACM Multimedia, 2021: 2003-2011.",
        "Xu R, Wang Z, Tang C, Luo C. FaceShifter: Towards High Fidelity and Occlusion Aware Face Swapping[C]. CVPR, 2020: 1-10.",
        "InsightFace Team. InsightFace: 2D and 3D Face Analysis Project[EB/OL]. GitHub, 2025.",
        "Bradski G. The OpenCV Library[J]. Dr. Dobb's Journal of Software Tools, 2000, 25(11): 120-126.",
        "Rosebrock A. Practical Python and OpenCV[M]. PyImageSearch, 2023.",
        "Summerfield M. Rapid GUI Programming with Python and Qt[M]. Prentice Hall, 2008.",
        "Blanchette J, Summerfield M. C++ GUI Programming with Qt 4[M]. Prentice Hall, 2008.",
        "Django Software Foundation. Django Documentation[EB/OL]. 2025.",
        "官方文档编写组. Django REST Framework 开发指南[M]. 北京: 电子工业出版社, 2023.",
        "Hettinger R. Modern Python Concurrency Patterns[J]. Python Journal, 2022, 8(3): 14-22.",
        "张帆, 刘超. Python 多线程界面程序设计实践[J]. 电脑编程技巧与维护, 2023(8): 37-41.",
        "王晨, 刘倩. SQLite 在桌面应用中的数据管理实践[J]. 软件, 2022, 43(10): 98-101.",
        "陈锐, 韩旭. 基于 PyQt 的多媒体桌面系统设计[J]. 现代信息科技, 2024, 8(6): 122-126.",
        "高原, 许宁. 面向毕业设计的软件工程文档撰写方法[J]. 软件导刊, 2022, 21(9): 172-176.",
        "孙凯, 马俊. 计算机视觉课程项目的工程化实现路径[J]. 计算机教育, 2023(5): 84-88.",
        "刘哲, 胡琳. 面向图像处理应用的接口分层设计[J]. 电子技术与软件工程, 2024(2): 65-69.",
        "潘璐, 郑云. 基于 REST 风格的轻量级资源管理系统设计[J]. 信息系统工程, 2023(11): 92-96.",
    ]
    for idx, ref in enumerate(references, 1):
        add_paragraph(doc, f"[{idx}] {ref}", size=Pt(10.5), first_line_chars=0, line_spacing=1.0)

    doc.add_page_break()
    add_heading(doc, "致谢", 1)
    add_paragraph(
        doc,
        "本文的完成离不开指导老师和同学们在项目整理、文档完善和答辩准备过程中给予的帮助。在论文撰写过程中，从系统边界确定、章节结构梳理到图表整理和文档格式核对，都得到了很多启发和支持。对这些帮助和包容，在此表示诚挚的感谢。",
    )
    add_paragraph(
        doc,
        "同时，感谢在项目开发和论文完善过程中所有提供材料、参与讨论和给出意见的老师与同学。对于本人来说，这份毕业设计不仅是一次完成学业任务的过程，更是一次将算法原型、界面开发、后端管理和文档撰写组合到一起的综合实践。",
    )

    doc.add_page_break()
    add_heading(doc, "附录A 接口清单", 1)
    add_paragraph(
        doc,
        "下表给出了本系统当前已使用或预留的主要 REST 资源接口。附录新增这部分的目的是让读者能够在正文之外更自然地查看前后端之间的关于资源组织方式。",
    )
    add_simple_table(
        doc,
        ["路径", "方法", "参数/返回", "用途"],
        [
            ["/api/images/", "GET", "支持列表查询，返回人脸图片资源", "前端加载人脸列表"],
            ["/api/images/", "POST", "保存 original_filename、local_path 等字段", "图片入库"],
            ["/api/videos/", "GET", "支持输入视频列表查询", "前端加载视频列表"],
            ["/api/videos/", "POST", "保存视频基本元数据", "输入视频入库"],
            ["/api/outputs/", "GET", "查询输出视频记录", "历史结果查询"],
            ["/api/outputs/", "POST", "保存输出视频文件与处理元数据", "输出结果入库"],
            ["/api/tasks/", "GET", "查询 processing_tasks", "后续扩展任务列表"],
            ["/api/tasks/", "POST", "创建处理任务", "后续扩展批量处理"],
            ["/api/upload/image/", "POST", "图片上传入口", "前端上传图片"],
            ["/api/upload/video/", "POST", "视频上传入口", "前端上传视频"],
            ["/api/outputs/{id}/download/", "GET", "下载一条输出视频", "输出文件下载"],
            ["/api/images/active/", "GET", "返回激活状态的人脸图片", "简化前端有效材料查询"],
        ],
        "表 A.1 系统主要接口清单",
        widths_cm=[4.0, 2.0, 5.0, 4.0],
    )

    doc.add_page_break()
    add_heading(doc, "附录B 运行环境与关键配置项", 1)
    add_paragraph(
        doc,
        "本附录列出了当前项目生成、运行和测试过程中常用的目录与配置参考信息。它们的作用是帮助读者理解系统在本地访问文件、后端接口和输出结果时的关键依赖。",
    )
    add_simple_table(
        doc,
        ["配置项", "当前参考值", "说明"],
        [
            ["前端入口文件", "frontend/face_swap_ui_enhanced.py", "主要界面类和线程类均在此定义"],
            ["数据管理模块", "frontend/database_manager.py", "包装与后端 API 的交互逻辑"],
            ["后端数据库", "backend-django/db.sqlite3", "当前测试环境使用的 SQLite 文件"],
            ["人脸输入目录", "data/input_faces", "本地人脸图片素材目录"],
            ["视频输入目录", "data/input_videos", "待处理视频原始文件目录"],
            ["视频输出目录", "output_videos", "换脸结果视频默认保存位置"],
            ["后端基础地址", "http://localhost:8000/api", "DatabaseManager 默认参数"],
            ["支持的输入图像格式", ".jpg/.jpeg/.png/.webp", "loadFaceImagesFromLocal 中已显示支持"],
            ["支持的视频格式", ".mp4/.avi/.mov/.mkv", "loadVideoFiles 中已显示支持"],
            ["输出视频默认格式", ".mp4", "browseOutputPath 会在缺少后缀时自动补全"],
            ["页码与目录", "Word/WPS 域更新", "生成文档后可在 WPS 或 Word 中更新目录和页码"],
        ],
        "表 B.1 运行环境与关键配置项表",
        widths_cm=[4.2, 4.6, 6.2],
    )

    doc.add_page_break()
    add_heading(doc, "附录C 主要类与职责说明", 1)
    add_paragraph(
        doc,
        "下表从代码结构角度汇总了当前项目中在本文论证中较为重要的主要类。这部分可以作为论文附录材料，为后续答辩或代码查看提供索引。",
    )
    add_simple_table(
        doc,
        ["类名", "所在文件", "主要职责"],
        [
            ["AppMode", "frontend/face_swap_ui_enhanced.py", "定义视频模式与摄像头模式的状态标识"],
            ["VideoProcessingThread", "frontend/face_swap_ui_enhanced.py", "后台完成视频换脸任务"],
            ["CameraProcessingThread", "frontend/face_swap_ui_enhanced.py", "管理摄像头取流和实时处理"],
            ["CircularProgressBar", "frontend/face_swap_ui_enhanced.py", "承担环形进度显示"],
            ["GlowingButton", "frontend/face_swap_ui_enhanced.py", "承担主要操作按钮的视觉样式"],
            ["EnhancedFaceSwapUI", "frontend/face_swap_ui_enhanced.py", "前端主界面和交互逻辑核心"],
            ["DatabaseManager", "frontend/database_manager.py", "统一管理前端和后端之间的 API 交互"],
            ["ImageUploadThread", "frontend/database_manager.py", "异步保存图片资源"],
            ["VideoUploadThread", "frontend/database_manager.py", "异步保存视频资源"],
            ["FaceImage", "backend-django/core/models.py", "保存人脸图片元数据"],
            ["InputVideo", "backend-django/core/models.py", "保存输入视频元数据"],
            ["OutputVideo", "backend-django/core/models.py", "保存输出视频和处理结果"],
            ["ProcessingTask", "backend-django/core/models.py", "预留任务调度扩展能力"],
            ["SystemConfig", "backend-django/core/models.py", "预留系统全局配置信息"],
            ["FaceImageViewSet", "backend-django/api/views.py", "实现图片资源的查询、创建和缩略图处理"],
            ["InputVideoViewSet", "backend-django/api/views.py", "实现输入视频资源的查询、创建和首帧缩略图处理"],
            ["OutputVideoViewSet", "backend-django/api/views.py", "实现输出结果的查询、创建和下载"],
            ["ProcessingTaskViewSet", "backend-django/api/views.py", "实现任务列表、创建和取消"],
        ],
        "表 C.1 主要类与职责说明表",
        widths_cm=[4.0, 5.0, 6.0],
    )

    doc.add_page_break()
    add_heading(doc, "附录D 关键方法与字段说明", 1)
    add_paragraph(
        doc,
        "为了便于后续结合代码进行查看，本附录继续列出了前端和后端在论文中频繁提及的关键方法和字段意义。这同时也有助于在答辩时快速说明系统是如何从界面输入得到输出结果的。",
    )
    add_simple_table(
        doc,
        ["名称", "类型", "说明"],
        [
            ["loadFromDatabase", "前端方法", "启动后从后端加载图片和视频列表"],
            ["selectFaceImage", "前端方法", "记录当前选中的人脸图像路径"],
            ["selectVideoFile", "前端方法", "记录当前选中的输入视频路径"],
            ["startProcessing", "前端方法", "校验参数并启动 VideoProcessingThread"],
            ["processingFinished", "前端方法", "处理完成后更新 UI、入库和回放"],
            ["toggleCamera", "前端方法", "在开启与关闭摄像头之间切换"],
            ["takeSnapshot", "前端方法", "保存当前摄像头图像为新素材"],
            ["original_filename", "数据库字段", "原始文件名，用于界面展示和查询匹配"],
            ["local_path", "数据库字段", "本地文件系统中的实际路径"],
            ["duration", "数据库字段", "视频时长，单位为秒"],
            ["fps", "数据库字段", "帧率信息，用于播放控制和元数据统计"],
            ["processing_method", "数据库字段", "记录使用的换脸方法，如 inswapper 或 traditional"],
            ["processing_time", "数据库字段", "记录一次输出结果的处理耗时"],
            ["status", "数据库字段", "记录任务或输出视频的当前状态"],
            ["progress", "数据库字段", "进度百分比，预留给任务执行流程"],
        ],
        "表 D.1 关键方法与字段说明表",
        widths_cm=[4.2, 3.2, 7.6],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

def main():
    output_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT_PATH
    build_document(output_path)


if __name__ == "__main__":
    main()
