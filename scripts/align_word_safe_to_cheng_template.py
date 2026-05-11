from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Emu, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
SRC_DOCX = ROOT / (
    "\u6362\u8138\u7cfb\u7edf_\u521d\u7a3f_\u5b66\u672f\u5316\u7248_"
    "\u7ec8\u7a3f\u7edf\u4e00\u6574\u6539\u7248_Word\u516c\u5f0f\u5b89\u5168\u7248.docx"
)
OUT_DOCX = ROOT / (
    "\u6362\u8138\u7cfb\u7edf_\u521d\u7a3f_\u5b66\u672f\u5316\u7248_"
    "\u7ec8\u7a3f\u7edf\u4e00\u6574\u6539\u7248_Word\u516c\u5f0f\u5b89\u5168\u7248_"
    "\u7a0b\u6cfd\u4e2d\u6a21\u677f\u5bf9\u9f50\u7248.docx"
)

CHINESE_NUMERALS = {
    1: "\u4e00",
    2: "\u4e8c",
    3: "\u4e09",
    4: "\u56db",
    5: "\u4e94",
    6: "\u516d",
    7: "\u4e03",
    8: "\u516b",
    9: "\u4e5d",
    10: "\u5341",
}

REFERENCE_HEADING = "\u53c2\u8003\u6587\u732e"
ACK_HEADING = "\u81f4\u8c22"
FRONT_TITLES = {"\u6458    \u8981", "ABSTRACT", "\u76ee    \u5f55"}
BODY_HEADING_1_RE = re.compile(r"^(\d+)\s+(.+)$")
BODY_HEADING_2_RE = re.compile(r"^\d+\.\d+\s+.+$")
BODY_HEADING_3_RE = re.compile(r"^\d+\.\d+\.\d+\s+.+$")
CAPTION_RE = re.compile(r"^([\u56fe\u8868])\s*([0-9]+)\.([0-9]+)\s*(.+)$")
REF_ENTRY_RE = re.compile(r"^\[\d+\]")

NEW_ABSTRACT = (
    "\u672c\u6587\u56f4\u7ed5\u684c\u9762\u7aef\u4eba\u8138\u66ff\u6362\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5b9e\u73b0"
    "\u5c55\u5f00\u7814\u7a76\uff0c\u9762\u5411\u4eba\u8138\u66ff\u6362\u5e94\u7528\u5728\u4ea4\u4e92\u6027\u3001\u5904\u7406\u6548\u7387"
    "\u548c\u5de5\u7a0b\u53ef\u7528\u6027\u65b9\u9762\u7684\u9700\u6c42\uff0c\u5728\u5df2\u6709\u7b97\u6cd5\u80fd\u529b\u57fa\u7840\u4e0a"
    "\u5b8c\u6210\u7cfb\u7edf\u5316\u96c6\u6210\u4e0e\u529f\u80fd\u5b9e\u73b0\u3002\u7cfb\u7edf\u91c7\u7528\u524d\u540e\u7aef\u5206\u5c42"
    "\u8bbe\u8ba1\uff0c\u4ee5 PyQt5 \u6784\u5efa\u684c\u9762\u4ea4\u4e92\u754c\u9762\uff0c\u7ed3\u5408 Django REST \u5b8c\u6210\u6570\u636e"
    "\u7ba1\u7406\u4e0e\u670d\u52a1\u652f\u6491\uff0c\u5e76\u4f9d\u6258 OpenCV \u4e0e InsightFace \u5b9e\u73b0\u89c6\u9891\u6362\u8138"
    "\u548c\u6444\u50cf\u5934\u5b9e\u65f6\u6362\u8138\u4e24\u7c7b\u6838\u5fc3\u4e1a\u52a1\u3002\u56f4\u7ed5\u7d20\u6750\u5bfc\u5165\u3001"
    "\u53c2\u6570\u8bbe\u7f6e\u3001\u4efb\u52a1\u542f\u52a8\u3001\u7ed3\u679c\u4fdd\u5b58\u4e0e\u56de\u653e\u7b49\u73af\u8282\uff0c\u672c\u6587"
    "\u5bf9\u7cfb\u7edf\u9700\u6c42\u5206\u6790\u3001\u603b\u4f53\u67b6\u6784\u3001\u5173\u952e\u6a21\u5757\u8bbe\u8ba1\u4e0e\u754c\u9762\u4ea4"
    "\u4e92\u6d41\u7a0b\u8fdb\u884c\u4e86\u7cfb\u7edf\u9610\u8ff0\uff0c\u5e76\u9488\u5bf9\u7ebf\u7a0b\u8c03\u5ea6\u3001\u5a92\u4f53\u8d44\u6e90"
    "\u7ba1\u7406\u548c\u7ed3\u679c\u5f52\u6863\u7b49\u5de5\u7a0b\u95ee\u9898\u7ed9\u51fa\u76f8\u5e94\u5b9e\u73b0\u65b9\u6848\u3002\u901a\u8fc7"
    "\u529f\u80fd\u6d4b\u8bd5\u4e0e\u8fd0\u884c\u9a8c\u8bc1\u53ef\u4ee5\u770b\u51fa\uff0c\u6240\u8bbe\u8ba1\u7cfb\u7edf\u80fd\u591f\u8f83\u4e3a"
    "\u5b8c\u6574\u5730\u652f\u6301\u684c\u9762\u7aef\u4eba\u8138\u66ff\u6362\u7684\u4e3b\u8981\u5e94\u7528\u6d41\u7a0b\uff0c\u5177\u6709\u8f83\u597d"
    "\u7684\u4ea4\u4e92\u4f53\u9a8c\u4e0e\u5b9e\u9645\u6f14\u793a\u80fd\u529b\uff0c\u53ef\u4e3a\u76f8\u5173\u56fe\u50cf\u5904\u7406\u7cfb\u7edf\u7684"
    "\u5de5\u7a0b\u5316\u5b9e\u73b0\u63d0\u4f9b\u53c2\u8003\u3002"
)

NEW_ENGLISH_ABSTRACT = (
    "This thesis focuses on the design and implementation of a desktop face-swapping system. "
    "Instead of proposing a brand-new swapping model, the study targets the practical needs of "
    "interactivity, processing efficiency, and engineering usability, and integrates existing "
    "algorithmic capabilities into a coherent software system. The system adopts a layered "
    "frontend-backend architecture: PyQt5 is used to build the desktop user interface, Django REST "
    "supports data management and service coordination, and OpenCV together with InsightFace enables "
    "two core workflows, namely video face swapping and real-time camera face swapping. Around material "
    "import, parameter configuration, task execution, result storage, and replay, the thesis presents "
    "the requirement analysis, overall architecture, key module design, and interaction process of the "
    "system, while also addressing engineering issues such as thread scheduling, media resource management, "
    "and result archiving. Functional testing and runtime verification show that the system can support "
    "the main workflow of desktop face swapping with solid interaction quality and practical demonstration "
    "value, providing a useful reference for the engineering implementation of related image-processing systems."
)

NEW_ACKNOWLEDGEMENTS = [
    "\u9996\u5148\uff0c\u8c28\u5411\u6307\u5bfc\u6559\u5e08\u4ee5\u53ca\u5728\u8bfe\u7a0b\u5b66\u4e60\u4e2d\u7ed9\u4e88\u6211\u5e2e\u52a9\u7684"
    "\u5404\u4f4d\u8001\u5e08\u81f4\u4ee5\u8bda\u631a\u7684\u8c22\u610f\u3002\u4ece\u8bba\u6587\u9009\u9898\u3001\u7ed3\u6784\u8bbe\u8ba1\u5230"
    "\u7cfb\u7edf\u5b9e\u73b0\u4e0e\u8bba\u6587\u4fee\u6539\uff0c\u5404\u4f4d\u8001\u5e08\u90fd\u7ed9\u4e88\u4e86\u8010\u5fc3\u6307\u5bfc\u4e0e"
    "\u5b9d\u8d35\u5efa\u8bae\uff0c\u4f7f\u6211\u80fd\u591f\u4e0d\u65ad\u5b8c\u5584\u7814\u7a76\u601d\u8def\u548c\u8868\u8fbe\u65b9\u5f0f\u3002",
    "\u5176\u6b21\uff0c\u611f\u8c22\u8eab\u8fb9\u7684\u540c\u5b66\u4eec\u5728\u9879\u76ee\u5f00\u53d1\u3001\u529f\u80fd\u8c03\u8bd5\u3001\u8d44\u6599"
    "\u6574\u7406\u548c\u7b54\u8fa9\u51c6\u5907\u8fc7\u7a0b\u4e2d\u7ed9\u4e88\u7684\u652f\u6301\u4e0e\u5e2e\u52a9\u3002\u4e0e\u5927\u5bb6\u7684\u4ea4"
    "\u6d41\u8ba8\u8bba\u4e0d\u4ec5\u5e2e\u52a9\u6211\u53d1\u73b0\u95ee\u9898\uff0c\u4e5f\u8ba9\u6211\u5728\u5b8c\u5584\u7cfb\u7edf\u7ec6\u8282\u548c"
    "\u68b3\u7406\u8bba\u6587\u5185\u5bb9\u65f6\u83b7\u5f97\u4e86\u8bb8\u591a\u542f\u53d1\u3002",
    "\u6700\u540e\uff0c\u611f\u8c22\u5bb6\u4eba\u5728\u6c42\u5b66\u548c\u8bba\u6587\u64b0\u5199\u8fc7\u7a0b\u4e2d\u7ed9\u4e88\u6211\u7684\u7406\u89e3\u3001"
    "\u9f13\u52b1\u4e0e\u966a\u4f34\u3002\u6b63\u662f\u7531\u4e8e\u4ed6\u4eec\u59cb\u7ec8\u5982\u4e00\u7684\u5173\u5fc3\u548c\u652f\u6301\uff0c\u6211\u624d"
    "\u80fd\u591f\u4ee5\u66f4\u52a0\u7a33\u5b9a\u7684\u5fc3\u6001\u5b8c\u6210\u6bd5\u4e1a\u8bbe\u8ba1\u4e0e\u8bba\u6587\u5199\u4f5c\u3002",
]


def output_candidates() -> list[Path]:
    candidates = [OUT_DOCX]
    stem = OUT_DOCX.stem
    suffix = OUT_DOCX.suffix
    for index in range(1, 6):
        candidates.append(OUT_DOCX.with_name(f"{stem}_修订版{index}{suffix}"))
    return candidates


def copy_source_to_writable_output() -> Path:
    last_error: PermissionError | None = None
    for candidate in output_candidates():
        try:
            shutil.copyfile(SRC_DOCX, candidate)
            return candidate
        except PermissionError as exc:
            last_error = exc
            continue
    if last_error is not None:
        raise last_error
    raise RuntimeError("No writable output path was available.")


NEW_ABSTRACT = (
    "\u968f\u7740\u77ed\u89c6\u9891\u521b\u4f5c\u3001\u6570\u5b57\u5a92\u4f53\u7f16\u8f91\u4e0e\u667a\u80fd\u89c6\u89c9\u5e94\u7528\u7684\u5feb\u901f\u53d1\u5c55\uff0c"
    "\u4eba\u8138\u66ff\u6362\u6280\u672f\u5df2\u4ece\u5355\u4e00\u7684\u7b97\u6cd5\u9a8c\u8bc1\u573a\u666f\u9010\u6b65\u8d70\u5411\u9762\u5411\u5b9e\u9645\u4f7f\u7528\u7684"
    "\u7cfb\u7edf\u5316\u5e94\u7528\u3002\u9488\u5bf9\u73b0\u6709\u4eba\u8138\u66ff\u6362\u5de5\u5177\u5728\u4ea4\u4e92\u6027\u4e0d\u8db3\u3001\u5904\u7406\u6d41\u7a0b\u5206\u6563\u3001"
    "\u7ed3\u679c\u7ba1\u7406\u4e0d\u4fbf\u4ee5\u53ca\u5de5\u7a0b\u53ef\u7ef4\u62a4\u6027\u4e0d\u5f3a\u7b49\u95ee\u9898\uff0c\u672c\u6587\u4ee5\u684c\u9762\u7aef\u4eba\u8138\u66ff\u6362"
    "\u7cfb\u7edf\u4e3a\u7814\u7a76\u5bf9\u8c61\uff0c\u56f4\u7ed5\u7cfb\u7edf\u9700\u6c42\u5206\u6790\u3001\u603b\u4f53\u67b6\u6784\u8bbe\u8ba1\u3001\u5173\u952e\u6a21\u5757\u5b9e\u73b0"
    "\u4e0e\u6d4b\u8bd5\u9a8c\u8bc1\u5c55\u5f00\u7814\u7a76\u3002\u7cfb\u7edf\u91c7\u7528\u524d\u540e\u7aef\u5206\u5c42\u7684\u5b9e\u73b0\u601d\u8def\uff0c\u4ee5 PyQt5 \u6784\u5efa"
    "\u684c\u9762\u4ea4\u4e92\u754c\u9762\uff0c\u4ee5 Django REST \u4e0e SQLite \u5b8c\u6210\u6570\u636e\u7ba1\u7406\u4e0e\u8d44\u6e90\u8c03\u5ea6\uff0c\u5e76\u7ed3\u5408 OpenCV "
    "\u4e0e InsightFace \u5b9e\u73b0\u89c6\u9891\u6362\u8138\u548c\u6444\u50cf\u5934\u5b9e\u65f6\u6362\u8138\u4e24\u7c7b\u4e3b\u8981\u4e1a\u52a1\u529f\u80fd\u3002\u5728\u5177\u4f53"
    "\u8bbe\u8ba1\u4e0a\uff0c\u672c\u6587\u5bf9\u7d20\u6750\u5bfc\u5165\u3001\u53c2\u6570\u8bbe\u7f6e\u3001\u4efb\u52a1\u542f\u52a8\u3001\u7ebf\u7a0b\u8c03\u5ea6\u3001\u7ed3\u679c\u4fdd\u5b58"
    "\u4e0e\u5386\u53f2\u56de\u653e\u7b49\u6d41\u7a0b\u8fdb\u884c\u4e86\u7cfb\u7edf\u5316\u68b3\u7406\uff0c\u91cd\u70b9\u89e3\u51b3\u4e86\u754c\u9762\u54cd\u5e94\u6027\u3001\u591a\u5a92\u4f53"
    "\u8d44\u6e90\u7ec4\u7ec7\u3001\u5904\u7406\u8fc7\u7a0b\u53ef\u8ddf\u8e2a\u6027\u4ee5\u53ca\u8f93\u51fa\u7ed3\u679c\u53ef\u7ba1\u7406\u6027\u7b49\u5173\u952e\u95ee\u9898\u3002\u4e3a\u4f7f"
    "\u7cfb\u7edf\u66f4\u5951\u5408\u5de5\u7a0b\u5e94\u7528\u9700\u6c42\uff0c\u8fd8\u5bf9\u524d\u7aef\u53cc\u6a21\u5f0f\u4ea4\u4e92\u3001\u540e\u7aef\u6570\u636e\u652f\u6491\u4e0e\u7ed3\u679c"
    "\u5f52\u6863\u673a\u5236\u8fdb\u884c\u4e86\u534f\u540c\u8bbe\u8ba1\uff0c\u4f7f\u7cfb\u7edf\u5728\u4fdd\u8bc1\u5904\u7406\u6548\u679c\u7684\u540c\u65f6\uff0c\u5177\u5907\u8f83\u597d\u7684"
    "\u4ea4\u4e92\u4f53\u9a8c\u4e0e\u4f7f\u7528\u8fde\u7eed\u6027\u3002\u901a\u8fc7\u529f\u80fd\u6d4b\u8bd5\u4e0e\u8fd0\u884c\u9a8c\u8bc1\u53ef\u4ee5\u770b\u51fa\uff0c\u6240\u8bbe\u8ba1\u7cfb\u7edf"
    "\u80fd\u591f\u7a33\u5b9a\u5b8c\u6210\u684c\u9762\u7aef\u4eba\u8138\u66ff\u6362\u7684\u4e3b\u8981\u4e1a\u52a1\u6d41\u7a0b\uff0c\u5728\u89c6\u9891\u5904\u7406\u3001\u6444\u50cf\u5934\u6f14\u793a"
    "\u3001\u7ed3\u679c\u5c55\u793a\u4e0e\u4fe1\u606f\u7ba1\u7406\u7b49\u65b9\u9762\u5747\u5177\u6709\u826f\u597d\u7684\u5b9e\u7528\u6027\u3002\u672c\u7814\u7a76\u5c06\u4eba\u8138\u66ff\u6362\u6280\u672f"
    "\u4e0e\u684c\u9762\u7aef\u8f6f\u4ef6\u5de5\u7a0b\u5b9e\u8df5\u76f8\u7ed3\u5408\uff0c\u4e0d\u4ec5\u5b8c\u6210\u4e86\u53ef\u7528\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5b9e\u73b0\uff0c\u4e5f\u4e3a\u76f8"
    "\u5173\u56fe\u50cf\u5904\u7406\u7cfb\u7edf\u7684\u5de5\u7a0b\u5316\u5efa\u8bbe\u4e0e\u6bd5\u4e1a\u8bbe\u8ba1\u7c7b\u8bfe\u9898\u63d0\u4f9b\u4e86\u53c2\u8003\u3002"
)

NEW_ENGLISH_ABSTRACT = (
    "With the rapid development of short-video creation, digital media editing, and intelligent vision "
    "applications, face-swapping technology is moving from isolated algorithm verification toward practical "
    "system-oriented deployment. Aiming at the common problems of insufficient interactivity, fragmented "
    "processing procedures, inconvenient result management, and limited maintainability in existing face-swapping "
    "tools, this thesis takes a desktop face-swapping system as the research object and studies its requirement "
    "analysis, overall architecture, key-module implementation, and testing process. The system adopts a layered "
    "frontend-backend design. PyQt5 is used to build the desktop interactive interface, Django REST and SQLite "
    "are employed to support data management and resource coordination, and OpenCV together with InsightFace is "
    "integrated to realize two major functions: video face swapping and real-time camera face swapping. In the "
    "implementation stage, the thesis systematically organizes the workflow of material import, parameter setting, "
    "task startup, thread scheduling, result storage, and historical replay, with particular attention to interface "
    "responsiveness, multimedia resource organization, process traceability, and output manageability. To make the "
    "system better aligned with engineering application requirements, the study further coordinates dual-mode "
    "frontend interaction, backend data support, and result archiving mechanisms, so that the system maintains both "
    "processing effectiveness and continuous usability. Functional testing and runtime verification show that the "
    "proposed system can stably complete the main workflow of desktop face swapping and performs well in video "
    "processing, camera demonstration, result presentation, and information management. This work combines face-"
    "swapping technology with desktop software engineering practice, and it can provide a useful reference for the "
    "engineering construction of related image-processing systems and similar graduation-design projects."
)


def output_candidates() -> list[Path]:
    candidates = [OUT_DOCX]
    stem = OUT_DOCX.stem
    suffix = OUT_DOCX.suffix
    for index in range(1, 6):
        candidates.append(OUT_DOCX.with_name(f"{stem}_\u4fee\u8ba2\u7248{index}{suffix}"))
    return candidates


def iter_block_items(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_rfonts(target, east_asia: str | None, ascii_font: str | None) -> None:
    if east_asia is None and ascii_font is None:
        return
    r_pr = target.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    if ascii_font:
        r_fonts.set(qn("w:ascii"), ascii_font)
        r_fonts.set(qn("w:hAnsi"), ascii_font)
        r_fonts.set(qn("w:cs"), ascii_font)
    if east_asia:
        r_fonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, *, east_asia: str | None, ascii_font: str | None, size_pt: float, bold: bool) -> None:
    style.font.name = ascii_font or east_asia
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    set_rfonts(style.element, east_asia=east_asia, ascii_font=ascii_font)


def set_run_font(run, *, east_asia: str | None, ascii_font: str | None, size_pt: float, bold: bool | None) -> None:
    run.font.name = ascii_font or east_asia
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    if ascii_font or east_asia:
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        if ascii_font:
            r_fonts.set(qn("w:ascii"), ascii_font)
            r_fonts.set(qn("w:hAnsi"), ascii_font)
            r_fonts.set(qn("w:cs"), ascii_font)
        if east_asia:
            r_fonts.set(qn("w:eastAsia"), east_asia)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(round(size_pt * 2))))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(round(size_pt * 2))))
    r_pr.append(sz_cs)


def format_runs(paragraph: Paragraph, *, east_asia: str | None, ascii_font: str | None, size_pt: float, bold: bool | None) -> None:
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size_pt=size_pt, bold=bold)


def is_caption_text(text: str) -> bool:
    normalized = text.strip()
    if not CAPTION_RE.fullmatch(normalized):
        return False
    if len(normalized) > 50:
        return False
    return not any(mark in normalized for mark in ("\u3002", "\uff1a", "\uff1b"))


def delete_content_after_acknowledgement(doc: Document) -> None:
    ack_found = False
    delete_from = None

    for block in iter_block_items(doc):
        if not isinstance(block, Paragraph):
            continue
        text = (block.text or "").strip()
        style_name = block.style.name if block.style else ""
        if text == ACK_HEADING:
            ack_found = True
            continue
        if ack_found and style_name == "Heading 1" and text and text != ACK_HEADING:
            delete_from = block._element
            break

    if delete_from is None:
        return

    current = delete_from
    while current is not None and current.tag != qn("w:sectPr"):
        next_element = current.getnext()
        remove_element(current)
        current = next_element

    for section in doc.sections:
        section.start_type = WD_SECTION_START.CONTINUOUS


def rename_heading_1(text: str) -> str:
    match = BODY_HEADING_1_RE.fullmatch(text.strip())
    if not match:
        return text
    number = int(match.group(1))
    title = match.group(2).strip()
    numeral = CHINESE_NUMERALS.get(number, str(number))
    return f"\u7b2c{numeral}\u7ae0{title}"


def normalize_heading_and_title_styles(doc: Document) -> None:
    set_style_font(doc.styles["Heading 1"], east_asia="\u9ed1\u4f53", ascii_font="SimHei", size_pt=15.9, bold=True)
    set_style_font(doc.styles["Heading 2"], east_asia="\u9ed1\u4f53", ascii_font="SimHei", size_pt=14.1, bold=True)
    if "Heading 3" in doc.styles:
        set_style_font(doc.styles["Heading 3"], east_asia="\u9ed1\u4f53", ascii_font="SimHei", size_pt=12.0, bold=True)
    if "toc 1" in doc.styles:
        set_style_font(doc.styles["toc 1"], east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=12.0, bold=False)
    if "toc 2" in doc.styles:
        set_style_font(doc.styles["toc 2"], east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=12.0, bold=False)


def format_front_title(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, east_asia="\u9ed1\u4f53", ascii_font="Times New Roman", size_pt=18.0, bold=True)


def format_heading_1(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, east_asia="\u9ed1\u4f53", ascii_font="SimHei", size_pt=15.9, bold=True)


def format_heading_2(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, east_asia="\u9ed1\u4f53", ascii_font="SimHei", size_pt=14.1, bold=True)


def format_heading_3(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, east_asia="\u9ed1\u4f53", ascii_font="SimHei", size_pt=12.0, bold=True)


def format_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=12.0, bold=False)


def format_reference_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=12.0, bold=False)


def convert_caption_text(text: str) -> str:
    match = CAPTION_RE.fullmatch(text.strip())
    if not match:
        return text
    prefix, chapter, seq, title = match.groups()
    return f"{prefix}{chapter}-{seq} {title.strip()}"


def is_formula_paragraph(paragraph: Paragraph) -> bool:
    xml = paragraph._element.xml
    return "m:oMath" in xml or "m:oMathPara" in xml


def is_standalone_formula(paragraph: Paragraph) -> bool:
    return is_formula_paragraph(paragraph) and not paragraph.text.strip()


def format_caption_paragraph(paragraph: Paragraph) -> None:
    paragraph.text = convert_caption_text(paragraph.text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    is_table_caption = paragraph.text.strip().startswith("\u8868")
    paragraph.paragraph_format.space_before = Pt(2 if is_table_caption else 0)
    paragraph.paragraph_format.space_after = Pt(2 if is_table_caption else 0)
    format_runs(
        paragraph,
        east_asia="\u5b8b\u4f53",
        ascii_font="Times New Roman",
        size_pt=10.5,
        bold=is_table_caption,
    )


def format_formula_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def set_formula_right_tab(paragraph: Paragraph, usable_width: Emu) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    existing_tabs = p_pr.find(qn("w:tabs"))
    if existing_tabs is not None:
        p_pr.remove(existing_tabs)

    tabs = OxmlElement("w:tabs")
    right_tab = OxmlElement("w:tab")
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), str(int(usable_width)))
    tabs.append(right_tab)
    p_pr.append(tabs)


def apply_equation_numbering(doc: Document) -> int:
    usable_width = Emu(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
    formula_count = 0

    for paragraph in doc.paragraphs:
        if not is_standalone_formula(paragraph):
            continue
        formula_count += 1
        format_formula_paragraph(paragraph)
        set_formula_right_tab(paragraph, usable_width)
        run = paragraph.add_run(f"\t\uff08{formula_count}\uff09")
        set_run_font(run, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=12.0, bold=False)

    return formula_count


def replace_abstract(doc: Document) -> None:
    abstract_title_seen = False
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text == "\u6458    \u8981":
            abstract_title_seen = True
            continue
        if not abstract_title_seen or not text:
            continue
        if text.startswith("\u5173\u952e\u8bcd"):
            break
        paragraph.text = NEW_ABSTRACT
        return


def replace_english_abstract(doc: Document) -> None:
    abstract_title_seen = False
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text == "ABSTRACT":
            abstract_title_seen = True
            continue
        if not abstract_title_seen or not text:
            continue
        if text.startswith("Keywords:"):
            break
        paragraph.text = NEW_ENGLISH_ABSTRACT
        return


def replace_acknowledgements(doc: Document) -> None:
    ack_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if (paragraph.text or "").strip() == ACK_HEADING:
            ack_index = index
            break

    if ack_index is None:
        return

    trailing = list(doc.paragraphs[ack_index + 1 :])
    if not trailing:
        return

    for paragraph, text in zip(trailing, NEW_ACKNOWLEDGEMENTS):
        paragraph.text = text

    for paragraph in trailing[len(NEW_ACKNOWLEDGEMENTS) :]:
        remove_element(paragraph._element)


def find_body_heading_index(doc: Document, heading_text: str) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if text == heading_text and "\t" not in paragraph.text and style_name == "Heading 2":
            return index
    return None


def rewrite_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if is_formula_paragraph(paragraph):
        return
    paragraph.text = text


def replace_paragraph_by_prefix(doc: Document, prefix: str, new_text: str) -> bool:
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text.startswith(prefix):
            rewrite_paragraph_text(paragraph, new_text)
            return True
    return False


def rewrite_related_method_sections(doc: Document) -> None:
    traditional_heading = find_body_heading_index(doc, "2.2 传统人脸替换方法")
    deep_heading = find_body_heading_index(doc, "2.3 基于深度学习的人脸替换方法与 InsightFace")

    if traditional_heading is not None:
        rewrite_paragraph_text(doc.paragraphs[traditional_heading], "2.2 基于OpenCV与三角剖分的传统人脸替换方法")
        replacements_22 = {
            1: "在本系统的传统处理链中，人脸替换的核心并不是神经网络推理，而是围绕 OpenCV 提供的几何变换与融合能力展开。代码实现上，传统路径在 process_frame_traditional 中完成单帧处理，并进一步调用 advanced_face_swap 或 simple_face_swap 执行局部变形与融合；其中高级方案以 Delaunay 三角剖分为核心，结合 OpenCV 的仿射变换和无缝融合函数完成整个人脸区域的替换。这一路径直接体现了传统换脸方法“关键点驱动、三角划分、像素级拼接”的基本特征。",
            2: "需要说明的是，传统路径虽然在检测阶段允许用户在界面中选择 dlib 或 OpenCV 级联检测器，但真正决定换脸效果的主体仍然是 OpenCV 图像处理流程本身。代码中无论使用哪一种检测入口，最终都要借助关键点集合构造三角网格，并通过逐三角形重映射完成纹理迁移。因此，从论文论述角度看，传统方法更应当围绕 OpenCV 的几何处理能力和 Delaunay 三角剖分机制展开，而不是围绕深度学习模型展开。",
            5: "从具体算法机制看，三角剖分的意义在于避免对整张人脸执行单一刚性变换，而是将人脸区域离散为多个局部三角单元，再分别计算对应的仿射映射关系。OpenCV 提供的 Subdiv2D、getAffineTransform、warpAffine 和掩模运算，使系统能够在逐三角形变换之后继续完成边界平滑与颜色过渡。对于第 k 个对应三角形，其局部几何映射可进一步写为",
            7: "若将源人脸关键点与目标人脸关键点投影到同一坐标系中，则传统换脸通常先依据这些关键点构造 Delaunay 三角形集合",
            10: "在工程实现上，系统会依据上述三角对应关系完成局部区域裁剪、仿射变形与结果叠加。若将源人脸关键点集合与目标人脸关键点集合分别记为",
            15: "在逐三角形映射完成之后，系统还需要对变形后的人脸纹理进行区域回填与边界融合。完成局部仿射变换后，基础融合可写为",
            20: "在具体像素重映射过程中，若目标三角形内部像素 p 满足",
            21: "结合 advanced_face_swap 中的 seamlessClone 与颜色校正处理，可以进一步减弱拼接边界和亮度差异，这也是传统方法在工程实践中较为常见的补偿思路[1][2][13][14]。",
            23: "表2-1 从实现依赖、资源消耗、复杂场景效果与系统适配性等角度比较了传统方法和深度学习方法。在本文系统中，前者主要对应 process_frame_traditional 所组织的 OpenCV 与 Delaunay 三角剖分方案；后者则主要对应 InsightFace FaceAnalysis 与 inswapper_128.onnx 所构成的深度学习处理路径。通过这一对照，可以更清楚地说明系统为何同时保留兼容性方案与主处理方案。",
        }
        for offset, text in replacements_22.items():
            rewrite_paragraph_text(doc.paragraphs[traditional_heading + offset], text)

    if deep_heading is not None:
        rewrite_paragraph_text(doc.paragraphs[deep_heading], "2.3 围绕dlib、InsightFace与inswapper_128的深度学习处理链")
        replacements_23 = {
            1: "从实际代码看，系统中的深度学习换脸并不是一个泛泛的概念，而是由 dlib 相关模型配置、InsightFace 人脸分析器和 inswapper_128.onnx 模型共同构成的工程处理链。其中，dlib 通过关键点模型文件与检测配置选项参与系统的人脸处理基础支撑；真正的深度学习主路径则由 InsightFace 的 FaceAnalysis 负责目标脸与源脸分析，再由 inswapper_128 模型执行身份迁移。与传统三角剖分方案相比，这一路径在复杂姿态、表情变化和光照扰动场景下通常能够获得更稳定、更自然的结果。",
            2: "在 face_swap.py 中，系统首先加载 shape_predictor_68_face_landmarks.dat 与 inswapper_128.onnx 等模型资源，并在界面层提供 dlib / OpenCV 检测器和 traditional / inswapper 换脸方法的可切换配置。就深度学习主路径而言，insightface_face_swap 会先调用 self.face_analyser.get 分别分析目标帧和源人脸图像，获得人脸框、关键点与身份相关特征；随后再调用 self.inswapper.get(result, face, source_face, paste_back=True) 完成逐脸替换，并在需要时执行颜色校正。这说明本文讨论的深度学习方法，核心应当围绕 InsightFace 分析流程与 inswapper_128 推理模型展开，同时结合 dlib 在系统检测配置中的辅助作用进行理解。",
            5: "若从方法抽象角度进行描述，本文所使用的深度学习换脸过程可以理解为人脸检测与分析、身份特征提取、目标属性保持以及生成映射的联合过程。设源人脸图像为 I_s，目标图像或目标帧为 I_t，则身份编码与生成过程可写为",
            6: "在上述表达式中，InsightFace 提供检测与身份表征基础，inswapper_128 负责将源身份特征迁移到目标帧之中，而 dlib 则主要在系统的检测配置与关键点支持层面提供辅助。因此，深度学习方法的关键不再是局部三角网格拼接，而是借助预训练表征与生成模型在高维特征空间内协调身份一致性与场景一致性。",
            8: "进一步地，身份保持项可用于约束 inswapper_128 输出结果与源人脸特征之间的一致性，因此身份保持项可进一步表示为",
            10: "从模型优化目标看，此类方法通常将总损失写为",
            11: "其中，E(.) 可理解为基于 InsightFace 识别能力提取的人脸嵌入；L_(rec) 用于约束内容重建或结构一致性，L_(per) 用于衡量高层感知差异，L_(adv) 则用于提升结果的真实感与分布一致性。结合实际代码可以看出，系统在深度学习主路径中并不重新训练模型，而是直接加载 inswapper_128.onnx，配合 FaceAnalysis 完成检测、识别与推理，从而在工程实现层面获得较高质量的人脸身份迁移结果[4][6][9][10][11][12]。",
            13: "表2-2 将 OpenCV、dlib、InsightFace、inswapper_128、PyQt5、Django REST 与 SQLite 等关键技术同系统职责进行对应，有助于从工程实现角度理解各组件之间的协作关系。尤其在人脸处理层，OpenCV 负责传统几何处理与媒体操作，dlib 提供关键点模型与检测配置支撑，InsightFace 负责人脸分析，而 inswapper_128 承担核心推理任务。",
        }
        for offset, text in replacements_23.items():
            rewrite_paragraph_text(doc.paragraphs[deep_heading + offset], text)


def update_method_tables(doc: Document) -> None:
    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []

        if header == ["比较维度", "传统方法", "深度学习方法"] and len(table.rows) >= 5:
            rows = [
                ("实现复杂度", "基于 OpenCV 与三角剖分，实现路径清晰", "依赖 InsightFace 与 inswapper_128 等模型资源，部署较复杂"),
                ("资源消耗", "CPU 侧也可运行，便于兼容部署", "更依赖 GPU 或较高推理性能"),
                ("复杂场景效果", "对姿态、遮挡和光照变化更敏感", "在复杂场景下通常更稳定"),
                ("系统适配性", "适合作为兼容与回退方案", "适合作为主要处理方案"),
            ]
            for row, values in zip(table.rows[1:], rows):
                for cell, value in zip(row.cells, values):
                    cell.text = value

        if header == ["关键技术", "在系统中的作用", "对应实现位置"]:
            target_rows = [
                ("OpenCV", "负责视频读写、Subdiv2D 三角剖分、仿射变换与融合输出", "传统换脸流程与媒体处理模块"),
                ("dlib", "提供 68 点关键点模型，并为检测配置提供支撑", "shape_predictor_68_face_landmarks.dat 与检测选项"),
                ("InsightFace", "负责 FaceAnalysis 检测、识别与人脸分析", "FaceAnalysis 初始化与分析流程"),
                ("inswapper_128", "通过 self.inswapper.get 执行核心换脸推理", "models/inswapper_128.onnx 与换脸调用逻辑"),
                ("PyQt5", "承担桌面界面、列表展示、状态反馈和模式切换", "EnhancedFaceSwapUI"),
                ("Django REST", "负责图片、视频、输出结果和任务的接口管理", "api 视图集与路由"),
                ("SQLite", "保存素材、任务与输出视频元数据", "core 模型定义"),
            ]

            while len(table.rows) - 1 < len(target_rows):
                table.add_row()

            for row, values in zip(table.rows[1:], target_rows):
                for cell, value in zip(row.cells, values):
                    cell.text = value


def relocate_inline_citations(doc: Document) -> None:
    replacements = [
        (
            "从技术演进路径来看，人脸替换大致经历了由传统几何配准与区域融合向深度学习驱动身份迁移方案逐步过渡的发展过程。",
            "从技术演进路径来看，人脸替换大致经历了由传统几何配准与区域融合向深度学习驱动身份迁移方案逐步过渡的发展过程。早期方法主要依赖检测、关键点定位、仿射变换和局部融合[1][2]，其优势在于实现链路相对清晰、外部依赖较少；生成对抗网络等深度学习框架的引入[3]，则推动了换脸方法从像素级拼接逐步走向身份表征驱动。随后，DeepFace[5]、FaceNet[4]、ArcFace[6] 等人脸识别表征方法，以及 FSGAN[9]、SimSwap[10]、FaceShifter[11] 和 InsightFace[12] 等身份迁移方案相继出现，使换脸任务在身份保持能力和复杂场景适应性方面获得了更为明显的提升。",
        ),
        (
            "不过，国内外已有研究更多将注意力集中在模型效果、数据集构建以及视觉评价指标上，而较少深入讨论桌面系统中的工作流组织问题。",
            "不过，国内外已有研究更多将注意力集中在模型效果、数据集构建以及视觉评价指标上，其中既包括面向伪造视频检测任务的研究[7]，也包括面向数据集构建与基准评价的公开工作[8]；相较之下，对桌面系统中的工作流组织问题讨论较少。对于真正需要交付给用户使用的原型系统而言，模型只是整体链路中的一个环节，还必须回应素材如何加载、失败如何提示、结果如何保存以及历史记录如何回查等问题。当前不少开源实现虽然已经能够生成较好的视觉结果，但在资源管理与界面组织方面仍明显带有实验性。",
        ),
        (
            "结合 advanced_face_swap 中的 seamlessClone 与颜色校正处理，可以进一步减弱拼接边界和亮度差异，这也是传统方法在工程实践中较为常见的补偿思路",
            "结合 advanced_face_swap 中的 seamlessClone[13][14] 与颜色校正处理，可以进一步减弱拼接边界和亮度差异，这也是传统换脸方法在工程实践中较为常见的补偿思路[1][2]。",
        ),
        (
            "其中，E(.) 可理解为基于 InsightFace 识别能力提取的人脸嵌入；L_(rec) 用于约束内容重建或结构一致性，L_(per) 用于衡量高层感知差异，L_(adv) 则用于提升结果的真实感与分布一致性。",
            "其中，E(.) 可理解为基于 InsightFace[12] 识别能力提取的人脸嵌入，这类身份表征思路与 FaceNet[4]、ArcFace[6] 等工作一脉相承；L_(rec) 用于约束内容重建或结构一致性，L_(per) 用于衡量高层感知差异，L_(adv) 则用于提升结果的真实感与分布一致性。结合实际代码可以看出，系统在深度学习主路径中并不重新训练模型，而是直接加载 inswapper_128.onnx，配合 FaceAnalysis 完成检测、识别与推理，从而在工程实现层面获得较高质量的人脸身份迁移结果，这与 FSGAN[9]、SimSwap[10]、FaceShifter[11] 等方法关注的身份迁移目标相一致。",
        ),
        (
            "OpenCV 是当前系统中最基础的工具之一。",
            "OpenCV[13][14] 是当前系统中最基础的工具之一。无论是前端列表中的视频缩略图生成，还是后端在登记输入视频与输出视频时读取分辨率、时长和帧率信息，都需要依赖 OpenCV 对视频文件的访问能力。对于工程系统而言，OpenCV 的价值并不局限于算法处理本身，更在于它为媒体元数据的统一获取提供了相对稳定的入口。",
        ),
        (
            "PyQt5 主要承担桌面端的人机交互层任务。",
            "PyQt5[15][16] 主要承担桌面端的人机交互层任务。项目中的增强版界面类集成了窗口初始化、控件布局、媒体播放、列表展示、处理按钮、模式切换、摄像头开关以及状态栏更新等多项逻辑。结合多媒体桌面系统设计中的常见做法[22]，相较于命令行程序，PyQt5 能够让用户以更低的认知成本完成素材选择与处理操作。",
        ),
        (
            "在数据管理层面，本文并未把所有逻辑都堆放在桌面端，而是引入 Django REST Framework 作为统一的资源接口层。",
            "在数据管理层面，本文并未把所有逻辑都堆放在桌面端，而是引入 Django REST Framework[17][18] 作为统一的资源接口层。图片、输入视频、输出视频与处理任务都能够以 REST 资源分层设计[25][26] 的形式被查询和记录，相关元数据再结合 SQLite 持久化方案[21] 进行归档。这种组织方式既有利于后续扩展到更多客户端，也有助于将交互界面与数据归档两类职责区分开来。",
        ),
        (
            "后端接口采用 REST 风格进行组织。",
            "后端接口采用 REST 风格[25][26] 进行组织，并以 Django REST Framework[17][18] 完成资源注册与访问控制的基础结构。当前项目已经注册 images、videos、outputs 和 tasks 四类资源，并保留了图片与视频上传入口。虽然本文更强调工程验证而非接口安全设计，但这种资源化组织方式仍为后续接入鉴权、分页、筛选与多用户逻辑提供了基础。",
        ),
        (
            "从代码结构上看，系统的并行协作主要围绕三类对象展开，即界面控制类 EnhancedFaceSwapUI、视频处理线程 VideoProcessingThread 与摄像头线程 CameraProcessingThread。",
            "从代码结构上看，系统的并行协作主要围绕三类对象展开，即界面控制类 EnhancedFaceSwapUI、视频处理线程 VideoProcessingThread 与摄像头线程 CameraProcessingThread。三者与 original_app 以及 DatabaseManager 共同构成系统的核心协作关系：EnhancedFaceSwapUI 负责承载用户交互和界面状态，两个 QThread[19][20] 负责执行长时间或持续运行任务，original_app 负责具体换脸算法逻辑，而 DatabaseManager 则承担资源管理与结果入库功能。",
        ),
        (
            "本章测试的定位是工程验证，而非算法竞赛式 benchmark。",
            "本章测试的定位是工程验证，而非算法竞赛式 benchmark。验证重点主要包括三个方面：其一，系统业务链路能否真正贯通；其二，界面在处理与切换过程中能否保持基本响应；其三，数据库记录与本地文件之间是否存在可观察且可解释的偏差。与软件工程文档撰写中的工程验证思路[23] 和计算机视觉课程项目常见的工程化实现路径[24] 相比，本文更关注当前项目在真实材料条件下的可运行性与可说明性。",
        ),
    ]

    for prefix, new_text in replacements:
        replace_paragraph_by_prefix(doc, prefix, new_text)


def set_page_break_before(paragraph: Paragraph, enabled: bool) -> None:
    paragraph.paragraph_format.page_break_before = enabled


def enforce_front_matter_pagination(doc: Document) -> None:
    first_body_heading_done = False

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if not text:
            continue

        if text == "\u6458    \u8981":
            set_page_break_before(paragraph, True)
            continue

        if text == "ABSTRACT":
            set_page_break_before(paragraph, True)
            continue

        if text == "\u76ee    \u5f55":
            set_page_break_before(paragraph, True)
            continue

        if not first_body_heading_done and style_name == "Heading 1" and text.startswith("\u7b2c\u4e00\u7ae0"):
            set_page_break_before(paragraph, True)
            first_body_heading_done = True


def apply_top_level_paragraph_formatting(doc: Document) -> None:
    in_body = False
    in_references = False
    in_ack = False

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text and not is_formula_paragraph(paragraph):
            continue

        if text in FRONT_TITLES:
            format_front_title(paragraph)
            continue

        if text == REFERENCE_HEADING:
            in_body = False
            in_references = True
            in_ack = False
            format_heading_1(paragraph)
            continue

        if text == ACK_HEADING:
            in_body = False
            in_references = False
            in_ack = True
            format_heading_1(paragraph)
            continue

        if BODY_HEADING_1_RE.fullmatch(text):
            paragraph.text = rename_heading_1(text)
            in_body = True
            in_references = False
            in_ack = False
            format_heading_1(paragraph)
            continue

        if BODY_HEADING_3_RE.fullmatch(text):
            format_heading_3(paragraph)
            continue

        if BODY_HEADING_2_RE.fullmatch(text):
            format_heading_2(paragraph)
            continue

        if is_caption_text(text):
            format_caption_paragraph(paragraph)
            continue

        if is_formula_paragraph(paragraph):
            format_formula_paragraph(paragraph)
            continue

        if in_body or in_ack:
            format_body_paragraph(paragraph)
        elif in_references and REF_ENTRY_RE.match(text):
            format_reference_paragraph(paragraph)


def ensure_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": 50, "bottom": 50, "left": 90, "right": 90}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def clear_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is not None:
        tc_pr.remove(shd)


def set_border(edge_parent, edge: str, *, val: str, size: int = 0, color: str = "000000") -> None:
    border = edge_parent.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        edge_parent.append(border)
    border.set(qn("w:val"), val)
    if val not in {"nil", "none"}:
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
    else:
        border.attrib.pop(qn("w:sz"), None)
        border.attrib.pop(qn("w:space"), None)
        border.attrib.pop(qn("w:color"), None)


def set_table_three_line_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    set_border(tbl_borders, "top", val="single", size=12)
    set_border(tbl_borders, "bottom", val="single", size=12)
    set_border(tbl_borders, "left", val="nil")
    set_border(tbl_borders, "right", val="nil")
    set_border(tbl_borders, "insideH", val="nil")
    set_border(tbl_borders, "insideV", val="nil")


def clear_table_style_formatting(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblStyle", "w:tblLook"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cnf_style = tr_pr.find(qn("w:cnfStyle"))
        if cnf_style is not None:
            tr_pr.remove(cnf_style)

        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cnf_style = tc_pr.find(qn("w:cnfStyle"))
            if cnf_style is not None:
                tc_pr.remove(cnf_style)


def set_cell_borders(cell, *, top: str = "nil", bottom: str = "nil", left: str = "nil", right: str = "nil") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    set_border(tc_borders, "top", val=top)
    set_border(tc_borders, "bottom", val=bottom if bottom in {"nil", "none"} else "single", size=8)
    set_border(tc_borders, "left", val=left)
    set_border(tc_borders, "right", val=right)


def set_table_layout_fixed(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_width(table: Table, width_emu: Emu) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_emu / 635)))


def format_table_cell_paragraph(paragraph: Paragraph, *, is_header: bool, center_body: bool) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_runs(paragraph, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=10.5, bold=False)


def format_tables(doc: Document) -> None:
    usable_width = Emu(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)

    for table in doc.tables:
        clear_table_style_formatting(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_layout_fixed(table)
        set_table_width(table, Emu(int(usable_width * 0.90)))
        set_table_three_line_borders(table)

        if table.rows:
            mark_header_row(table.rows[0])
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            table.rows[0].height = Pt(18)

        for row_index, row in enumerate(table.rows):
            is_header = row_index == 0
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(18)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                ensure_cell_margins(cell)
                clear_cell_shading(cell)
                set_cell_borders(cell, bottom="single" if is_header else "nil")
                for paragraph in cell.paragraphs:
                    format_table_cell_paragraph(paragraph, is_header=is_header, center_body=True)


def apply_equation_numbering_with_word(docx_path: Path) -> int:
    import win32com.client  # type: ignore

    word = None
    document = None
    formula_count = 0
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(docx_path))
        usable_width = (
            document.PageSetup.PageWidth
            - document.PageSetup.LeftMargin
            - document.PageSetup.RightMargin
        )

        for paragraph in document.Paragraphs:
            rng = paragraph.Range
            plain_text = rng.Text.replace("\r", "").replace("\x07", "").replace("\t", "")
            if rng.OMaths.Count == 0 or plain_text.strip():
                continue

            formula_count += 1
            fmt = rng.ParagraphFormat
            fmt.Alignment = 1
            fmt.LeftIndent = 0
            fmt.RightIndent = 0
            fmt.FirstLineIndent = 0
            fmt.SpaceBefore = 0
            fmt.SpaceAfter = 0
            fmt.TabStops.ClearAll()
            fmt.TabStops.Add(usable_width, 2, 0)

            insert_range = document.Range(rng.End - 1, rng.End - 1)
            insert_range.Text = f"\t\uff08{formula_count}\uff09"

        document.Repaginate()
        document.Save()
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
        if word is not None:
            word.Quit()

    return formula_count


def collect_toc_entries_with_word(docx_path: Path) -> list[tuple[int, str, int]]:
    import win32com.client  # type: ignore

    word = None
    document = None
    entries: list[tuple[int, str, int]] = []
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(docx_path))
        document.Repaginate()
        started_main = False

        for paragraph in document.Paragraphs:
            text = paragraph.Range.Text.replace("\r", "").replace("\x07", "").strip()
            if not text:
                continue
            outline_level = int(paragraph.OutlineLevel)

            if not started_main:
                if outline_level == 1 and text.startswith("\u7b2c\u4e00\u7ae0"):
                    started_main = True
                else:
                    continue

            if outline_level == 1:
                if not (text.startswith("\u7b2c") or text in {REFERENCE_HEADING, ACK_HEADING}):
                    continue
                level = 1
            elif outline_level == 2:
                if not BODY_HEADING_2_RE.fullmatch(text):
                    continue
                level = 2
            else:
                continue

            page = int(paragraph.Range.Information(3))
            entries.append((level, text, page))

            if text == ACK_HEADING:
                break
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
        if word is not None:
            word.Quit()

    return entries


def rebuild_static_toc(doc: Document, entries: list[tuple[int, str, int]]) -> None:
    toc_title_index = None
    first_body_index = None

    for index, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if text == "\u76ee    \u5f55":
            toc_title_index = index
            continue
        if toc_title_index is not None and style_name == "Heading 1" and text.startswith("\u7b2c\u4e00\u7ae0"):
            first_body_index = index
            break

    if toc_title_index is None or first_body_index is None:
        raise RuntimeError("Failed to locate TOC title or first body heading.")

    first_body_heading = doc.paragraphs[first_body_index]
    to_remove = list(doc.paragraphs[toc_title_index + 1 : first_body_index])
    for paragraph in to_remove:
        remove_element(paragraph._element)

    usable_width = Emu(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)

    for level, text, page in entries:
        style_name = "toc 1" if level == 1 else "toc 2"
        inserted = first_body_heading.insert_paragraph_before(f"{text}\t{page}", style=style_name)
        inserted.alignment = WD_ALIGN_PARAGRAPH.LEFT
        inserted.paragraph_format.first_line_indent = Pt(0)
        inserted.paragraph_format.space_before = Pt(0)
        inserted.paragraph_format.space_after = Pt(0)
        inserted.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        format_runs(inserted, east_asia="\u5b8b\u4f53", ascii_font="Times New Roman", size_pt=12.0, bold=False)


def main() -> None:
    if not SRC_DOCX.exists():
        raise FileNotFoundError(SRC_DOCX)

    output_path = copy_source_to_writable_output()
    doc = Document(str(output_path))

    delete_content_after_acknowledgement(doc)
    rewrite_related_method_sections(doc)
    update_method_tables(doc)
    relocate_inline_citations(doc)
    replace_abstract(doc)
    replace_english_abstract(doc)
    replace_acknowledgements(doc)
    normalize_heading_and_title_styles(doc)
    apply_top_level_paragraph_formatting(doc)
    enforce_front_matter_pagination(doc)
    format_tables(doc)
    formula_count = apply_equation_numbering(doc)
    doc.save(str(output_path))

    toc_entries: list[tuple[int, str, int]] = []
    try:
        toc_entries = collect_toc_entries_with_word(output_path)
    except Exception as exc:
        print(f"WARNING=Skipped TOC rebuild because Word COM is unavailable: {exc}")

    final_doc = Document(str(output_path))
    if toc_entries:
        rebuild_static_toc(final_doc, toc_entries)
        final_doc.save(str(output_path))

    print(f"OUTPUT={output_path}")
    print(f"FORMULAS={formula_count}")
    print(f"TABLES={len(final_doc.tables)}")


if __name__ == "__main__":
    main()
