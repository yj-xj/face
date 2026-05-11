from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
BODY_FIRST_INDENT = Pt(24)
LINE_SPACING = Pt(20)
SPECIAL_TITLES = {"ABSTRACT", "参考文献", "致谢"}

PARAGRAPH_REWRITES: list[tuple[str, str]] = [
    (
        "随着短视频创作、数字媒体编辑与智能视觉应用的快速发展，人脸替换技术已从单一的算法验证场景逐步走向面向实际使用的系统化应用。",
        "本文讨论的对象并不是从零开始搭建的换脸模型，而是一套已经形成前后端雏形的人脸替换系统。论文以现有仓库中的 frontend、backend-django、数据库记录和本地媒体文件为依据，重新梳理系统的业务链路、模块边界与实现细节。前端采用 PyQt5 构建统一桌面界面，围绕视频模式与摄像头模式组织交互流程；后端以 Django REST 和 SQLite 管理人脸图片、输入视频、输出结果及任务信息；在人脸处理层，则同时保留基于 OpenCV 与三角剖分的传统路径，以及由 InsightFace 与 inswapper_128 组成的主处理路径。论文重点说明素材加载、参数设置、线程调度、结果保存和回放验证等环节如何衔接为可运行的完整流程，并结合当前已有的数据库记录与样例素材，对系统的功能完成度、结果留痕能力和现实问题进行分析。就目前版本而言，该系统已经能够支撑桌面端换脸演示、离线视频处理和实时摄像头展示等主要任务，同时也暴露出历史路径治理、结果补登记和任务调度仍待完善的问题。",
    ),
    (
        "With the rapid development of short-video creation, digital media editing, and intelligent vision applications, face-swapping technology is moving from isolated algorithm verification toward practical system-oriented deployment.",
        "This thesis does not treat face swapping as a model built from scratch. Instead, it examines a desktop face-swapping system that already has a working frontend, backend, database records, and local media assets, and then reorganizes those materials into a complete software-engineering narrative. The desktop side is implemented with PyQt5 and supports both video mode and camera mode in a single interface. On the backend, Django REST and SQLite are used to manage face images, input videos, output records, and task-related metadata. At the processing layer, the project keeps both a traditional OpenCV-based path and an InsightFace plus inswapper_128 path so that compatibility and visual quality can be balanced in different runtime conditions. The thesis focuses on how material loading, parameter preparation, threaded processing, result archiving, and replay verification are connected into a continuous workflow rather than being scattered across isolated demo steps. Based on the codebase, database contents, and existing sample files, the study further analyzes what the current system can already accomplish and which engineering issues still remain, especially path maintenance, result backfilling, and task scheduling. The final document therefore emphasizes practical system behavior and implementation logic more than abstract model discussion alone.",
    ),
    (
        "本章围绕课题提出的现实背景、研究价值以及全文任务安排展开论述。",
        "本章先从短视频编辑和生成式视觉应用的现实需求切入，说明为什么毕业设计不能只停留在单次换脸演示，而需要面向可操作的桌面系统展开论述。随后，结合现有研究路线、仓库中已经形成的前后端结构以及本文拟解决的问题，交代论文的研究边界、主要工作和章节安排，为后文的需求分析与实现说明奠定基础。",
    ),
    (
        "本章将从系统实现所依赖的关键技术出发，介绍人脸替换任务的基本定义、典型处理流程，以及传统方法与深度学习方法之间的差异，并进一步说明 OpenCV、PyQt5、Django REST Framework 与 SQLite 在本系统中的具体作用。",
        "本章不单纯罗列技术名词，而是围绕系统真实运行所依赖的处理链展开说明。内容先界定人脸替换任务在当前项目中的完整流程，再分别讨论传统方案、InsightFace 主路径、OpenCV 媒体处理、PyQt5 交互机制以及 Django REST Framework 与 SQLite 的数据管理职责，从而为后续需求分析和设计实现提供统一的技术语境。",
    ),
    (
        "本章将结合真实使用过程对系统需求展开分析，重点说明视频换脸与摄像头演示两类业务场景下的功能目标及约束条件。",
        "本章把需求分析建立在真实使用过程之上，而不是停留在抽象功能清单。论文将围绕视频换脸和摄像头演示两类核心场景，说明用户在素材选择、参数设置、处理启动、结果保存和历史回放等环节中的实际诉求，并进一步分析这些诉求对应的功能边界、非功能约束和实现前提。",
    ),
    (
        "本章将在前文需求分析的基础上，对系统整体结构进行抽象与规划。",
        "本章在需求分析的基础上，对当前项目的整体结构做工程化整理。正文将依次说明设计目标、分层架构、模块职责、接口组织、数据库关系和本机部署形态，把原本分散在前端、后端与脚本中的实现内容归纳为边界清楚、职责明确的系统方案。",
    ),
    (
        "本章将围绕系统从界面组织到处理线程协作的关键实现展开说明。",
        "本章聚焦系统已经落地的关键实现，而不是停留在设计层描述。内容将从前端主界面、双模式交互、素材加载、视频处理、摄像头处理、结果归档以及线程协作几个方面展开，说明现有代码如何把换脸能力组织成一套能够连续运行、能够回放结果、也能够保留处理记录的桌面应用。",
    ),
    (
        "本章以工程验证为导向，对系统当前版本的运行情况展开测试与分析。",
        "本章围绕当前版本的真实运行结果开展工程验证。测试重点不是和公开 benchmark 做数值对比，而是结合现有数据库记录、本地媒体文件、界面交互流程和输出结果，检查系统在素材加载、处理执行、结果回放和信息留痕等环节上的实际表现，并据此分析系统已经达到的完成度与仍然存在的短板。",
    ),
    (
        "本章将对全文工作进行归纳，并从工程实现视角总结系统已经形成的特点与价值。",
        "本章在回顾全文工作的基础上，重新归纳这套人脸替换系统已经形成的工程特点与现实价值。相较于单纯强调算法效果，本章更关注系统目前如何支撑视频模式、摄像头模式、结果归档和历史回放，以及这些能力在毕业设计场景中意味着什么，同时也会明确指出后续需要继续补足的工程环节。",
    ),
    (
        "本章测试的定位是工程验证，而非算法竞赛式 benchmark。",
        "本节中的“测试”并不是把系统放进统一 benchmark 做横向对比，而是检查当前项目在真实素材、真实路径和真实数据库记录条件下能否稳定运行。关注点主要集中在三方面：一是素材加载、处理启动、结果保存这条链路是否真正打通；二是 PyQt5 界面在视频模式、摄像头模式和结果回放之间切换时能否保持响应；三是数据库记录与本地文件是否一致，尤其是历史 local_path 和输出结果之间是否存在可解释的偏差。",
    ),
    (
        "测试环境延续了项目开发阶段的单机模式。",
        "测试环境沿用项目开发阶段的单机部署方式。前端通过 frontend/main.py 启动桌面界面，后端由本机 Django REST 服务提供资源接口，数据库采用 backend-django/db.sqlite3，本地媒体素材和输出结果则分别保存在数据目录与 output_videos 目录中。这样的环境虽然不追求分布式能力，却最能反映当前项目在答辩演示和日常调试中的真实运行边界。",
    ),
    (
        "截至本文撰写时，系统数据库中已经记录 19 条人脸图片、13 条输入视频与 5 条输出视频信息，同时输出目录中实际保存着 10 个 mp4 文件。",
        "截至本文撰写时，数据库中已经保留 19 条人脸图片记录、13 条输入视频记录和 5 条输出视频记录，而 output_videos 目录中实际还能看到 10 个 mp4 文件。这样的数据状态说明系统已经积累了可供验证的真实材料，同时也直接暴露出“文件已经生成但记录尚未完全补齐”的工程问题。下文据此结合环境配置、统计图、功能验证表和案例图，对系统当前的完成度进行说明。",
    ),
    (
        "测试方案是依据当前项目的真实调用链设计的，而不是脱离实现随意拆分。",
        "测试方案直接对应当前项目的真实调用链，而不是脱离实现随意拆分功能点。论文把 loadFaceImages、loadVideos、startProcessing、toggleCamera、takeSnapshot、processingFinished 与 saveOutputVideoToDatabase 等关键过程串联起来考察，重点观察素材是否顺利进入系统、线程是否正确启动、结果是否能够回放，以及处理完成后能否留下可查询的记录。",
    ),
    (
        "从功能验证结果来看，当前系统已经覆盖主要处理流程。",
        "从功能验证结果来看，当前系统已经具备较完整的主流程。前端既可以从数据库读取人脸素材和视频素材，也可以在后端不可用时退回到本地目录扫描；视频模式能够创建 VideoProcessingThread 执行离线处理，并在完成后回放结果、尝试登记输出记录；摄像头模式则支持开启取流、关闭线程和保存快照。结合测试过程可以看出，这套系统已经不再是一次性脚本调用，而是拥有明确入口、处理反馈和结果留痕的桌面原型。不过，验证中也保留了真实问题，例如历史路径失效时界面会跳过无效条目而不是整体崩溃，这说明系统具备一定的异常容忍能力，但数据库记录与本地文件之间仍缺少更严格的同步机制。",
    ),
    (
        "根据当前记录，输入视频平均时长约为 11.24 秒，平均帧率约为 38.75 FPS；输出视频平均时长约为 3.96 秒，单条输出记录的平均处理时间约为 1.98 秒。",
        "从当前样本统计来看，输入视频平均时长约为 11.24 秒，平均帧率约为 38.75 FPS；数据库中已有输出记录对应的视频平均时长约为 3.96 秒，单条输出记录的平均处理时间约为 1.98 秒。这里的数值并不承担 benchmark 指标的意义，但它们足以说明当前系统主要服务于中短时长样本，更接近答辩演示和工程验证场景，而不是长视频批处理环境。",
    ),
    (
        "相较于时长与帧率本身，更值得关注的是资源一致性差异。",
        "比时长和帧率更值得关注的是资源一致性。现有记录中，带有 local_path 的人脸图片共有 9 条且全部可用；带有 local_path 的输入视频共有 8 条，其中 1 条已经失效；与此同时，output_videos 目录中的 mp4 文件数量又比数据库输出记录多出 5 个。这组数据说明系统当前最明显的问题并不是“完全无法处理”，而是历史路径治理和结果补登记尚未形成稳定机制。",
    ),
    (
        "从工程实践的角度来看，这类问题并不少见。",
        "从工程实践的角度看，这类偏差并不罕见。当前系统已经能够在数据不完全一致的情况下继续运行，例如列表加载时会主动跳过失效路径，界面也不会因单条错误记录而立刻失去响应；但如果后续希望把系统从答辩演示原型进一步发展为可长期维护的工具，就必须补足路径整理、结果补录和任务状态同步这些基础能力。",
    ),
    (
        "为了使测试分析不局限于统计表与文字说明，本文结合 image 目录中新增的样例图像，对传统方法与 InsightFace 的输出结果进行了重新整理，并生成两组对比图用于说明系统在真实照片源人脸与绘制源人脸条件下的表现差异。",
        "为避免案例分析只停留在统计表和文字判断，本文从 image 目录中整理出代表性样例，分别展示真实照片源人脸和绘制源人脸两种条件下的输出结果。这样处理的目的，不是把这部分内容包装成严格控制变量的视觉实验，而是借助真实样例说明系统在不同素材条件下的可用性、稳定性和视觉特征。",
    ),
    (
        "结合前文验证结果可以看出，当前系统最突出的不足主要集中在资源记录、结果登记与任务调度三个方面。",
        "综合前文验证结果，当前系统最突出的短板仍然集中在资源记录、结果登记和任务调度三个方面。输入素材的历史路径还存在失效条目，output_videos 目录中的部分实际结果尚未回填到数据库，而 processing_tasks 虽然已经在模型层预留，却还没有真正承担批量排队、失败重试和取消控制等职责。这些问题都不是算法效果层面的缺陷，而是工程闭环尚未完全补齐的表现。",
    ),
    (
        "后续改进应当围绕上述问题逐项展开。",
        "后续改进应当围绕上述问题逐项推进。第一类工作是路径治理，需要把素材和结果逐步收敛到更稳定的目录规则中，并为旧记录补充修复脚本；第二类工作是结果补登记，需要让 output_videos 中已经生成的文件能够被重新扫描和纳入数据库；第三类工作则是把 processing_tasks 从预留模型发展为真正可用的任务队列，同时补充更清楚的日志分类和错误反馈界面。只有这些基础能力逐步补齐，系统才更适合在毕业设计阶段之后继续使用。",
    ),
    (
        "本文围绕既有人脸替换项目的梳理、补充与论文表达展开工作，基于真实代码、数据库记录、媒体文件和界面材料，完成了从需求分析、总体设计到详细实现与工程测试的系统说明。",
        "本文的工作重点不是重新提出一套换脸网络，而是把既有人脸替换项目整理为结构完整、论证清楚的系统型毕业设计。论文以真实代码、数据库记录、媒体文件和界面材料为基础，完成了从需求分析、总体设计到详细实现与工程测试的整体说明，并把素材加载、参数准备、线程处理、结果回放和数据留痕这些原本分散在代码中的行为重新组织成可解释的系统流程。",
    ),
    (
        "从实现结果来看，当前项目已经形成相对完整的桌面端处理流程。",
        "从已经完成的实现来看，当前项目具备较完整的桌面端处理流程。EnhancedFaceSwapUI 负责双模式界面和状态联动，DatabaseManager 承担前后端资源协同，VideoProcessingThread 与 CameraProcessingThread 分别支撑离线处理和实时处理，而 saveOutputVideoToDatabase 等逻辑又把处理结果纳入统一记录体系。这说明项目已经超出单次算法演示的范围，具备系统实现型毕业设计应有的基本形态。",
    ),
    (
        "本系统的主要特点并不在于提出新的换脸网络，而在于将既有换脸方法整合进一套实际可运行的桌面原型。",
        "本系统最明显的特点，不在于提出新的换脸网络，而在于把既有换脸方法组织成一套确实能够运行的桌面原型。视频模式与摄像头模式被纳入同一界面，并通过状态切换、线程控制和结果回放维持连续的交互过程；图片、视频和输出结果又被统一纳入后端与数据库管理流程，使系统能够保存并查询多次处理记录。与此同时，系统同时保留传统方法与 InsightFace 两条处理路径，因此在不同运行条件下仍能保持基本处理能力。",
    ),
    (
        "论文呈现方式的优化同样构成了本次工作的一个特点。",
        "论文呈现方式的调整同样是本次整理工作的重要部分。正文中的流程图、结构图和案例图不再依赖表格拼接，而是结合真实界面截图、脚本生成图和已有样例进行统一排版，这使文档的版面更接近正式论文的表达方式。虽然这种处理不改变系统功能本身，却明显提高了结构说明的清晰度，也有助于减少阅读过程中因图文组织混乱带来的理解成本。",
    ),
    (
        "后续工作可首先从资源整理入手，通过调整入库规则、本地目录组织方式以及历史记录修复策略，尽量减少 local_path 失效和输出文件孤立等现象。",
        "后续工作可以先从资源整理入手。现阶段最直接的任务，是通过调整入库规则、本地目录组织方式和历史记录修复策略，尽量减少 local_path 失效和输出文件孤立现象；与此同时，还需要把 processing_tasks 从当前的预留模型发展为真正可用的处理队列，使其能够承担批量提交、失败重试、取消控制和更细致的状态展示。",
    ),
    (
        "从更长的使用周期来看，系统还可以继续扩展到权限控制、日志查看、调用审计与合规提示等方面。",
        "如果从更长的使用周期来看，系统还可以继续扩展到权限控制、日志查看、调用审计与合规提示等方面。人脸替换系统能够运行只是起点，后续更关键的是让处理过程可追溯、问题可回查、结果可解释。这些工作未必需要在毕业阶段一次性全部完成，但它们确实构成了项目继续深入时必须面对的方向。",
    ),
]

SECTION_OVERVIEWS = {
    "2.2 基于OpenCV与三角剖分的传统人脸替换方法": "本节围绕当前项目中的传统处理路径展开说明。内容将先交代 process_frame_traditional、advanced_face_swap 等入口如何组织传统换脸流程，再说明 Delaunay 三角剖分与局部仿射变换如何完成纹理迁移，最后讨论 seamlessClone 与颜色校正等补偿步骤在边界处理中的作用，从而把传统方案在代码中的真实职责说明清楚。",
    "2.3 围绕dlib、InsightFace与inswapper_128的深度学习处理链": "本节讨论系统中的深度学习主路径如何在实际代码中落地。正文将依次说明 dlib、InsightFace 与 inswapper_128 在工程链路中的分工，解释身份表征与目标属性保持的基本思路，并结合损失表达与推理流程说明该方案为什么能够成为当前系统的主要处理路径。",
    "5.7 关键线程与类协作实现": "本节从类与线程协作的角度解释桌面端原型为什么能够在长耗时处理中保持基本响应。内容先说明界面控制类、处理线程与底层引擎各自承担的职责，再分析信号传递、状态清理和线程切换如何共同支撑视频模式与摄像头模式的连续运行。",
    "6.1 测试环境与测试目标": "本节先交代测试所面对的真实环境，再说明本轮验证到底想检查什么。具体而言，正文会依次说明测试定位、运行环境组成以及样本范围，使后续功能验证、稳定性观察和案例分析都建立在可追溯的前提之上。",
}

FIGURE_EXPLANATIONS = {
    "图2-1 人脸替换典型处理流程图": "该图把当前项目中的换脸任务拆解为素材进入系统、线程启动、换脸处理、结果输出与后续回放几个关键阶段。它说明本文讨论的对象并不是孤立的一次模型推理，而是一条由前端交互、底层引擎和结果管理共同组成的完整处理链。",
    "图3-1 主要用户业务场景图": "该图直观展示了视频模式与摄像头模式两类核心业务场景的差异。前者更强调稳定输出、结果保存与后续回放，后者则更关注实时显示、目标人脸切换和界面响应速度，因此后续需求分析必须同时兼顾两类场景。",
    "图3-2 系统用例关系图": "该图从用户操作顺序的角度呈现了各项用例之间的依赖关系。通过它可以看出，视频模式和摄像头模式虽然共享素材选择与结果查看等基础动作，但在前置条件和执行路径上仍然存在清晰区分，这为后续界面状态联动提供了依据。",
    "图4-1 系统总体架构图": "该图从分层角度概括了用户交互层、业务控制层、媒体处理层和数据持久化层之间的关系。它强调系统的可用性并不是由单一模块决定的，而取决于界面逻辑、处理线程、底层引擎和数据记录是否能够稳定协同。",
    "图4-2 数据流与接口流示意图": "该图进一步展示了素材、参数、处理结果与元数据在不同模块之间的流动路径。借助这一视角可以看到，前端并非直接处理全部资源，而是依靠后端接口和本地文件共同完成组织与回放，这也是系统分层设计的重要体现。",
    "图4-3 数据库关系示意图": "该图将 FaceImage、InputVideo、OutputVideo 与 ProcessingTask 之间的关联关系以图形方式呈现出来。相较于单独阅读字段表，关系示意图更能说明结果记录如何回溯到源素材，以及任务管理能力未来准备沿着怎样的结构继续扩展。",
    "图4-4 系统部署结构图": "该图说明了当前系统采用的本机部署形态，即桌面前端、本地 Django 服务、SQLite 数据库与媒体目录共同运行在同一台设备上。这样的部署方式虽然不追求生产级分布式能力，却非常符合毕业设计阶段对调试便利性和演示稳定性的要求。",
    "图5-1 系统主界面截图": "该图展示了系统主界面在单窗口条件下整合素材选择、预览显示、参数控制与状态反馈的整体布局。这样的组织方式能够缩短用户在一次演示中的操作路径，使换脸处理、结果查看与参数调整尽量保持在同一工作面内完成。",
    "图5-2 主界面局部功能区域截图": "该图对主界面中的局部功能区域进行了放大，便于观察模式切换、素材选择、状态提示和控制按钮之间的配合关系。通过这一局部视角，可以更直观地理解系统如何在同一前端中兼容视频模式和摄像头模式两类工作流。",
    "图5-3 视频模式处理流程图": "该图展示了视频模式从素材校验、参数同步到线程执行和结果写出的完整顺序。它对应正文中 startProcessing、VideoProcessingThread、processingFinished 与 saveOutputVideoToDatabase 等关键调用关系，能够帮助读者把抽象流程与实际代码入口对应起来。",
    "图5-4 视频模式输入与输出样例": "该图将视频模式中的输入素材与输出结果并置展示，有助于把前面的流程说明落实到真实样例上。它的意义不仅在于展示视觉效果，更在于证明前端选择、后台处理、结果写出和回放验证之间已经形成可以执行的闭环。",
    "图5-5 摄像头模式处理流程图": "该图突出摄像头模式在持续采集、实时显示和动态切换方面的特殊要求。与视频模式相比，这一路径并不以单次输出文件为目标，而更强调线程持续运行期间的界面响应、状态反馈和人脸切换控制。",
    "图5-6 摄像头模式运行界面截图": "该图展示了摄像头模式运行时的界面状态，包括实时画面、控制按钮和状态反馈区域。通过该图可以更直观地理解用户在实时处理过程中如何启停换脸、观察结果并完成快照保存。",
    "图5-7 关键线程与模块协作时序图": "该图从时序角度展示了界面控制类、处理线程、底层引擎与数据管理组件之间的调用关系。它说明系统之所以能够在长耗时处理过程中维持基本响应，依赖的不只是线程本身，而是线程职责划分、信号传递和状态清理机制的共同作用。",
    "图6-1 系统样本数据统计图": "该图直观反映了当前系统中人脸图片、输入视频与输出视频三类资源的数量分布。结合数据库记录和本地目录状态，可以看出本轮测试并非建立在理想化大样本之上，而是以现有真实项目材料作为工程验证基础。",
    "图6-2 最近输入视频的时长与帧率分布图": "该图把最近输入视频的时长与帧率分布情况可视化，便于理解当前测试样本的大致规模和多样性。它说明系统目前主要面对中短时长视频，也解释了为什么本轮测试更强调流程打通和结果回放，而不是长视频批量处理能力。",
    "图6-3 真实照片源人脸条件下的代表性结果对比图": "该图展示了真实照片源人脸条件下的代表性输出结果。通过这些样例可以观察到，系统在常规照片素材上更容易保持身份特征和基础融合效果，因此这一组图主要用于说明当前主路径在常见素材条件下的可用表现。",
    "图6-4 绘制源人脸条件下的代表性结果对比图": "该图展示了绘制源人脸条件下的代表性输出结果。与真实照片相比，这类素材在纹理细节、光照关系和面部结构表达上更不规则，因此该组图更能反映系统在非常规输入条件下面临的适配难点与结果差异。",
}


def find_source() -> Path:
    matches = [p for p in ROOT.glob("*修订版6.docx") if not p.name.startswith("~$")]
    if not matches:
        raise FileNotFoundError("Could not find the non-temp revision 6 docx.")
    return matches[0]


def build_output_path(source: Path) -> Path:
    match = re.search(r"修订版(\d+)$", source.stem)
    if match:
        next_rev = int(match.group(1)) + 1
        prefix = source.stem[: match.start()]
        candidate = source.with_name(prefix + f"修订版{next_rev}" + source.suffix)
        while candidate.exists():
            next_rev += 1
            candidate = source.with_name(prefix + f"修订版{next_rev}" + source.suffix)
        return candidate
    candidate = source.with_name(source.stem + "_终检版" + source.suffix)
    counter = 1
    while candidate.exists():
        candidate = source.with_name(source.stem + f"_终检版{counter}" + source.suffix)
        counter += 1
    return candidate


def normalized_title(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def is_special_title(text: str) -> bool:
    compact = normalized_title(text)
    return compact in {"摘要", "目录"} or text.strip() in SPECIAL_TITLES


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def paragraph_has_math(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))


def is_formula_number(text: str) -> bool:
    return bool(re.fullmatch(r"（\d+）", text.strip()))


def is_caption(text: str) -> bool:
    return text.startswith("图") or text.startswith("表")


def is_reference_entry(text: str) -> bool:
    return bool(re.match(r"^\[\d+\]", text.strip()))


def is_keyword_line(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("关键词：") or stripped.startswith("Keywords:")


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def replace_paragraph_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.bold = bold


def find_paragraph_by_exact_text(doc: DocumentObject, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if (paragraph.text or "").strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_by_prefix(doc: DocumentObject, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if (paragraph.text or "").strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found for prefix: {prefix}")


def next_nonempty_content_paragraph(paragraph: Paragraph) -> Paragraph | None:
    current = paragraph._p.getnext()
    while current is not None:
        if current.tag != qn("w:p"):
            current = current.getnext()
            continue
        candidate = Paragraph(current, paragraph._parent)
        if (candidate.text or "").strip() or paragraph_has_drawing(candidate) or paragraph_has_math(candidate):
            return candidate
        current = current.getnext()
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def clear_fixed_line_spacing(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.attrib.pop(qn("w:line"), None)
        spacing.attrib.pop(qn("w:lineRule"), None)


def set_exact_20pt(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = LINE_SPACING


def normalize_english_colons(text: str) -> str:
    pattern = re.compile(r"([A-Za-z][A-Za-z0-9_./()+\- ]{0,40})：")

    def repl(match: re.Match[str]) -> str:
        return match.group(1).rstrip() + ": "

    return pattern.sub(repl, text)


def rewrite_target_paragraphs(doc: DocumentObject) -> int:
    replacements = 0
    for prefix, new_text in PARAGRAPH_REWRITES:
        paragraph = find_paragraph_by_prefix(doc, prefix)
        replace_paragraph_text(paragraph, new_text, bold=False)
        replacements += 1
    return replacements


def insert_section_overviews(doc: DocumentObject) -> int:
    inserted = 0
    for heading_text, overview_text in SECTION_OVERVIEWS.items():
        heading = find_paragraph_by_exact_text(doc, heading_text)
        next_para = next_nonempty_content_paragraph(heading)
        if next_para is None:
            continue
        next_text = (next_para.text or "").strip()
        if next_text == overview_text:
            continue
        if next_para.style.name == "Heading 3":
            paragraph = insert_paragraph_before(next_para, overview_text, style="Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            inserted += 1
        elif next_para.style.name != "Heading 3":
            replace_paragraph_text(next_para, overview_text, bold=False)
    return inserted


def insert_figure_explanations(doc: DocumentObject) -> int:
    updated = 0
    for caption_text, explanation_text in FIGURE_EXPLANATIONS.items():
        caption = find_paragraph_by_exact_text(doc, caption_text)
        next_para = next_nonempty_content_paragraph(caption)
        if next_para is None:
            insert_paragraph_after(caption, explanation_text, style="Normal")
            updated += 1
            continue

        next_text = (next_para.text or "").strip()
        next_style = next_para.style.name if next_para.style else ""
        if (
            next_text
            and not paragraph_has_drawing(next_para)
            and next_style not in {"Heading 1", "Heading 2", "Heading 3"}
            and not is_caption(next_text)
            and not is_reference_entry(next_text)
        ):
            replace_paragraph_text(next_para, explanation_text, bold=False)
        else:
            insert_paragraph_after(caption, explanation_text, style="Normal")
        updated += 1
    return updated


def remove_empty_headings(doc: DocumentObject) -> int:
    to_remove = [
        paragraph
        for paragraph in doc.paragraphs
        if (paragraph.style.name if paragraph.style else "") in {"Heading 1", "Heading 2", "Heading 3"}
        and not (paragraph.text or "").strip()
    ]
    for paragraph in to_remove:
        remove_paragraph(paragraph)
    return len(to_remove)


def apply_page_breaks(doc: DocumentObject) -> int:
    updated = 0
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style = paragraph.style.name if paragraph.style else ""
        if style != "Heading 1":
            continue
        if (text.startswith("第") and "章" in text[:4]) or text in {"参考文献", "致谢"}:
            paragraph.paragraph_format.page_break_before = True
            updated += 1
    return updated


def format_document(doc: DocumentObject) -> int:
    in_content = False
    image_count = 0
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style = paragraph.style.name if paragraph.style else ""

        if normalized_title(text) == "摘要":
            in_content = True

        if not in_content:
            continue

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        if paragraph_has_drawing(paragraph):
            image_count += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            clear_fixed_line_spacing(paragraph)
            continue

        if paragraph_has_math(paragraph):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            clear_fixed_line_spacing(paragraph)
            continue

        if not text:
            set_exact_20pt(paragraph)
            continue

        if style in {"Heading 1", "Heading 2", "Heading 3"} or is_special_title(text):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_exact_20pt(paragraph)
            continue

        if is_caption(text) or is_reference_entry(text) or is_keyword_line(text) or is_formula_number(text):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_exact_20pt(paragraph)
            continue

        paragraph.paragraph_format.first_line_indent = BODY_FIRST_INDENT
        if paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_exact_20pt(paragraph)
    return image_count


def normalize_colons(doc: DocumentObject) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph) or paragraph_has_math(paragraph):
            continue
        text = paragraph.text or ""
        new_text = normalize_english_colons(text)
        if new_text != text:
            bold = any(bool(run.bold) for run in paragraph.runs if run.text.strip())
            replace_paragraph_text(paragraph, new_text, bold=bold)
            changed += 1
    return changed


def verify_output(doc: DocumentObject) -> tuple[int, int, int]:
    pagebreaks = 0
    missing_overviews = 0
    figure_gaps = 0

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style = paragraph.style.name if paragraph.style else ""
        if style == "Heading 1" and ((text.startswith("第") and "章" in text[:4]) or text in {"参考文献", "致谢"}):
            if paragraph.paragraph_format.page_break_before:
                pagebreaks += 1

    for heading_text in SECTION_OVERVIEWS:
        heading = find_paragraph_by_exact_text(doc, heading_text)
        next_para = next_nonempty_content_paragraph(heading)
        if next_para is None or next_para.style.name == "Heading 3":
            missing_overviews += 1

    for caption_text in FIGURE_EXPLANATIONS:
        caption = find_paragraph_by_exact_text(doc, caption_text)
        next_para = next_nonempty_content_paragraph(caption)
        next_text = (next_para.text or "").strip() if next_para is not None else ""
        if (
            next_para is None
            or paragraph_has_drawing(next_para)
            or next_para.style.name in {"Heading 1", "Heading 2", "Heading 3"}
            or is_caption(next_text)
        ):
            figure_gaps += 1

    return pagebreaks, missing_overviews, figure_gaps


def main() -> None:
    source = find_source()
    output = build_output_path(source)
    shutil.copy2(source, output)

    doc = Document(str(output))

    removed_headings = remove_empty_headings(doc)
    rewritten = rewrite_target_paragraphs(doc)
    inserted_overviews = insert_section_overviews(doc)
    updated_figures = insert_figure_explanations(doc)
    pagebreak_count = apply_page_breaks(doc)
    colon_changes = normalize_colons(doc)
    image_count = format_document(doc)
    doc.save(str(output))

    final_doc = Document(str(output))
    verified_pagebreaks, missing_overviews, figure_gaps = verify_output(final_doc)

    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    print(f"REMOVED_EMPTY_HEADINGS={removed_headings}")
    print(f"REWRITTEN_PARAGRAPHS={rewritten}")
    print(f"INSERTED_SECTION_OVERVIEWS={inserted_overviews}")
    print(f"UPDATED_FIGURE_EXPLANATIONS={updated_figures}")
    print(f"PAGE_BREAK_TARGETS_SET={pagebreak_count}")
    print(f"ENGLISH_COLON_FIXES={colon_changes}")
    print(f"IMAGE_PARAGRAPHS_NORMALIZED={image_count}")
    print(f"VERIFY_PAGEBREAKS={verified_pagebreaks}")
    print(f"VERIFY_MISSING_OVERVIEWS={missing_overviews}")
    print(f"VERIFY_FIGURE_GAPS={figure_gaps}")


if __name__ == "__main__":
    main()
