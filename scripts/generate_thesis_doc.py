# -*- coding: utf-8 -*-
from generate_thesis_doc_v2 import main

if __name__ == "__main__":
    main()
    raise SystemExit

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
ASSETS_DIR = BASE_DIR / "thesis_assets"
DEFAULT_OUTPUT_PATH = BASE_DIR / "sdu_faceswap_thesis_fixed.docx"


def set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_format(paragraph, line_spacing=1.5, first_line_chars=2, space_before=0, space_after=0, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line_spacing == 1.5 else WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.5 if line_spacing == 1.5 else 1
    fmt.first_line_indent = Pt(24) if first_line_chars == 2 else Pt(0)
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align


def add_text_paragraph(doc, text, size=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       first_line_chars=2, line_spacing=1.5, east_asia="宋体", ascii_font="Times New Roman"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_fonts(r, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    set_paragraph_format(
        p,
        line_spacing=line_spacing,
        first_line_chars=first_line_chars,
        align=align,
    )
    return p


def add_center_title(doc, text, size=Pt(18), east_asia="黑体", bold=True, ascii_font="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_fonts(r, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    set_paragraph_format(p, first_line_chars=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level, 3)}"]
    r = p.add_run(text)
    if level == 1:
        set_run_fonts(r, east_asia="黑体", size=Pt(16), bold=True)
        set_paragraph_format(p, first_line_chars=0, space_before=10, space_after=6)
    elif level == 2:
        set_run_fonts(r, east_asia="黑体", size=Pt(14), bold=True)
        set_paragraph_format(p, first_line_chars=0, space_before=6, space_after=6)
    else:
        set_run_fonts(r, east_asia="黑体", size=Pt(12), bold=True)
        set_paragraph_format(p, first_line_chars=0, space_before=6, space_after=6)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_fonts(r, east_asia="宋体", size=Pt(10.5), bold=True)
    set_paragraph_format(p, line_spacing=1, first_line_chars=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_picture(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    return p


def add_field(paragraph, field_code, placeholder=""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    if placeholder:
        placeholder_run = paragraph.add_run(placeholder)
        set_run_fonts(placeholder_run, east_asia="宋体", size=Pt(12), bold=False)
    end_run = paragraph.add_run()
    end_run._r.append(fld_end)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_fonts(run, east_asia="Times New Roman", ascii_font="Times New Roman", size=Pt(9), bold=False)


def set_page_number_format(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    for node in sect_pr.xpath("./w:pgNumType"):
        sect_pr.remove(node)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))
    sect_pr.append(pg)


def format_section(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_fonts(run, east_asia="宋体", size=Pt(10.5), bold=False)


def build_document(output_path):
    doc = Document()
    for section in doc.sections:
        format_section(section)

    doc.core_properties.title = "人脸替换系统的设计与实现"
    doc.core_properties.author = "叶俊"

    for _ in range(4):
        doc.add_paragraph()
    add_center_title(doc, "山东大学本科毕业论文（设计）", Pt(22))
    for _ in range(2):
        doc.add_paragraph()
    add_center_title(doc, "人脸替换系统的设计与实现", Pt(18))
    add_center_title(doc, "Design and Implementation of a Face Swapping System Based on PyQt5 and Django", Pt(16))
    for _ in range(4):
        doc.add_paragraph()

    for item in [
        "姓名：叶俊",
        "学号：202200201151",
        "学院：软件学院",
        "年级：2022级",
        "指导教师：待填写",
        "完成时间：2026年4月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        set_run_fonts(r, east_asia="宋体", size=Pt(14), bold=False)
        set_paragraph_format(p, first_line_chars=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_center_title(doc, "摘    要", Pt(18))
    add_text_paragraph(
        doc,
        "随着深度学习与数字媒体技术的快速发展，面向普通用户的图像生成和视频编辑工具逐步从专业实验环境走向桌面端应用。"
        "人脸替换技术作为计算机视觉、图像处理与人机交互交叉融合的典型场景，既具有较高的研究价值，也具有明确的工程实现意义。"
        "针对传统人脸替换流程操作繁琐、前后端数据割裂、处理结果难以统一管理的问题，本文结合课程设计项目，设计并实现了一个基于 PyQt5 与 Django 的人脸替换系统。"
        "系统以前端桌面界面为交互入口，以 Django REST 服务为数据管理中枢，结合 OpenCV、InsightFace 与本地文件存储机制，实现了人脸图片导入、输入视频选择、视频模式换脸、摄像头模式切换、处理结果保存与输出视频归档等功能。"
        "在系统设计过程中，本文围绕可用性、扩展性和工程可维护性三个维度展开工作：一是在界面层采用视频模式与摄像头模式双界面组织方式，降低用户操作复杂度；二是在业务层引入数据库管理模块，将图片、视频与输出结果统一纳入后端接口管理；三是在处理层兼顾传统方法与 InsightFace 方法，提升系统对不同运行环境的适配能力。"
        "结合项目当前数据库样本统计结果，系统已完成 19 张人脸图片、13 个输入视频与 5 个输出视频的管理记录，可支持本地文件路径保存、缩略图展示与结果回放。"
        "本文最终完成了系统需求分析、总体架构设计、关键模块实现、数据结构设计与测试验证，并通过项目图片样例与界面截图展示了系统在实际场景中的应用效果。"
        "研究结果表明，该系统能够较为完整地支撑桌面端人脸替换的主要业务流程，具备一定的演示价值与后续扩展基础。"
        "但同时，系统在批量任务调度、模型性能评估、异常恢复与隐私风险控制方面仍有进一步完善空间。"
    )
    p = doc.add_paragraph()
    r1 = p.add_run("关键词：")
    set_run_fonts(r1, east_asia="黑体", size=Pt(12), bold=True)
    r2 = p.add_run("人脸替换；PyQt5；Django；InsightFace；桌面应用")
    set_run_fonts(r2, east_asia="宋体", size=Pt(12), bold=False)
    set_paragraph_format(p, first_line_chars=0)

    doc.add_page_break()

    add_center_title(doc, "ABSTRACT", Pt(18))
    add_text_paragraph(
        doc,
        "With the rapid development of deep learning and digital media technologies, image generation and video editing tools have gradually moved from research laboratories to desktop applications for ordinary users. "
        "Face swapping is a representative scenario that integrates computer vision, image processing, and human-computer interaction. It therefore has both academic value and practical engineering significance. "
        "To address the problems of complicated operation flow, fragmented data management, and weak result archiving in traditional desktop prototypes, this thesis designs and implements a face swapping system based on PyQt5 and Django. "
        "The system takes a PyQt5 desktop interface as the interaction entry, uses Django REST services as the data management hub, and combines OpenCV, InsightFace, and local file storage to support face image import, input video selection, video-mode swapping, camera-mode switching, result saving, and processed video archiving. "
        "From the perspective of usability, scalability, and maintainability, the system introduces a dual-mode interface, a unified database management module, and a hybrid processing strategy that supports both traditional methods and InsightFace-based methods. "
        "According to the current project records, the system already manages 19 face images, 13 input videos, and 5 output videos, which provides a realistic basis for prototype verification. "
        "The implementation shows that the proposed system can support the main workflow of desktop face swapping and provides a solid foundation for future work such as batch task scheduling, stronger performance evaluation, error recovery, and privacy protection.",
        east_asia="Times New Roman",
        ascii_font="Times New Roman",
    )
    p = doc.add_paragraph()
    r1 = p.add_run("Key Words: ")
    set_run_fonts(r1, east_asia="Times New Roman", ascii_font="Times New Roman", size=Pt(12), bold=True)
    r2 = p.add_run("face swapping, PyQt5, Django, InsightFace, desktop application")
    set_run_fonts(r2, east_asia="Times New Roman", ascii_font="Times New Roman", size=Pt(12), bold=False)
    set_paragraph_format(p, first_line_chars=0)

    sec_toc = doc.add_section(WD_SECTION.NEW_PAGE)
    format_section(sec_toc)
    sec_toc.header.is_linked_to_previous = False
    sec_toc.footer.is_linked_to_previous = False
    set_page_number_format(sec_toc, "upperRoman", 1)
    footer = sec_toc.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    add_center_title(doc, "目    录", Pt(18))
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "右键更新目录")

    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    format_section(sec_body)
    sec_body.header.is_linked_to_previous = False
    sec_body.footer.is_linked_to_previous = False
    set_page_number_format(sec_body, "decimal", 1)
    header = sec_body.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run("山东大学本科毕业论文（设计）")
    set_run_fonts(hr, east_asia="宋体", size=Pt(9), bold=False)
    footer = sec_body.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    body_items = [
        ("1 绪论", 1),
        ("1.1 研究背景与意义", 2),
        ("近年来，人工智能生成内容技术快速进入图像编辑、视频制作和数字娱乐等应用场景。与静态图片编辑相比，视频人脸替换同时涉及目标人脸检测、关键点定位、区域融合、颜色校正、帧间一致性和结果回放等多个环节，因此其工程实现难度更高。对于软件工程专业的本科毕业设计而言，人脸替换系统能够覆盖前端交互、后端服务、数据管理、算法集成与系统测试等完整链路，是一个兼具研究性与实践性的综合型课题。", "p"),
        ("本项目以桌面端场景为出发点，试图解决传统实验性程序“能跑但不好用”的问题。许多开源原型往往仅提供命令行脚本或简单界面，用户在导入素材、切换模式、保存结果和管理输出时需要频繁切换目录与参数，这会显著增加使用成本。本文通过将 PyQt5 桌面界面与 Django REST 后端结合，使界面操作、任务处理和数据归档能够在统一系统中完成，从而提升系统的可操作性、可维护性和可展示性。", "p"),
        ("1.2 国内外研究现状", 2),
        ("从研究发展脉络来看，人脸替换技术经历了从基于几何对齐与泊松融合的传统方法，到基于深度神经网络的高保真生成方法的演化过程。早期方案主要依赖人脸检测、关键点匹配和局部融合，对光照变化、姿态偏移与遮挡情况较为敏感；随着 FaceNet、ArcFace、VGGFace2 等人脸表示学习方法的发展，深度特征对身份保持和语义一致性的支持显著增强，为高质量换脸奠定了基础[4-6]。近几年，DeepFaceLab、FaceShifter、SimSwap 等框架进一步推动了换脸技术的工程化与应用化[8-10]。", "p"),
        ("与此同时，围绕深度伪造的检测、可信使用和伦理治理也成为重要研究方向。FaceForensics++、人脸伪造检测和图像篡改识别相关研究指出，视觉真实感的增强并不意味着系统可以脱离安全边界独立存在[7][11][12]。因此，本文在实现系统功能时，将工程可用性放在首位，而不以追求极端视觉效果为唯一目标，重点关注桌面交互、数据可追踪、模块解耦和可扩展集成，这也更符合本科毕业设计的实践定位。", "p"),
        ("1.3 研究内容与论文结构", 2),
        ("围绕“人脸替换系统的设计与实现”这一主题，本文完成了以下几项工作：第一，结合现有代码与项目结构，对系统的业务场景、用户操作路径与功能边界进行了分析；第二，设计了由 PyQt5 前端、Django REST 后端、SQLite 数据库和本地文件系统构成的总体架构；第三，分析并实现了素材管理、视频处理、结果回放、摄像头模式切换和输出归档等关键模块；第四，基于当前项目数据库样本与界面截图对系统进行了验证和展示。", "p"),
        ("全文共分为六章。第一章介绍研究背景、意义与相关研究现状；第二章阐述系统涉及的关键技术与需求分析；第三章给出总体架构、业务流程和数据设计；第四章说明系统关键模块的实现过程；第五章结合样本数据、功能测试与界面截图展示系统效果；第六章总结本文工作并给出后续改进方向。", "p"),
        ("2 相关技术与需求分析", 1),
        ("2.1 关键技术基础", 2),
        ("系统前端采用 PyQt5 实现。相较于单纯的 Web 页面，桌面端程序在访问本地视频、摄像头设备和文件路径时更直接，且便于构建视频播放控件、缩略图列表和模式切换动画。项目中的 EnhancedFaceSwapUI 类负责主界面初始化、视频模式与摄像头模式切换、素材列表展示、处理进度更新和文件上传等核心交互逻辑，构成了系统的人机交互层。", "p"),
        ("系统后端采用 Django 与 Django REST Framework 构建。后端为图片、输入视频、输出视频和处理任务提供统一接口，并通过模型层将元数据存储到 SQLite 数据库中。相较于将所有逻辑都写在桌面端，REST 化的后端设计有利于解耦界面与数据管理，使上传、读取、归档和后续扩展变得更加清晰。项目中的 FaceImageViewSet、InputVideoViewSet 与 OutputVideoViewSet 分别负责三类核心资源的 CRUD 逻辑。", "p"),
        ("在图像与视频处理方面，系统结合 OpenCV 与 InsightFace 完成人脸检测、视频帧读取、颜色空间转换、缩略图生成和换脸结果融合。OpenCV 适合处理视频解码、帧提取和基本图像处理任务；InsightFace 则为高质量身份特征建模与人脸替换提供支持。考虑到不同环境下模型可用性存在差异，系统同时保留传统方法作为替代路径，以提升程序在多种硬件与软件环境中的可运行性。", "p"),
        ("在数据管理方面，系统通过数据库表记录图片、视频与结果文件的路径、分辨率、帧率、时长和文件大小等元数据，并配合本地文件系统保存真实素材文件。这种“数据库记录+本地路径”的策略既避免了桌面端重复扫描目录带来的效率损失，又增强了素材查询和结果追踪能力，适合原型系统的快速实现。", "p"),
        ("2.2 功能需求分析", 2),
        ("从用户视角出发，系统需要至少支持五类核心功能。其一，用户能够导入或选择待替换的人脸图片，并以缩略图方式直观浏览；其二，用户能够导入视频并自动生成输出路径；其三，系统能够启动视频换脸处理并给出进度反馈；其四，系统能够在摄像头模式与视频模式之间平滑切换；其五，处理结果能够被保存、回放并记录到数据库中。", "p"),
        ("除显式功能外，系统还需要满足若干非功能性需求。首先是易用性，界面应尽量减少命令式输入，使用按钮、列表和状态栏降低操作门槛；其次是可维护性，各模块之间应保持相对独立，便于后续修改；再次是可扩展性，算法层和数据层应预留替换空间，支持增加新模型或新接口；最后是可追踪性，已导入与已处理的文件都应具备统一的记录方式，避免因目录混乱而丢失结果。", "p"),
        ("2.3 可行性分析", 2),
        ("从技术可行性看，本系统所需的关键技术均具备成熟的开源基础。PyQt5 可满足复杂桌面界面的构建需求，Django REST Framework 能较稳定地支撑数据接口，OpenCV 具备完善的视频处理能力，InsightFace 提供了成熟的人脸分析与换脸模型。因此，在本科毕业设计的时间和资源条件下，开发一个可演示、可验证、可扩展的人脸替换原型系统是可行的。", "p"),
        ("从工程可行性看，项目已形成相对完整的目录结构，包括前端、后端、模型文件、输入数据、输出视频和日志目录。现有数据库中已保存 19 张图片、13 个输入视频和 5 个输出视频记录，这说明系统原型已经具备一定的数据基础和运行痕迹。本文的工作重点不是从零开始构造算法，而是在已有实现上进一步梳理系统结构、提升交互组织方式，并将其整理为符合毕业论文要求的工程成果。", "p"),
    ]

    for item in body_items:
        if item[1] == "p":
            add_text_paragraph(doc, item[0])
        else:
            add_heading(doc, item[0], item[1])

    add_heading(doc, "3 系统分析与总体设计", 1)
    add_heading(doc, "3.1 总体架构设计", 2)
    add_text_paragraph(doc, "系统总体采用“桌面前端+REST 后端+本地算法处理+数据库存档”的分层模式。前端界面负责素材展示、参数控制和结果预览；后端负责图片、视频与输出对象的统一管理；核心换脸引擎负责视频帧处理、人脸检测与结果生成；SQLite 和本地文件系统共同承担数据持久化任务。相较于单体式脚本结构，这种分层模式更有利于后续维护与模块替换。")
    add_picture(doc, ASSETS_DIR / "diagram_architecture.png", 15.5)
    add_caption(doc, "图 3.1 系统总体架构图")
    add_text_paragraph(doc, "从调用关系来看，用户首先在前端界面上完成图片和视频的选择；随后界面将路径或文件信息交给数据库管理模块；当用户触发处理任务时，前端调用换脸引擎执行视频模式或摄像头模式的核心逻辑；处理结束后，结果文件被保存到本地输出目录，并通过后端接口记录到输出视频表。整个流程覆盖了交互层、业务层和数据层的完整闭环。")
    add_heading(doc, "3.2 业务流程设计", 2)
    add_text_paragraph(doc, "在视频模式下，系统的业务流程可概括为“选择人脸图片、选择输入视频、自动生成输出路径、设置处理参数、启动视频处理、保存并回放结果”六个步骤。界面层通过列表项选择与按钮点击收集用户输入，核心线程 VideoProcessingThread 负责在后台执行处理任务，并通过信号机制实时更新进度条与状态文本，避免长时间计算阻塞界面主线程。")
    add_text_paragraph(doc, "在摄像头模式下，系统首先初始化摄像头采集线程，并在需要时读取用户选中的目标人脸图像。摄像头线程持续获取视频流帧，在启用处理开关后将最新帧送入换脸方法，并将处理后的结果回传界面显示区域。这一机制虽然更强调实时性而非高精度离线渲染，但能有效展示系统的交互能力和线程组织能力。")
    add_heading(doc, "3.3 数据结构与数据库设计", 2)
    add_text_paragraph(doc, "后端模型层定义了 FaceImage、InputVideo、OutputVideo、ProcessingTask 和 SystemConfig 等核心实体。其中，FaceImage 记录人脸图片文件名、本地路径、宽高和大小信息；InputVideo 记录视频分辨率、时长、帧率与文件大小；OutputVideo 则在此基础上增加处理方法、处理状态和处理进度等信息。通过这些元数据，系统能够对素材和结果进行统一管理。")
    add_caption(doc, "表 3.1 核心数据表设计")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "数据表"
    table.rows[0].cells[1].text = "主要字段"
    table.rows[0].cells[2].text = "作用说明"
    for row in [
        ("face_images", "original_filename, local_path, width, height, file_size", "记录输入人脸图片元数据"),
        ("input_videos", "original_filename, local_path, duration, width, height, fps", "记录待处理视频信息"),
        ("output_videos", "filename, file, processing_method, status, progress", "记录输出结果与处理状态"),
        ("processing_tasks", "task_type, status, progress, processing_params", "扩展任务调度与过程管理"),
    ]:
        cells = table.add_row().cells
        for idx, txt in enumerate(row):
            cells[idx].text = txt
    style_table(table)
    add_text_paragraph(doc, "数据库设计与本地文件路径绑定是本系统的一个工程特色。当前后端接口在创建图片和视频对象时，会直接读取本地路径获取元数据并存入数据库，而不是强制用户重新上传二进制文件。这种方式降低了桌面端使用门槛，也使前端能够基于数据库记录快速恢复素材列表，提高了启动后的可用性。")
    add_heading(doc, "3.4 界面设计", 2)
    add_text_paragraph(doc, "增强版前端界面以“左侧内容显示、右侧控制面板、顶部模式切换”的布局组织主要操作元素。左侧用于视频播放、摄像头显示和结果预览，右侧用于图片选择、视频选择、输出路径设置和高级参数调整，顶部提供视频模式与摄像头模式切换按钮。该布局遵循“主任务在左、控制项在右”的习惯，能够减少界面跳转带来的认知负担。")
    add_picture(doc, BASE_DIR / "image" / "ui_demo.png", 15.5)
    add_caption(doc, "图 3.2 系统主界面截图")
    add_text_paragraph(doc, "界面设计还引入了状态栏、圆形进度条、播放控制栏和模式切换动画等元素，使系统不仅具备基础功能，也具备较好的演示性。对于本科毕业设计而言，这种明确的可视化反馈有助于在答辩或展示时直观说明处理流程和系统状态。")

    add_heading(doc, "4 系统实现", 1)
    add_heading(doc, "4.1 前端模块实现", 2)
    add_text_paragraph(doc, "前端入口位于 frontend/main.py，主要负责创建 QApplication 对象、配置高分屏属性并设置统一样式表。主界面类 EnhancedFaceSwapUI 在初始化过程中完成窗口属性设置、数据库管理器绑定、界面组件创建、文件列表加载和媒体播放器初始化等工作。由此，界面逻辑不再分散在多个脚本中，而是集中到统一的主窗口类中进行组织。")
    add_text_paragraph(doc, "在前端实现中，素材列表采用 QListWidget 呈现，图片与视频均以缩略图加文本形式显示，便于用户快速定位素材。用户点击列表项后，系统会记录所选图片和视频路径，并自动计算输出路径。对于视频结果播放，系统同时提供媒体播放器和基于 OpenCV 的备选播放方式，以增强兼容性。这些设计体现了工程实现中对“可操作”和“可运行”的优先考虑。")
    add_heading(doc, "4.2 后端接口实现", 2)
    add_text_paragraph(doc, "后端基于 Django REST Framework 的视图集实现资源管理。FaceImageViewSet.create() 与 InputVideoViewSet.create() 以本地路径为输入，读取文件元数据后创建数据库对象；OutputVideoViewSet.create() 则在接收处理结果文件后写入输出目录，并补充缩略图与状态信息。与纯文件夹扫描相比，这种接口化方式使素材读写具有更清晰的边界，也更便于后续接入分页、搜索和权限控制。")
    add_caption(doc, "表 4.1 主要 REST 接口设计")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "接口路径"
    table.rows[0].cells[1].text = "请求方法"
    table.rows[0].cells[2].text = "功能说明"
    for row in [
        ("/api/images/", "GET / POST", "获取图片列表与登记人脸图片"),
        ("/api/videos/", "GET / POST", "获取视频列表与登记输入视频"),
        ("/api/outputs/", "GET / POST", "保存与查询输出视频结果"),
        ("/api/tasks/", "GET / POST", "处理任务扩展接口"),
        ("/api/docs/", "GET", "查看 Swagger 文档"),
    ]:
        cells = table.add_row().cells
        for idx, txt in enumerate(row):
            cells[idx].text = txt
    style_table(table)
    add_text_paragraph(doc, "后端配置文件 settings.py 使用 SQLite 作为默认开发数据库，并启用了 rest_framework、corsheaders、django_filters 与 drf_yasg 等组件。结合 Swagger 文档页面，开发者可以更方便地调试接口与查看数据结构。这种后端组织方式不仅服务于当前桌面项目，也为未来可能的网页端、小程序端或其他客户端接入提供了接口基础。")
    add_heading(doc, "4.3 换脸处理流程实现", 2)
    add_text_paragraph(doc, "系统在算法层采取兼顾可用性与表现力的策略。对于具备相关依赖的环境，系统优先选择 InsightFace 路径，通过人脸分析器与换脸模型生成效果更自然的输出；当模型不可用或环境受限时，系统可以切换到传统处理方法，保证程序仍然能够完成基本功能。这样做避免了原型系统完全依赖单一模型造成的脆弱性。")
    add_text_paragraph(doc, "视频处理流程主要包括视频读取、目标人脸指定、逐帧处理、进度反馈、输出写入和结果保存六个阶段。为了防止界面卡死，项目将视频处理逻辑封装到 VideoProcessingThread 中，由后台线程定期向界面发送进度与状态信号。相较于同步执行方式，这种多线程模式更符合桌面应用的交互要求，也更利于后续将处理流程扩展为可取消的任务。")
    add_heading(doc, "4.4 视频模式与摄像头模式实现", 2)
    add_text_paragraph(doc, "视频模式强调离线处理与结果导出。用户可以从已加载视频列表中选择素材，并在右侧控制面板中设置输出路径、平滑度、颜色校正和检测器等参数。处理完成后，结果不仅保存在本地目录，还可通过后端接口写入 output_videos 表，实现结果留痕和后续管理。")
    add_text_paragraph(doc, "摄像头模式则更强调实时交互。通过 CameraProcessingThread，系统持续获取摄像头帧，并允许用户在开启换脸开关后应用目标人脸。虽然实时模式不以离线高质量渲染为目标，但它对线程调度、帧显示与界面响应提出了更高要求，因此能够较好地体现系统在交互层的综合能力。")
    add_heading(doc, "4.5 样例结果展示", 2)
    add_text_paragraph(doc, "为了展示系统处理效果，本文直接选取项目 image 目录中的样例图片进行对比说明。下列两组样例展示了不同输入素材在系统处理后的视觉输出结果，可用于辅助说明系统已经具备基础的人脸替换能力。需要说明的是，样例图更多用于系统展示与功能说明，不应替代更大规模的客观算法评测。")
    add_picture(doc, ASSETS_DIR / "compare_case_a.png", 15.2)
    add_caption(doc, "图 4.1 样例 A 的输入图像与换脸结果")
    add_picture(doc, ASSETS_DIR / "compare_case_b.png", 12.5)
    add_caption(doc, "图 4.2 样例 B 的输入图像与换脸结果")

    add_heading(doc, "5 系统测试与结果分析", 1)
    add_heading(doc, "5.1 测试环境与样本概况", 2)
    add_text_paragraph(doc, "本文的测试与展示主要基于当前项目目录、数据库记录和界面原型完成。系统采用 Windows 环境运行，数据层使用 backend-django/db.sqlite3 中的样本记录。根据统计结果，系统当前已记录 19 张人脸图片、13 个输入视频和 5 个输出视频，这些数据足以支撑原型系统的功能验证与展示。")
    add_picture(doc, ASSETS_DIR / "chart_asset_counts.png", 14.5)
    add_caption(doc, "图 5.1 系统样本数据概览")
    add_text_paragraph(doc, "对于输入视频样本，数据库中保留了分辨率、时长和帧率等元数据。统计最近 5 个视频记录可见，输入视频时长与帧率存在明显差异，这意味着系统在实际使用中需要兼顾短视频处理与不同帧率素材的兼容性。该特点也说明在工程实现中，统一的视频读取和播放控制逻辑是必要的。")
    add_picture(doc, ASSETS_DIR / "chart_video_metadata.png", 14.5)
    add_caption(doc, "图 5.2 最近 5 个输入视频的时长与帧率信息")
    add_heading(doc, "5.2 功能测试分析", 2)
    add_text_paragraph(doc, "围绕系统目标，本文从素材加载、路径生成、模式切换、结果归档和界面反馈等方面对系统进行了面向功能的验证。由于当前原型系统更强调工程可用性，因此测试重点不是算法精度指标，而是各业务环节是否能够连贯运行、界面是否能正确反馈、数据是否能够被完整记录。")
    add_caption(doc, "表 5.1 功能验证结果")
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "测试项目"
    table.rows[0].cells[1].text = "验证方式"
    table.rows[0].cells[2].text = "结果"
    table.rows[0].cells[3].text = "说明"
    for row in [
        ("图片列表加载", "读取数据库与本地图片目录", "符合预期", "可显示缩略图与文件名"),
        ("视频列表加载", "读取数据库与视频元数据", "符合预期", "支持缩略图与路径选择"),
        ("输出路径生成", "选择视频后自动生成输出名", "符合预期", "降低了用户手动输入成本"),
        ("模式切换", "视频模式与摄像头模式切换", "符合预期", "具备独立控制栏与状态提示"),
        ("结果归档", "输出视频写入数据库", "符合预期", "支持后续回放与查询"),
    ]:
        cells = table.add_row().cells
        for idx, txt in enumerate(row):
            cells[idx].text = txt
    style_table(table)
    add_text_paragraph(doc, "从测试结果看，系统的主要业务流程已经打通，尤其是在素材选择、输出管理与界面反馈方面表现较为完整。现有原型已经不仅是单纯的算法脚本，而是具备了基本软件系统形态。这一结果说明将桌面界面、后端接口和本地处理逻辑进行整合，是提升毕业设计项目完成度的重要手段。")
    add_heading(doc, "5.3 结果讨论", 2)
    add_text_paragraph(doc, "通过当前样例与截图可以看出，系统在人脸替换任务中已经能够产生可辨识的视觉结果，并具备基本的界面可视化能力。对于本科毕业设计而言，这类结果足以支撑“系统已实现并可运行”的结论。同时，系统在数据归档、界面组织和模式切换方面的设计，也使其具备一定的软件工程规范性，而不仅停留在算法实验层面。")
    add_text_paragraph(doc, "需要注意的是，数据库中的 processing_time 字段在当前项目中主要用于结果记录与展示，尚不能直接视为严格基准测试数据。因此，本文在结果分析中更多使用样本数量、界面流程、元数据分布和功能验证来支撑结论，而不夸大算法性能。这种处理方式有助于保证论文表述的客观性和可信度。")
    add_heading(doc, "5.4 存在问题与改进方向", 2)
    add_text_paragraph(doc, "尽管系统已经完成了原型实现，但仍存在若干不足：其一，当前系统的性能测试与质量评估还不够系统，缺少更大规模的数据集对比；其二，批量任务、异常恢复与任务取消等机制尚未完善；其三，隐私保护、数据脱敏和深度伪造风险提示尚需进一步纳入系统设计；其四，界面虽然已具备较好的演示性，但在跨平台兼容和细节打磨方面仍有改进空间。")

    add_heading(doc, "6 总结与展望", 1)
    add_text_paragraph(doc, "本文围绕“人脸替换系统的设计与实现”这一主题，结合现有项目代码与样本数据，完成了系统的需求分析、总体设计、关键实现与测试展示工作。论文以 PyQt5 与 Django 的协同架构为主线，分析了前端交互、后端接口、数据库管理和换脸处理逻辑之间的关系，并通过图表、界面截图和样例结果说明了系统的实际运行效果。")
    add_text_paragraph(doc, "总体来看，系统已经能够满足桌面端人脸替换原型的基本使用需求，具备素材管理、结果归档和多模式交互能力，达到了毕业设计阶段“可设计、可实现、可展示、可分析”的目标。后续研究可继续从三方面推进：一是增强算法层评测与模型切换能力，二是完善任务调度、异常恢复与批量处理机制，三是在工程层面补充权限控制、隐私保护与合规提示，使系统从原型进一步走向更加稳健的应用形态。")

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] 全国信息与文献标准化技术委员会. 信息与文献 参考文献著录规则: GB/T 7714-2015[S]. 北京: 中国标准出版社, 2015.",
        "[2] KING D E. Dlib-ml: A machine learning toolkit[J]. Journal of Machine Learning Research, 2009, 10: 1755-1758.",
        "[3] VIOLA P, JONES M. Rapid object detection using a boosted cascade of simple features[C]//Proceedings of the IEEE Computer Society Conference on Computer Vision and Pattern Recognition. Kauai: IEEE, 2001: I-511-I-518.",
        "[4] SCHROFF F, KALENICHENKO D, PHILBIN J. FaceNet: A unified embedding for face recognition and clustering[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Boston: IEEE, 2015: 815-823.",
        "[5] DENG J, GUO J, XUE N, et al. ArcFace: Additive angular margin loss for deep face recognition[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Long Beach: IEEE, 2019: 4690-4699.",
        "[6] CAO Q, SHEN L, XIE W, et al. VGGFace2: A dataset for recognising faces across pose and age[C]//2018 13th IEEE International Conference on Automatic Face and Gesture Recognition. Xi’an: IEEE, 2018: 67-74.",
        "[7] ROSSLER A, COZZOLINO D, VERDOLIVA L, et al. FaceForensics++: Learning to detect manipulated facial images[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. Seoul: IEEE, 2019.",
        "[8] LI L, BAO J, YANG H, et al. FaceShifter: Towards high fidelity and occlusion aware face swapping[EB/OL]. arXiv:1912.13457, 2019.",
        "[9] CHEN R, CHEN X, NI B, et al. SimSwap: An efficient framework for high fidelity face swapping[C]//Proceedings of the 28th ACM International Conference on Multimedia. New York: ACM, 2020.",
        "[10] PETROV I, CHEREPKOV A, AVRADYUK A, et al. DeepFaceLab: Integrated, flexible and extensible face-swapping framework[EB/OL]. arXiv:2005.05535, 2020.",
        "[11] LI Y, LYU S. Exposing deepfake videos by detecting face warping artifacts[C]//IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops. Long Beach: IEEE, 2019.",
        "[12] LI Y, CHANG M C, LYU S. In ictu oculi: Exposing AI generated fake face videos by detecting eye blinking[C]//2018 IEEE International Workshop on Information Forensics and Security. Hong Kong: IEEE, 2018: 1-7.",
        "[13] ZHOU P, HAN X, MORARIU V I, et al. Two-stream neural networks for tampered face detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition Workshops. Honolulu: IEEE, 2017: 1831-1839.",
        "[14] HE K, ZHANG X, REN S, et al. Deep residual learning for image recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 770-778.",
        "[15] GOODFELLOW I, POUGET-ABADIE J, MIRZA M, et al. Generative adversarial nets[C]//Advances in Neural Information Processing Systems. Montreal, 2014: 2672-2680.",
        "[16] RONNEBERGER O, FISCHER P, BROX T. U-Net: Convolutional networks for biomedical image segmentation[C]//Medical Image Computing and Computer-Assisted Intervention. Cham: Springer, 2015: 234-241.",
        "[17] ZHU J Y, PARK T, ISOLA P, et al. Unpaired image-to-image translation using cycle-consistent adversarial networks[C]//Proceedings of the IEEE International Conference on Computer Vision. Venice: IEEE, 2017: 2223-2232.",
        "[18] OpenCV Team. OpenCV Documentation[EB/OL]. https://docs.opencv.org/, 2026-04-01.",
        "[19] Django Software Foundation. Django Documentation[EB/OL]. https://docs.djangoproject.com/, 2026-04-01.",
        "[20] Riverbank Computing. PyQt5 Reference Guide[EB/OL]. https://www.riverbankcomputing.com/software/pyqt/, 2026-04-01.",
        "[21] Python Software Foundation. Python 3 Documentation[EB/OL]. https://docs.python.org/3/, 2026-04-01.",
        "[22] SQLite Consortium. SQLite Documentation[EB/OL]. https://www.sqlite.org/docs.html, 2026-04-01.",
        "[23] InsightFace. InsightFace: 2D and 3D Face Analysis Project[EB/OL]. https://github.com/deepinsight/insightface, 2026-04-01.",
    ]
    for ref in references:
        add_text_paragraph(doc, ref, size=Pt(10.5), first_line_chars=0, line_spacing=1)

    add_heading(doc, "致    谢", 1)
    add_text_paragraph(doc, "在本系统设计、实现与论文整理过程中，我得到了课程指导、项目资料与开源社区资源的帮助。感谢指导教师在选题方向、系统组织和论文写作规范方面给予的建议；感谢相关课程中关于软件工程、数据库、计算机视觉和界面开发的知识积累，为本课题的完成提供了基础；同时感谢开源社区提供的 PyQt5、Django、OpenCV 与 InsightFace 等技术资源。由于本人水平有限，文中难免存在不足，恳请各位老师批评指正。")

    add_heading(doc, "附录 A 主要接口清单", 1)
    add_text_paragraph(doc, "为便于后续系统维护与功能扩展，本文在附录中给出项目当前使用的主要接口清单。接口路径依据项目中的 Django 路由配置整理，既可供测试使用，也可为后续客户端扩展提供参考。")
    add_caption(doc, "表 A.1 接口清单")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "资源类型"
    table.rows[0].cells[1].text = "接口路径"
    table.rows[0].cells[2].text = "说明"
    for row in [
        ("人脸图片", "/api/images/", "列表查询、登记图片"),
        ("输入视频", "/api/videos/", "列表查询、登记视频"),
        ("输出视频", "/api/outputs/", "保存与获取输出结果"),
        ("处理任务", "/api/tasks/", "任务管理扩展接口"),
        ("接口文档", "/api/docs/", "Swagger 页面"),
    ]:
        cells = table.add_row().cells
        for idx, txt in enumerate(row):
            cells[idx].text = txt
    style_table(table)

    for section in doc.sections:
        format_section(section)
    doc.save(str(output_path))


if __name__ == "__main__":
    output_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT_PATH
    build_document(output_path)
    reopened = Document(str(output_path))
    for idx, para in enumerate(reopened.paragraphs[:12], 1):
        text = para.text.strip()
        if text:
            print(idx, text[:80])
    print(output_path)
