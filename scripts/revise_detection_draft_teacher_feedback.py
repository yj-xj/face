from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
DOCX_PATH = ROOT / "换脸系统_检测稿.docx"
BACKUP_PATH = ROOT / "换脸系统_检测稿_老师意见修改前备份.docx"
OUTPUT_PATH = ROOT / "换脸系统_检测稿_按老师意见修改版.docx"
ASSET_DIR = ROOT / "thesis_assets_generated"


REPLACEMENTS = {
    "本章不准备单独罗列技术名词，而是围绕系统真实运行所依赖的处理链来展开说明。内容先界定当前项目中的人脸替换任务，再分别讨论传统方案、InsightFace 主路径、OpenCV 媒体处理、PyQt5 交互机制，以及 Django REST Framework 与 SQLite 的数据管理职责，为后续需求分析和设计实现提供统一的技术背景。":
        "第二章围绕系统运行所依赖的技术链路展开，重点说明人脸替换任务的处理流程、传统 OpenCV 路径、InsightFace 主路径、媒体处理机制、桌面交互机制以及数据管理方式。通过对相关技术职责的梳理，为后续需求分析、总体设计与关键实现提供技术基础。",
    "本节围绕当前项目中的传统处理路径展开说明。前文会先交代 传统方法单帧处理、高级换脸处理 等入口如何组织传统换脸流程，再说明 Delaunay 三角剖分与局部仿射变换如何完成纹理迁移，最后讨论 无缝融合函数 与颜色校正等补偿步骤在边界处理中的作用，从而把传统方案在代码中的实际职责讲清楚。":
        "本节从代码入口、几何变换和边界补偿三个层面分析传统处理路径。系统通过传统单帧处理、高级换脸处理等入口组织 OpenCV 换脸流程，利用 Delaunay 三角剖分和局部仿射变换完成纹理迁移，并结合无缝融合与颜色校正降低边界突兀感。",
    "本节主要说明系统中的深度学习主路径是怎样在实际代码里落地的。下面会依次交代 dlib、InsightFace 和 inswapper_128 在工程链路中的分工，解释身份表征与目标属性保持的基本思路，并结合损失表达与推理流程说明该方案为什么会成为当前系统的主要处理路径。":
        "本节围绕深度学习主路径的工程组成展开，分析 dlib、InsightFace 与 inswapper_128 在处理链中的职责分工，说明身份表征、目标属性保持和推理流程之间的关系，并阐明该路径作为系统主要换脸方案的原因。",
    "本章的需求分析不是从抽象功能清单出发，而是放在真实使用过程里来讨论。论文将围绕视频换脸和摄像头演示两类核心场景，说明用户在素材选择、参数设置、处理启动、结果保存和历史回放等环节中的实际诉求，并进一步分析这些诉求对应的功能边界、非功能约束和实现前提。":
        "第三章以系统真实使用过程为分析对象，围绕视频换脸和摄像头演示两类核心场景，归纳用户在素材选择、参数设置、处理启动、结果保存和历史回放等环节中的需求，并进一步分析相关功能边界、非功能约束与实现条件。",
    "本章将系统总体设计与详细实现放在同一章展开。前半部分先从架构分层、模块职责、接口组织、数据库关系和部署方式说明整体设计，后半部分再结合前端界面、双模式交互、素材管理、处理流程与线程协作分析关键实现，以呈现系统从方案设计到落地运行的完整链路。":
        "第四章从总体设计与详细实现两个层面展开。前半部分围绕架构分层、模块职责、接口组织、数据库关系和部署方式说明系统方案，后半部分结合前端界面、双模式交互、素材管理、处理流程与线程协作分析关键实现，呈现系统从设计到运行的完整链路。",
    "本节在需求分析的基础上，对当前项目的整体结构做进一步的工程化整理。后文将依次说明设计目标、分层架构、模块职责、接口组织、数据库关系和本机部署形态，把前端、后端和脚本中的实现内容归纳为边界相对清楚、职责较为明确的系统方案。":
        "在需求分析基础上，本节对当前项目的整体结构进行工程化整理，依次说明设计目标、分层架构、模块职责、接口组织、数据库关系和本机部署形态，并将前端、后端和脚本中的实现内容归纳为边界清晰、职责明确的系统方案。",
    "这样的架构既能让系统在桌面端完成交互，也能把资源记录统一纳入后端管理。相比在前端直接操作文件系统，这种方式更利于维护素材与结果之间的对应关系，后续补充统计和查询能力时也会更方便。":
        "该架构既支持桌面端交互，也能够将资源记录统一纳入后端管理。相较于在前端直接操作文件系统，这种方式更有利于维护素材与结果之间的对应关系，并为后续统计和查询能力扩展提供基础。",
    "这样划分模块，主要是为了尽量降低耦合程度。虽然当前前端中仍保留了不少业务逻辑，但数据库交互和算法调用已经被拆分为相对独立的组件。这种组织方式更便于在不大幅调整界面的前提下替换后端实现，或切换不同的换脸方法。":
        "上述模块划分有助于降低系统耦合度。虽然当前前端仍承担部分业务控制逻辑，但数据库交互和算法调用已经拆分为相对独立的组件，因而便于在不大幅调整界面的前提下替换后端实现或切换换脸方法。",
    "这种部署结构的好处在于依赖关系清晰、搭建成本较低、调试周期也相对较短；相应地，不足之处是对本地路径依赖较强，多设备迁移和多用户共享都会比较困难。基于这一现实情况，论文在总结与展望中把资源路径整理、任务调度完善和部署方式调整列为后续重点。":
        "该部署结构具有依赖关系清晰、搭建成本较低和调试周期较短等优势；其局限在于对本地路径依赖较强，多设备迁移和多用户共享较为困难。基于这一现实情况，论文在总结与展望中将资源路径整理、任务调度完善和部署方式调整列为后续重点。",
    "本节关注的是系统已经落地的关键实现，而不是设计层面的抽象描述。下面将从前端主界面、双模式交互、素材加载、视频处理、摄像头处理、结果归档和线程协作几个方面展开，说明现有代码是如何把换脸能力组织成桌面原型的。":
        "4.2 节围绕系统已经落地的关键实现展开，重点分析前端主界面、双模式交互、素材加载、视频处理、摄像头处理、结果归档和线程协作等内容，说明现有代码如何将换脸能力组织为可运行的桌面应用原型。",
    "在具体交互实现中，系统通过 应用模式枚举 枚举类区分视频模式与摄像头模式，模式切换 方法则负责在两种状态之间切换可见控件、按钮样式以及预览区域的工作方式。这里的关键不在于界面表面上完成了切换，而在于切换后状态清理要正确、旧线程不能残留、播放器之间也不能发生冲突。因此，系统保留了单一主窗口，但把两类工作流程的控制逻辑明确区分开来。":
        "在具体交互实现中，系统通过应用模式枚举区分视频模式与摄像头模式，并由模式切换方法控制可见控件、按钮样式和预览区域状态。该实现的重点在于保证状态清理正确、旧线程及时释放以及播放器之间不发生资源冲突。因此，系统虽然采用单一主窗口，但对两类工作流程的控制逻辑进行了明确区分。",
    "本节从类与线程协作的角度，解释桌面端原型为什么能够在长耗时处理中保持基本响应。内容先说明界面控制类、处理线程和底层引擎各自承担的职责，再分析信号传递、状态清理和线程切换如何共同支撑视频模式与摄像头模式的连续运行。":
        "本节从类与线程协作角度分析桌面端原型在长耗时处理中保持基本响应的机制，重点说明界面控制类、处理线程和底层引擎的职责划分，并讨论信号传递、状态清理和线程切换对视频模式与摄像头模式连续运行的支撑作用。",
    "这个设计虽然细节较多，但从系统可用性的角度看，它会直接影响用户能否区分“摄像头未启动”“摄像头已启动但未开启换脸”和“正在取流并处理”这三种状态。":
        "上述设计直接影响用户对系统状态的判断，使其能够区分“摄像头未启动”“摄像头已启动但未开启换脸”和“正在取流并处理”三种运行状态。",
    "本章围绕当前版本的实际运行结果开展工程验证。测试重点不是和公开基准测试做数值对比，而是结合现有数据库记录、媒体文件、界面交互流程和输出结果，检查系统在素材加载、处理执行、结果回放和信息留痕等环节上的表现，并据此分析目前已经完成的部分和仍然存在的问题。":
        "第五章基于当前版本的实际运行结果开展工程验证。测试重点不在于与公开基准进行数值对比，而在于结合数据库记录、媒体文件、界面交互流程和输出结果，检查系统在素材加载、处理执行、结果回放和信息留痕等环节上的表现，并分析当前实现的完成情况与主要问题。",
    "本节先交代测试所面对的真实环境，再说明这一轮验证究竟想检查什么。下面会依次说明测试定位、运行环境组成和样本范围，使后续功能验证、稳定性观察与案例分析都建立在可追溯的前提之上。":
        "本节明确测试环境、验证目标与样本范围，使后续功能验证、稳定性观察与案例分析建立在可追溯的运行条件之上。",
    "本节中的“测试”并不是把系统放进统一 基准测试 做横向对比，而是检查当前项目在真实素材、真实路径和真实数据库记录条件下能否稳定运行。关注点主要集中在三方面：一是素材加载、处理启动和结果保存这条链路是否真正打通；二是 PyQt5 界面在视频模式、摄像头模式与结果回放之间切换时能否保持响应；三是数据库记录与本地文件是否一致，尤其是历史 本地路径字段 和输出结果之间是否存在可以解释的偏差。":
        "本文测试定位为工程可用性验证，重点检查当前项目在真实素材、真实路径和真实数据库记录条件下能否稳定运行。验证内容主要包括素材加载、处理启动和结果保存链路是否打通，PyQt5 界面在视频模式、摄像头模式与结果回放之间切换时能否保持响应，以及数据库记录与本地文件之间是否保持一致。",
    "测试方案直接对应当前项目的真实调用链，而不是脱离实现去随意拆分功能点。论文把 加载人脸图片、加载视频、开始处理、开启摄像头、结束处理 和 保存视频 等关键过程串联起来考察，重点观察素材是否顺利进入系统、线程是否正确启动、结果是否能够回放，以及处理完成后能否留下可查询的记录。":
        "测试方案以当前项目的真实调用链为依据，将加载人脸图片、加载视频、开始处理、开启摄像头、结束处理和保存视频等关键过程串联考察，重点观察素材进入系统、线程启动、结果回放和处理记录保存等环节是否能够连续完成。",
    "为避免案例分析只停留在统计表和文字判断上，论文从 image 目录中整理出代表性样例，分别展示真实照片源人脸和绘制源人脸两种条件下的输出结果。这样安排的目的，不是把这一部分包装成严格控制变量的视觉实验，而是借助真实样例说明系统在不同素材条件下的可用性、稳定性和视觉特征。":
        "典型案例分析选取 image 目录中的代表性样例，分别展示真实照片源人脸和绘制源人脸两种条件下的输出结果。该部分并不作为严格控制变量的视觉实验，而是借助真实样例说明系统在不同素材条件下的可用性、稳定性和视觉特征。",
    "本章在回顾全文工作的基础上，对这套人脸替换系统已经形成的工程特点与现实价值再作归纳。相比单纯强调算法效果，这里更关注系统目前如何支撑视频模式、摄像头模式、结果归档和历史回放，以及这些能力在毕业设计场景中的实际意义，同时也会明确指出后续仍需补足的工程环节。":
        "第六章在回顾全文工作的基础上，对人脸替换系统已经形成的工程特点与现实价值进行归纳。相较于单纯强调算法效果，本章更关注系统对视频模式、摄像头模式、结果归档和历史回放的支撑能力，并指出后续仍需补足的工程环节。",
}


TABLE_NOTE_REPLACEMENTS = {
    "表2-1 对比了 OpenCV 与 Delaunay 三角剖分传统路径、以及 InsightFace 与 inswapper_128 深度学习路径在实现依赖、效果表现和系统适配性上的差异。":
        "由表2-1可见，OpenCV 与 Delaunay 三角剖分传统路径、InsightFace 与 inswapper_128 深度学习路径在实现依赖、效果表现和系统适配性方面存在明显差异。",
    "表2-2 汇总了 OpenCV、dlib、InsightFace、inswapper_128、PyQt5、Django REST 与 SQLite 在系统中的具体职责。":
        "由表2-2可见，OpenCV、dlib、InsightFace、inswapper_128、PyQt5、Django REST 与 SQLite 分别承担媒体处理、人脸分析、界面交互和数据管理等职责。",
    "表3-1 汇总了素材加载、处理启动、结果保存、历史回放与状态提示等核心功能需求。":
        "由表3-1可见，系统功能需求集中在素材加载、处理启动、结果保存、历史回放与状态提示等核心环节。",
    "表3-2 汇总了界面可用性、处理稳定性、结构可维护性与扩展性等非功能需求。":
        "由表3-2可见，系统非功能需求主要包括界面可用性、处理稳定性、结构可维护性与扩展性。",
    "表3-3 汇总了技术基础、工程条件与实现风险，对系统可行性进行了整体评估。":
        "由表3-3可见，当前项目在技术基础和工程条件方面具备较高可行性，同时仍需关注模型依赖和路径失效等实现风险。",
    "表4-1 汇总了前端界面、数据库管理组件、换脸引擎与后端资源层的主要职责分工。":
        "由表4-1可见，前端界面、数据库管理组件、换脸引擎与后端资源层之间已经形成较为清晰的职责分工。",
    "表4-2 汇总了图片、视频、输出结果与任务等核心资源接口的访问方式和用途。":
        "由表4-2可见，图片、视频、输出结果与任务等核心资源均通过相对统一的接口路径进行访问和管理。",
    "表4-3 汇总了 人脸图片表、输入视频表、输出视频表 与 处理任务表 等数据表的关键字段和用途。":
        "由表4-3可见，人脸图片表、输入视频表、输出视频表与处理任务表分别承担素材记录、结果留痕和任务扩展等数据管理职责。",
    "表4-4 汇总了素材加载、后端交互、目录扫描与异步处理所对应的关键方法和职责。":
        "由表4-4可见，素材加载、后端交互、目录扫描与异步处理等功能已经被拆分为相对明确的方法职责。",
    "表4-5 汇总了输出结果从处理完成、元数据登记到列表展示和结果回放的主要环节。":
        "由表4-5可见，输出结果归档流程覆盖处理完成、元数据登记、列表展示和结果回放等主要环节。",
    "表5-1 汇总了测试所使用的硬件、软件与运行环境配置。":
        "由表5-1可见，测试环境采用本地单机部署方式，硬件、软件与运行条件均能够支撑原型系统验证。",
    "表5-2 汇总了当前系统的数据库记录、本地有效路径、输出文件数量以及视频统计信息。":
        "由表5-2可见，当前系统已经积累一定数量的数据库记录和本地媒体文件，同时也暴露出部分历史路径与输出登记不一致的问题。",
    "表5-3 汇总了功能完整性、界面稳定性、数据一致性和异常处理等测试项目的验收指标。":
        "由表5-3可见，测试方案围绕功能完整性、界面稳定性、数据一致性和异常处理等方面设置验收指标。",
    "表5-4 汇总了视频处理、摄像头演示、结果保存与历史回放等核心功能的验证结果。":
        "由表5-4可见，视频处理、摄像头演示、结果保存与历史回放等核心功能在当前版本中均已通过验证。",
    "表5-5 汇总了输入样本属性、处理耗时、回放表现与资源一致性等稳定性观察结果。":
        "由表5-5可见，当前系统在中短视频样本上的处理与回放表现较为稳定，但资源一致性仍需后续完善。",
}


IMAGE_REPLACEMENTS = {
    "word/media/image8.png": ASSET_DIR / "图4-4_系统部署结构图_纯中文版.png",
    "word/media/image11.png": ASSET_DIR / "图4-7_视频模式处理流程图_纯中文版.png",
    "word/media/image13.png": ASSET_DIR / "图4-9_摄像头模式处理流程图_纯中文版.png",
}


TABLE_TITLE_RE = re.compile(r"^表\d+-\d+\s+.+表$")
FIGURE_TITLE_RE = re.compile(r"^图\d+-\d+\s+")


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        italic = first_run.italic
    else:
        font_name = None
        font_size = None
        bold = None
        italic = None

    paragraph.text = new_text
    if paragraph.runs:
        run = paragraph.runs[0]
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        run.bold = bold
        run.italic = italic


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.findall(".//" + qn("w:t"))).strip()


def is_empty_paragraph(element) -> bool:
    return element.tag == qn("w:p") and not paragraph_text(element) and not element.findall(".//" + qn("w:drawing"))


def next_nonempty_body_child(body, start_index: int):
    children = list(body.iterchildren())
    for child in children[start_index + 1 :]:
        if child.tag == qn("w:p") and not paragraph_text(child) and not child.findall(".//" + qn("w:drawing")):
            continue
        return child
    return None


def previous_nonempty_body_child(body, start_index: int):
    children = list(body.iterchildren())
    for child in reversed(children[:start_index]):
        if child.tag == qn("w:p") and not paragraph_text(child) and not child.findall(".//" + qn("w:drawing")):
            continue
        return child
    return None


def move_table_titles_above(doc: Document) -> int:
    moved = 0
    body = doc.element.body

    for child in list(body.iterchildren()):
        if child.tag != qn("w:tbl"):
            continue

        children = list(body.iterchildren())
        index = children.index(child)
        prev_child = previous_nonempty_body_child(body, index)
        if prev_child is not None and prev_child.tag == qn("w:p") and TABLE_TITLE_RE.match(paragraph_text(prev_child)):
            continue

        next_child = next_nonempty_body_child(body, index)
        if next_child is None or next_child.tag != qn("w:p"):
            continue

        text = paragraph_text(next_child)
        if not TABLE_TITLE_RE.match(text):
            continue

        child.addprevious(next_child)
        moved += 1

    return moved


def remove_stray_punctuation(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "、":
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1
    return removed


def format_captions(doc: Document) -> tuple[int, int]:
    table_count = 0
    figure_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if TABLE_TITLE_RE.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            table_count += 1
        elif FIGURE_TITLE_RE.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)
            figure_count += 1
    return table_count, figure_count


def revise_text(doc: Document) -> tuple[int, int]:
    main_changes = 0
    table_notes = 0
    merged = {**REPLACEMENTS, **TABLE_NOTE_REPLACEMENTS}
    for paragraph in doc.paragraphs:
        old = paragraph.text.strip()
        new = merged.get(old)
        if not new:
            continue
        replace_paragraph_text(paragraph, new)
        if old in REPLACEMENTS:
            main_changes += 1
        else:
            table_notes += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return main_changes, table_notes


def replace_images(docx_path: Path) -> int:
    replacements = {}
    for target_name, asset_path in IMAGE_REPLACEMENTS.items():
        if asset_path.exists():
            replacements[target_name] = asset_path.read_bytes()

    if not replacements:
        return 0

    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename in replacements:
                data = replacements[item.filename]
            dst.writestr(item, data)

    temp_path.replace(docx_path)
    return len(replacements)


def validate(docx_path: Path) -> list[str]:
    issues: list[str] = []
    doc = Document(str(docx_path))
    body = doc.element.body
    children = list(body.iterchildren())

    for index, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        # The cover metadata table is not a body table and does not need a
        # thesis-style "表x-x" title.
        if index < 40:
            continue
        prev_child = previous_nonempty_body_child(body, index)
        if prev_child is None or prev_child.tag != qn("w:p") or not TABLE_TITLE_RE.match(paragraph_text(prev_child)):
            issues.append(f"table_without_title_above_at_body_{index}")

        next_child = next_nonempty_body_child(body, index)
        if next_child is not None and next_child.tag == qn("w:p") and TABLE_TITLE_RE.match(paragraph_text(next_child)):
            issues.append(f"table_title_still_below_at_body_{index}")

    figure_titles = [p.text.strip() for p in doc.paragraphs if FIGURE_TITLE_RE.match(p.text.strip())]
    table_titles = [p.text.strip() for p in doc.paragraphs if TABLE_TITLE_RE.match(p.text.strip())]
    # The document has 17 numbered body figures; the extra embedded image is
    # the school logo on the cover.
    if len(figure_titles) != 17:
        issues.append(f"figure_title_count={len(figure_titles)}")
    if len(table_titles) != 15:
        issues.append(f"table_title_count={len(table_titles)}")

    residual_patterns = ["本章不准备", "下面会", "前文会", "后文将", "这里的关键", "这样安排的目的", "究竟想检查什么"]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for pattern in residual_patterns:
        if pattern in full_text:
            issues.append(f"residual_phrase={pattern}")

    with zipfile.ZipFile(docx_path, "r") as zf:
        for target_name in IMAGE_REPLACEMENTS:
            try:
                size = len(zf.read(target_name))
            except KeyError:
                issues.append(f"missing_image={target_name}")
                continue
            if size < 500_000:
                issues.append(f"image_not_replaced_or_small={target_name}:{size}")

    return issues


def find_docx() -> Path:
    if DOCX_PATH.exists():
        return DOCX_PATH
    matches = [p for p in ROOT.glob("*检测稿.docx") if not p.name.startswith("~$")]
    if matches:
        return matches[0]
    matches = [p for p in ROOT.glob("*.docx") if p.stat().st_size == 6578555 and not p.name.startswith("~$")]
    if matches:
        return matches[0]
    raise FileNotFoundError("未找到检测稿 docx 文件。")


def main() -> None:
    source = find_docx()
    if not BACKUP_PATH.exists():
        shutil.copy2(source, BACKUP_PATH)

    shutil.copy2(source, OUTPUT_PATH)
    doc = Document(str(OUTPUT_PATH))

    moved_tables = move_table_titles_above(doc)
    removed = remove_stray_punctuation(doc)
    main_changes, table_notes = revise_text(doc)
    table_caption_count, figure_caption_count = format_captions(doc)

    doc.save(str(OUTPUT_PATH))
    replaced_images = replace_images(OUTPUT_PATH)

    issues = validate(OUTPUT_PATH)

    print(f"SOURCE={source}")
    print(f"BACKUP={BACKUP_PATH}")
    print(f"OUTPUT={OUTPUT_PATH}")
    print(f"MOVED_TABLE_TITLES={moved_tables}")
    print(f"REMOVED_STRAY_PARAGRAPHS={removed}")
    print(f"MAIN_TEXT_REWRITES={main_changes}")
    print(f"TABLE_NOTE_REWRITES={table_notes}")
    print(f"TABLE_CAPTIONS={table_caption_count}")
    print(f"FIGURE_CAPTIONS={figure_caption_count}")
    print(f"REPLACED_IMAGES={replaced_images}")
    print(f"VALIDATION_ISSUES={len(issues)}")
    for issue in issues:
        print(f"ISSUE={issue}")


if __name__ == "__main__":
    main()
