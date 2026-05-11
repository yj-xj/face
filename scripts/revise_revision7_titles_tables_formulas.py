from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
NUMBER_RE = re.compile(r"^（\d+(?:-\d+)?）$")
CHAPTER_HEADING_RE = re.compile(r"^(第[一二三四五六七八九十]+章)\s*(.+)$")


def find_source() -> Path:
    matches = [p for p in ROOT.glob("*修订版7.docx") if not p.name.startswith("~$")]
    if not matches:
        raise FileNotFoundError("Could not find the non-temp revision 7 docx.")
    return matches[0]


def build_output_path(source: Path) -> Path:
    match = re.search(r"修订版(\d+)$", source.stem)
    if match:
        next_rev = int(match.group(1)) + 1
        prefix = source.stem[: match.start()]
        candidate = source.with_name(prefix + f"修订版{next_rev}" + source.suffix)
        while candidate.exists():
            next_rev += 1
            candidate = source.with_name(prefix + f"修订版{next_rev}" + source.suffix)
        return candidate
    candidate = source.with_name(source.stem + "_表格公式修订版" + source.suffix)
    counter = 1
    while candidate.exists():
        candidate = source.with_name(source.stem + f"_表格公式修订版{counter}" + source.suffix)
        counter += 1
    return candidate


def paragraph_has_math(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))


def normalize_chapter_heading_text(text: str) -> str:
    stripped = (text or "").strip()
    match = CHAPTER_HEADING_RE.match(stripped)
    if not match:
        return stripped
    return f"{match.group(1)} {match.group(2).strip()}"


def update_chapter_headings(doc: DocumentObject) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name != "Heading 1":
            continue
        new_text = normalize_chapter_heading_text(paragraph.text)
        old_text = (paragraph.text or "").strip()
        if new_text and new_text != old_text:
            paragraph.text = new_text
            changed += 1
    return changed


def set_run_font(run, *, east_asia: str, ascii_font: str, size_pt: float, bold: bool) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size_pt)
    run.bold = bold


def format_formula_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def set_formula_right_tab(paragraph: Paragraph, usable_width: Emu) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    existing_tabs = p_pr.find(qn("w:tabs"))
    if existing_tabs is not None:
        p_pr.remove(existing_tabs)

    tabs = OxmlElement("w:tabs")
    right_tab = OxmlElement("w:tab")
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), str(int(usable_width)))
    tabs.append(right_tab)
    p_pr.append(tabs)


def clear_formula_number_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)


def renumber_formulas_by_chapter(doc: DocumentObject) -> int:
    usable_width = Emu(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
    chapter_no = 0
    chapter_formula_count = 0
    changed = 0

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        text = (paragraph.text or "").strip()

        if style_name == "Heading 1" and CHAPTER_HEADING_RE.match(text):
            chapter_no += 1
            chapter_formula_count = 0
            continue

        if chapter_no == 0 or not paragraph_has_math(paragraph):
            continue

        if text and not NUMBER_RE.fullmatch(text):
            continue

        chapter_formula_count += 1
        clear_formula_number_runs(paragraph)
        format_formula_paragraph(paragraph)
        set_formula_right_tab(paragraph, usable_width)
        run = paragraph.add_run(f"\t（{chapter_no}-{chapter_formula_count}）")
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=12.0, bold=False)
        changed += 1

    return changed


def ensure_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": 50, "bottom": 50, "left": 90, "right": 90}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clear_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is not None:
        tc_pr.remove(shd)


def set_border(edge_parent, edge: str, *, val: str, size: int = 0, color: str = "000000") -> None:
    border = edge_parent.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        edge_parent.append(border)
    border.set(qn("w:val"), val)
    if val not in {"nil", "none"}:
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
    else:
        border.attrib.pop(qn("w:sz"), None)
        border.attrib.pop(qn("w:space"), None)
        border.attrib.pop(qn("w:color"), None)


def set_table_three_line_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    set_border(tbl_borders, "top", val="single", size=12)
    set_border(tbl_borders, "bottom", val="single", size=12)
    set_border(tbl_borders, "left", val="nil")
    set_border(tbl_borders, "right", val="nil")
    set_border(tbl_borders, "insideH", val="nil")
    set_border(tbl_borders, "insideV", val="nil")


def clear_table_style_formatting(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblStyle", "w:tblLook"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cnf_style = tr_pr.find(qn("w:cnfStyle"))
        if cnf_style is not None:
            tr_pr.remove(cnf_style)

        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cnf_style = tc_pr.find(qn("w:cnfStyle"))
            if cnf_style is not None:
                tc_pr.remove(cnf_style)


def set_cell_borders(cell, *, top: str = "nil", bottom: str = "nil", left: str = "nil", right: str = "nil") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    set_border(tc_borders, "top", val=top)
    set_border(tc_borders, "bottom", val=bottom if bottom in {"nil", "none"} else "single", size=8)
    set_border(tc_borders, "left", val=left)
    set_border(tc_borders, "right", val=right)


def set_table_layout_fixed(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_width(table: Table, width_emu: Emu) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_emu / 635)))


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def format_table_cell_paragraph(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run.font.size = Pt(10.5)
        run.bold = False


def is_cover_info_table(table: Table) -> bool:
    if not table.rows or len(table.columns) != 2:
        return False
    header = [cell.text.strip() for cell in table.rows[0].cells]
    return header == ["姓    名", "叶俊"]


def format_body_tables(doc: DocumentObject) -> int:
    usable_width = Emu(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
    changed = 0

    for table in doc.tables:
        if is_cover_info_table(table):
            continue

        clear_table_style_formatting(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_layout_fixed(table)
        set_table_width(table, Emu(int(usable_width * 0.90)))
        set_table_three_line_borders(table)

        if table.rows:
            mark_header_row(table.rows[0])
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            table.rows[0].height = Pt(18)

        for row_index, row in enumerate(table.rows):
            is_header = row_index == 0
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(18)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                ensure_cell_margins(cell)
                clear_cell_shading(cell)
                set_cell_borders(cell, bottom="single" if is_header else "nil")
                for paragraph in cell.paragraphs:
                    format_table_cell_paragraph(paragraph)
        changed += 1

    return changed


def verify_formula_numbers(doc: DocumentObject) -> list[str]:
    found: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if NUMBER_RE.fullmatch(text):
            found.append(text)
    return found


def main() -> None:
    source = find_source()
    output = build_output_path(source)
    shutil.copy2(source, output)

    doc = Document(str(output))
    heading_changes = update_chapter_headings(doc)
    table_changes = format_body_tables(doc)
    formula_changes = renumber_formulas_by_chapter(doc)
    doc.save(str(output))

    final_doc = Document(str(output))
    formula_numbers = verify_formula_numbers(final_doc)

    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    print(f"HEADING_CHANGES={heading_changes}")
    print(f"TABLES_FORMATTED={table_changes}")
    print(f"FORMULAS_RENUMBERED={formula_changes}")
    print(f"FORMULA_NUMBERS={formula_numbers}")


if __name__ == "__main__":
    main()
