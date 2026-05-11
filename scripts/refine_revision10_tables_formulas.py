from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"E:\face")
SOURCE_PATTERN = "*修订版10.docx"


PARAGRAPH_REWRITES = {
    67: "需要说明的是，传统路径虽然在检测阶段允许用户在界面中选择 dlib 或 OpenCV 级联检测器，但真正决定换脸效果的主体仍然是 OpenCV 图像处理流程本身。代码中无论使用哪一种检测入口，最终都要借助关键点集合构造三角网格，并通过逐三角形重映射完成纹理迁移。因此，从论文论述角度看，传统方法更应当围绕 OpenCV 的几何处理能力和 Delaunay 三角剖分机制展开，而不是围绕深度学习模型展开。若将源三角形内部点坐标记为 (x_s, y_s)，目标三角形中的对应点坐标记为 (x_t, y_t)，则局部仿射映射可写为式（2-1）和式（2-2）。",
    73: "从具体算法机制看，三角剖分的意义在于避免对整张人脸执行单一刚性变换，而是将人脸区域离散为多个局部三角单元，再分别计算对应的仿射映射关系。OpenCV 提供的 Subdiv2D、getAffineTransform、warpAffine 和掩模运算，使系统能够在逐三角形变换之后继续完成边界平滑与颜色过渡。若将全部匹配三角形构成的集合记为 T，则其可表示为式（2-3）。",
    76: "式（2-3）中的 τ_k 表示第 k 个匹配三角单元。若进一步将目标人脸关键点集合与源人脸关键点集合分别记为式（2-4）和式（2-5），则每个三角单元都可以在两组关键点之间建立一一对应关系。",
    81: "在工程实现上，系统会依据上述三角对应关系完成局部区域裁剪、仿射变形与结果叠加。式（2-4）和式（2-5）分别给出了目标侧与源侧的关键点采样结果，它们共同为后续局部纹理映射提供了几何约束基础。",
    82: "进一步地，每个三角单元都会计算一组局部仿射参数，用于把源三角形纹理稳定映射到目标三角形区域。",
    85: "式（2-6）给出了无缝融合时常见的梯度约束目标，用于减弱拼接边界处的梯度突变。若采用更直接的遮罩混合方式，则输出像素还可以写为式（2-7）。",
    89: "在边界补偿之前，系统还需要先利用重心坐标完成目标三角形内部像素到源三角形纹理坐标的回填。若将源三角形中的对应纹理位置记为 q，则其可表示为式（2-8）。",
    92: "其中，λ_1、λ_2、λ_3 表示同一组三角形重心坐标，它们需要满足式（2-9）的约束，以保证插值位置仍然位于三角形内部。",
    97: "当目标三角形内部像素位置 p 用式（2-10）表示时，系统便可借助同一组重心坐标把目标侧像素稳定映射回源侧纹理位置，再结合 seamlessClone 与颜色校正完成边界补偿。",
    99: "表2-1 传统方法与深度学习方法对比表",
    100: "表2-1 对比了 OpenCV 与 Delaunay 三角剖分传统路径、以及 InsightFace 与 inswapper_128 深度学习路径在实现依赖、效果表现和系统适配性上的差异。",
    105: "在 face_swap.py 中，系统首先加载 shape_predictor_68_face_landmarks.dat 与 inswapper_128.onnx 等模型资源，并在界面层提供 dlib / OpenCV 检测器和 traditional / inswapper 换脸方法的可切换配置。就深度学习主路径而言，insightface_face_swap 会先调用 self.face_analyser.get 分别分析目标帧和源人脸图像，获得人脸框、关键点与身份相关特征；随后再调用 self.inswapper.get(result, face, source_face, paste_back=True) 完成逐脸替换，并在需要时执行颜色校正。这说明本文讨论的深度学习方法，核心应当围绕 InsightFace 分析流程与 inswapper_128 推理模型展开，同时结合 dlib 在系统检测配置中的辅助作用进行理解。若将源身份编码记为 e_s，生成结果记为 I_hat，则深度学习主路径可进一步概括为式（2-11）和式（2-12）。",
    111: "基于式（2-11）和式（2-12），深度学习的人脸替换可以进一步抽象为身份特征提取与目标属性保持的联合映射过程。设源人脸图像为 I_s，目标图像或目标帧为 I_t，则输出结果既需要保留源身份特征，也需要尽量保持目标帧中的姿态、表情和光照条件。",
    112: "在这一处理框架中，e_s 主要承担身份信息约束，而 I_t 中的姿态、表情、光照与背景结构则作为目标属性被尽量保留。因此，深度学习方法的关键不再只是局部几何对齐，而是如何在高维特征空间内同时协调身份一致性与目标场景一致性。",
    115: "进一步地，为同时兼顾身份一致性、内容重建、感知相似性与结果真实感，常用的综合优化目标可进一步表示为",
    119: "式（2-14）中的各项损失分别对应身份保持、重建约束、感知一致性与结果真实感，它们共同决定了深度学习换脸结果的稳定性与自然度。",
    122: "表2-2 汇总了 OpenCV、dlib、InsightFace、inswapper_128、PyQt5、Django REST 与 SQLite 在系统中的具体职责。",
    144: "表3-1 汇总了素材加载、处理启动、结果保存、历史回放与状态提示等核心功能需求。",
    149: "表3-2 汇总了界面可用性、处理稳定性、结构可维护性与扩展性等非功能需求。",
    160: "表3-3 汇总了技术基础、工程条件与实现风险，对系统可行性进行了整体评估。",
    178: "表4-1 汇总了前端界面、数据库管理组件、换脸引擎与后端资源层的主要职责分工。",
    186: "表4-2 汇总了图片、视频、输出结果与任务等核心资源接口的访问方式和用途。",
    191: "表4-3 汇总了 FaceImage、InputVideo、OutputVideo 与 ProcessingTask 等数据表的关键字段和用途。",
    222: "表4-4 汇总了素材加载、后端交互、目录扫描与异步处理所对应的关键方法和职责。",
    248: "表4-5 汇总了输出结果从处理完成、元数据登记到列表展示和结果回放的主要环节。",
    257: "从静态结构角度看，界面控制类、处理线程、数据管理组件与底层引擎之间已经形成较清晰的职责分工，这与前面的协作时序说明相互对应。",
    269: "表5-1 汇总了测试所使用的硬件、软件与运行环境配置。",
    274: "表5-2 汇总了当前系统的数据库记录、本地有效路径、输出文件数量以及视频统计信息。",
    279: "表5-3 汇总了功能完整性、界面稳定性、数据一致性和异常处理等测试项目的验收指标。",
    283: "表5-4 汇总了视频处理、摄像头演示、结果保存与历史回放等核心功能的验证结果。",
    289: "表5-5 汇总了输入样本属性、处理耗时、回放表现与资源一致性等稳定性观察结果。",
}

INSERT_AFTER_212_TEXT = (
    "式（2-11）说明输出结果由目标帧 I_t 与源身份编码 e_s 共同决定，"
    "式（2-12）则说明该身份编码由源人脸 I_s 经过身份特征提取网络 F_id 获得。"
    "二者共同对应代码中 FaceAnalysis 提取特征、inswapper_128 完成身份迁移的处理关系。"
)


def find_source() -> Path:
    matches = sorted(p for p in ROOT.glob(SOURCE_PATTERN) if not p.name.startswith("~$"))
    if not matches:
        raise FileNotFoundError("未找到修订版10文档。")
    return matches[0]


def build_output_path(source: Path) -> Path:
    match = re.search(r"修订版(\d+)$", source.stem)
    if match:
        next_rev = int(match.group(1)) + 1
        prefix = source.stem[: match.start()]
        candidate = source.with_name(f"{prefix}修订版{next_rev}{source.suffix}")
        while candidate.exists():
            next_rev += 1
            candidate = source.with_name(f"{prefix}修订版{next_rev}{source.suffix}")
        return candidate
    candidate = source.with_name(source.stem + "_修订版11" + source.suffix)
    counter = 12
    while candidate.exists():
        candidate = source.with_name(source.stem + f"_修订版{counter}" + source.suffix)
        counter += 1
    return candidate


def copy_run_format(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb

    src_rpr = src_run._element.rPr
    if src_rpr is None:
        return
    src_fonts = src_rpr.find(qn("w:rFonts"))
    if src_fonts is None:
        return

    dst_rpr = dst_run._element.get_or_add_rPr()
    dst_fonts = dst_rpr.find(qn("w:rFonts"))
    if dst_fonts is None:
        dst_fonts = OxmlElement("w:rFonts")
        dst_rpr.append(dst_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        value = src_fonts.get(qn(key))
        if value:
            dst_fonts.set(qn(key), value)


def copy_paragraph_format(src: Paragraph, dst: Paragraph) -> None:
    dst.style = src.style
    dst.alignment = src.alignment
    src_pf = src.paragraph_format
    dst_pf = dst.paragraph_format

    attrs = (
        "left_indent",
        "right_indent",
        "first_line_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacing_rule",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
    )
    for attr in attrs:
        try:
            setattr(dst_pf, attr, getattr(src_pf, attr))
        except Exception:
            continue


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    template_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.text = new_text
    if template_run and paragraph.runs:
        copy_run_format(template_run, paragraph.runs[0])


def insert_paragraph_after(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.add_run(text)
    copy_paragraph_format(template, paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = template.paragraph_format.first_line_indent or Pt(24)
    if template.runs and paragraph.runs:
        copy_run_format(template.runs[0], paragraph.runs[0])
    return paragraph


def verify(doc: DocumentObject) -> tuple[list[str], list[str]]:
    residual_table_style = []
    for idx in [100, 122, 144, 149, 160, 178, 186, 191, 222, 248, 269, 274, 279, 283, 289]:
        text = (doc.paragraphs[idx].text or "").strip()
        if "表 " in text and "." in text:
            residual_table_style.append(f"{idx}:{text}")

    formula_checks = []
    for idx in [67, 73, 76, 81, 82, 85, 89, 92, 97, 105, 111, 112, 115, 119]:
        text = (doc.paragraphs[idx].text or "").strip()
        if not text:
            formula_checks.append(f"{idx}:EMPTY")

    return residual_table_style, formula_checks


def main() -> None:
    source = find_source()
    output = build_output_path(source)
    shutil.copy2(source, output)

    doc = Document(str(output))

    for index, new_text in PARAGRAPH_REWRITES.items():
        replace_paragraph_text(doc.paragraphs[index], new_text)

    anchor = doc.paragraphs[109]
    template = doc.paragraphs[111]
    insert_paragraph_after(anchor, INSERT_AFTER_212_TEXT, template)

    doc.save(str(output))

    checked = Document(str(output))
    residual_table_style, formula_checks = verify(checked)

    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    print(f"PARAGRAPHS_UPDATED={len(PARAGRAPH_REWRITES)}")
    print("INSERTED_AFTER_2_12=1")
    print(f"RESIDUAL_TABLE_STYLE_COUNT={len(residual_table_style)}")
    for item in residual_table_style:
        print(f"RESIDUAL_TABLE_STYLE={item}")
    print(f"FORMULA_CHECK_ISSUE_COUNT={len(formula_checks)}")
    for item in formula_checks:
        print(f"FORMULA_CHECK_ISSUE={item}")


if __name__ == "__main__":
    main()
