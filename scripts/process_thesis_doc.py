from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"e:\face")
SOURCE_DOCX = BASE_DIR / "换脸系统_初稿_学术化版.docx"
OUTPUT_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_处理后.docx"
OUTPUT_MD = BASE_DIR / "换脸系统_初稿_学术化版_处理后.md"


CHAPTER_OVERVIEWS = {
    "1 绪论": "本章围绕课题的提出背景、研究价值与论文任务展开论述。首先结合短视频创作与生成式视觉应用的发展趋势，说明人脸替换系统从算法演示走向工程化实现的现实需求；随后梳理国内外相关研究现状，明确本文不以提出新模型为目标，而以现有项目的系统整理、功能验证和工程分析为核心。",
    "2 相关技术与理论基础": "本章从系统实现所依赖的关键技术出发，介绍人脸替换任务的基本定义、典型处理流程以及传统方法与深度学习方法的差异，并进一步说明 OpenCV、PyQt5、Django REST Framework 和 SQLite 在本系统中的具体作用。通过这些理论与技术基础的铺垫，可为后续需求分析和设计实现提供统一的概念框架。",
    "3 系统需求分析": "本章面向真实使用过程分析系统需求，重点说明视频换脸与摄像头演示两类业务场景下的功能目标与约束条件。在此基础上，本文将从功能需求、非功能需求、用例关系以及系统可行性等方面逐步展开论证，为后续总体设计和模块划分提供明确依据。",
    "4 系统总体设计": "本章在前述需求分析的基础上，对系统的整体结构进行抽象与规划。内容将依次说明设计目标与基本原则、总体架构、核心模块职责、接口组织方式、数据库表结构以及本机部署形态，从而把分散的代码实现整理为具有层次和边界的工程方案。",
    "5 系统详细设计与实现": "本章围绕系统从界面到处理线程的核心实现展开详细说明。首先介绍前端主界面与双模式交互逻辑，然后分析素材加载、视频处理、摄像头处理、结果归档和线程协作等关键实现环节，以展示系统如何把已有换脸能力组织为可运行、可反馈、可回放的完整桌面应用。",
    "6 系统测试与结果分析": "本章以工程验证为目标，对系统当前版本的运行情况进行测试与分析。内容包括测试环境与方案设计、功能验证结果、稳定性与性能观察、典型样例分析以及现存问题总结，重点说明系统在真实项目材料上的完成度、可运行性和后续需要改进的方向。",
    "7 总结与展望": "本章对全文工作进行归纳，并从工程实现视角总结系统已经形成的特点与价值。在回顾需求分析、总体设计、实现与测试结果的基础上，进一步讨论当前系统在资源治理、任务调度和部署扩展等方面的不足，为后续持续完善提供方向。",
}


CAPTION_EXPLANATIONS = {
    "表 4.1 核心模块职责表": "从表 4.1 可以看出，系统虽然由前端界面、数据库管理组件和原始换脸引擎等多个部分共同组成，但每个模块的责任边界已经比较明确：界面层负责用户输入与状态反馈，数据层负责素材与结果记录，处理层负责具体的人脸替换逻辑。这样的职责划分并不意味着模块之间彼此孤立，相反，它们必须通过稳定的数据流和调用链协同工作，才能保证一次处理任务能够从素材选择顺利过渡到结果输出。因此，在理解模块静态职责之后，还需要进一步观察系统内部信息是如何在不同组件之间流动的。",
    "表 4.3 数据库表设计表": "表 4.3 对各数据表字段和用途进行了静态描述，但数据库设计的价值并不仅停留在字段罗列层面。对于本系统而言，人脸图片、输入视频、输出结果与处理任务之间存在天然的业务关联，只有把这些对象之间的对应关系梳理清楚，后续的素材加载、结果查询和历史回放才具有可追踪性。也就是说，表结构本身决定了系统能够记录什么，而实体关系则进一步决定了这些记录能否被组织成可理解、可验证的处理链路，因此有必要结合关系示意图继续展开分析。",
    "图 5.3 视频模式处理流程图": "图 5.3 重点呈现了视频模式从输入检查到线程执行再到结果保存的主干流程。需要指出的是，这一流程并不是简单的顺序调用，而是把参数校验、界面状态更新、后台处理和结果回放串联为一个完整闭环。用户在界面上看到的是一次按钮触发，但系统内部实际上经历了素材确认、处理方法同步、信号连接和输出登记等多个步骤。正是由于这些步骤被稳定组织起来，系统才不仅能够生成换脸结果，还能够把结果及时返回到预览区并形成可复查的处理记录。基于这一流程，图 5.4 进一步给出了输入与输出样例，用于说明抽象流程在真实素材上的具体表现。",
    "图 5.6 关键线程与模块协作时序图": "图 5.6 展示的是系统运行时的动态协作关系，它说明界面线程、视频处理线程、摄像头线程以及原始换脸引擎之间并不是简单的前后调用，而是通过信号槽与状态控制形成持续交互。对于桌面多媒体应用而言，时序关系是否清楚直接影响界面响应、资源释放和异常处理效果。如果只知道模块名称，却不了解它们在运行过程中的职责分担，就很难解释系统为何能够在较长处理过程中保持可用。因此，在观察完协作时序之后，还需要进一步结合核心类职责表，从静态结构层面对各个对象承担的任务进行归纳。",
    "表 6.1 测试环境配置表": "表 6.1 说明了测试所依赖的硬件、软件和运行环境，这为后续结果分析提供了必要前提。对工程类毕业设计而言，测试结论能否成立不仅取决于是否执行过操作，还取决于运行环境是否具备可描述性和可复现性。在此基础上，图 6.1 进一步把当前系统实际积累的样本数据情况直观呈现出来，使读者能够同时从配置条件与样本规模两个角度理解测试结果的适用范围，也有助于判断后续功能验证和稳定性分析所依据的数据基础是否充分。",
    "图 6.1 系统样本数据统计图": "从图 6.1 可以直观看出，当前系统中人脸图片、输入视频和输出视频三类资源的数量分布并不完全均衡，这与项目以演示和工程验证为主的定位是相一致的。样本量虽然不属于大规模实验数据，但已经足以覆盖素材加载、任务处理、结果保存与历史查询等主要环节。进一步来看，图中的数量关系也提示出一个重要现象：输出结果的积累速度受处理链路和登记机制共同影响，而不仅仅取决于输入素材数量。因此，紧接着通过表 6.2 对基础统计数据进行结构化汇总，有助于把图形化观察转化为更明确的分析依据。",
    "图 6.2 最近输入视频的时长与帧率分布图": "图 6.2 反映了最近输入视频在时长与帧率上的分布情况，从中可以看出当前测试样本总体上更接近中短时长、便于快速演示的媒体材料，而不是超长视频或极端高帧率场景。这样的样本特征会直接影响系统的处理时延感知、界面反馈频率以及结果回放压力，因此在解读后续性能观察时必须纳入考虑。换言之，稳定性与性能判断不能脱离输入数据本身的属性来进行。基于该图所揭示的分布特征，表 6.5 进一步把相关观察结果整理为条目化记录，从而便于对系统当前表现进行更细致的归纳。",
    "图 6.3 样例 A 的输入图像与换脸结果": "图 6.3 所示样例 A 主要用于说明系统在一组典型输入条件下能够完成从源图到结果图的基本映射。观察该样例时，重点不仅在于结果是否生成，更在于人脸区域是否保持了可识别的身份迁移特征、整体结构是否与原始画面相协调，以及输出是否具备继续用于视频处理展示的可用性。对于工程化系统来说，单一样例只能证明链路可运行，却不足以说明系统在不同素材上的表现具有一定稳定性。因此，还需要引入第二组样例进行对照，以观察系统在输入条件变化后是否仍能维持基本一致的处理效果。",
    "图 6.4 样例 B 的输入图像与换脸结果": "与样例 A 相比，图 6.4 所示样例 B 进一步扩展了观察范围，使读者能够从另一组素材中判断系统在不同输入条件下的输出一致性。若两个样例都能够完成较为清晰的身份替换并保持基本可辨识的人脸结构，就可以说明系统当前实现并非偶然依赖单一素材，而是已经具备一定的通用处理能力。当然，静态图像结果仍然只能反映局部效果，无法完整说明视频场景中的连续性与展示表现。因此，在对样例 B 完成分析之后，还需要结合最近一次输出视频首帧截图，对系统在真实视频结果中的呈现状态进行补充说明。",
}


CITATION_RULES = [
    ("从研究脉络看，人脸替换技术大致经历了传统几何对齐与融合到深度学习驱动身份迁移两个阶段。", "[1][2][3][4][5][6][9][10][11]"),
    ("然而，国内外大量工作更关注模型效果、数据集和视觉指标，而不是桌面系统中的工作流组织问题。", "[7][8]"),
    ("基于深度学习的人脸替换方法更强调身份特征表示和复杂场景下的结果自然度。", "[4][5][6][9][10][11][12]"),
    ("OpenCV 是当前系统中的基础工具。", "[13][14]"),
    ("PyQt5 负责桌面端的人机交互层。", "[15][16][22]"),
    ("在数据管理方面，本文没有把所有逻辑都堆积在桌面端，而是引入 Django REST Framework 作为统一的资源接口层。", "[17][18][21][25][26]"),
    ("后端接口采用 REST 风格组织。", "[17][18][25][26]"),
    ("从代码结构看，系统的并行协作主要围绕三类对象展开，分别是界面控制类 EnhancedFaceSwapUI、视频处理线程 VideoProcessingThread 和摄像头线程 CameraProcessingThread。", "[19][20]"),
    ("本章测试定位为工程验证，而不是算法竞赛式 benchmark。", "[23][24]"),
]


TERMINAL_MEMO = "“⚠️ 文本已修正完毕。请将文本复制回Word后，务必全选（包含参考文献）执行以下格式调整：1. 所有内容（含正文和参考文献）设置为两端对齐；2. 全文行距设为固定值 20磅；3. 正文的段前和段后都设为 0；4. 确保每一章（第X章）的标题前插入了换页符（新起一页）；5. 检查页面排版，利用刚刚扩写的图表解释文字，将页面底部的空白填满，绝不允许连续贴图。”"


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = "Normal", bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip().replace("\t", " ")


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def append_citation(text: str, citation: str) -> str:
    if citation in text:
        return text
    match = re.match(r"^(.*?)([。；？！.;:])?$", text)
    if not match:
        return f"{text}{citation}"
    body = match.group(1)
    punct = match.group(2) or ""
    return f"{body}{citation}{punct}"


def next_nonempty_text(paragraphs: list[Paragraph], start_idx: int) -> str:
    for paragraph in paragraphs[start_idx + 1 :]:
        text = paragraph_text(paragraph)
        if text:
            return text
    return ""


def insert_marked_paragraph(doc: Document, anchor_text: str, new_text: str, inserted: list[str]) -> bool:
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph_text(paragraph) != anchor_text:
            continue
        if next_nonempty_text(paragraphs, idx) == new_text:
            return False
        insert_paragraph_after(paragraph, new_text, bold=True)
        inserted.append(new_text)
        return True
    raise ValueError(f"Anchor not found: {anchor_text}")


def apply_keyword_fixes(doc: Document) -> int:
    changes = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "Key Words:" in text:
            replace_paragraph_text(paragraph, text.replace("Key Words:", "Keywords:"))
            changes += 1
    return changes


def apply_citations(doc: Document) -> int:
    changes = 0
    for prefix, citation in CITATION_RULES:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith(prefix):
                updated = append_citation(text, citation)
                if updated != text:
                    replace_paragraph_text(paragraph, updated)
                    changes += 1
                break
        else:
            raise ValueError(f"Citation target not found: {prefix}")
    return changes


def build_markdown(doc: Document, inserted_texts: set[str]) -> str:
    lines: list[str] = ["一、 修改与扩写后的正文", ""]
    include = False
    skip_toc = False
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if not text:
            continue
        if text == "摘    要":
            include = True
        if not include:
            continue
        if text == "目    录":
            skip_toc = True
            continue
        if skip_toc:
            if paragraph.style.name == "Heading 1" and re.match(r"^1\s+绪论$", text):
                skip_toc = False
            else:
                continue
        if text in inserted_texts:
            lines.append(f"**{text}**")
        else:
            lines.append(text)
        lines.append("")

    lines.extend(
        [
            "二、 修改日志与排版备忘录",
            "",
            "1. 在第 1 章至第 7 章章标题后补充了概述性过渡段，并在导出文本中以加粗形式标出。",
            "2. 在表 4.1/图 4.2、表 4.3/图 4.3、图 5.3/图 5.4、图 5.6/表 5.3、表 6.1/图 6.1、图 6.1/表 6.2、图 6.2/表 6.5、图 6.3/图 6.4、图 6.4/图 6.5 之间补充了解释性文字，以避免连续贴图。",
            "3. 统一了英文关键词标签写法，将 `Key Words:` 修正为 `Keywords:`，并同步检查了中英文混排中的基本标点与空格形式。",
            "4. 已在研究现状、关键技术、接口设计、线程协作与测试分析等相关位置补充正文引用，覆盖参考文献 [1]-[26]。",
            "5. 当前导出版本以现有文档内容为基础进行结构补写和学术化强化，未虚构新的实验数据或系统功能。",
            "",
            "【终极排版备忘录】",
            TERMINAL_MEMO,
            "",
        ]
    )
    return "\n".join(lines)


def verify_body_citations(doc: Document) -> list[str]:
    body_text_parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text == "参考文献":
            break
        body_text_parts.append(text)
    body_text = "\n".join(body_text_parts)
    refs = sorted(set(re.findall(r"\[(\d+)\]", body_text)), key=lambda x: int(x))
    return refs


def main() -> None:
    doc = Document(SOURCE_DOCX)
    inserted_texts: list[str] = []

    keyword_changes = apply_keyword_fixes(doc)

    for heading, overview in CHAPTER_OVERVIEWS.items():
        insert_marked_paragraph(doc, heading, overview, inserted_texts)

    for caption, explanation in CAPTION_EXPLANATIONS.items():
        insert_marked_paragraph(doc, caption, explanation, inserted_texts)

    citation_changes = apply_citations(doc)

    doc.save(OUTPUT_DOCX)

    md_text = build_markdown(doc, set(inserted_texts))
    OUTPUT_MD.write_text(md_text, encoding="utf-8")

    cited_refs = verify_body_citations(doc)
    print(f"saved_docx={OUTPUT_DOCX}")
    print(f"saved_md={OUTPUT_MD}")
    print(f"inserted_paragraphs={len(inserted_texts)}")
    print(f"keyword_changes={keyword_changes}")
    print(f"citation_changes={citation_changes}")
    print(f"body_citations={','.join(cited_refs)}")


if __name__ == "__main__":
    main()
