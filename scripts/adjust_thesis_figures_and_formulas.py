from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"e:\face")
SOURCE_MD = BASE_DIR / "换脸系统_初稿_学术化版_终稿统一整改版.md"
SOURCE_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_终稿统一整改版.docx"
OUTPUT_MD = BASE_DIR / "换脸系统_初稿_学术化版_终稿统一整改版_图位公式调整版.md"
OUTPUT_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_终稿统一整改版_图位公式调整版.docx"

TRADITIONAL_FORMULAS = [
    "若将源人脸关键点集合记为 P_s = {p_i^s}_{i=1}^n，目标人脸关键点集合记为 P_t = {p_i^t}_{i=1}^n，则传统三角剖分换脸通常先基于关键点进行 Delaunay 三角剖分，得到三角形集合 T = {tau_k}_{k=1}^m，并在源图与目标图之间建立逐三角形对应关系。对于第 k 个对应三角形，其局部几何映射一般可表示为 [x_t, y_t]^T = A_k [x_s, y_s]^T + b_k，其中 A_k 表示 2 x 2 仿射变换矩阵，b_k 表示平移向量。该表达式反映了传统方法将整体换脸问题分解为多个局部线性形变子问题的基本思想。",
    "在具体像素重映射过程中，若目标三角形内部像素 p 满足 p = lambda_1 v_1^t + lambda_2 v_2^t + lambda_3 v_3^t，且 lambda_1 + lambda_2 + lambda_3 = 1、lambda_i >= 0，则可利用同一组重心坐标在源三角形中定位对应像素 q = lambda_1 v_1^s + lambda_2 v_2^s + lambda_3 v_3^s，并完成纹理采样。完成局部仿射变换后，基础融合常写为 I_out(p) = M(p) I_warp(p) + (1 - M(p)) I_t(p)，其中 M(p) 表示融合掩模，I_warp(p) 表示变形后的人脸纹理，I_t(p) 表示目标图像像素。若进一步采用无缝融合，则可通过最小化 int_Omega ||grad f - grad I_warp||^2 dOmega 来减弱边界拼接痕迹，这也是传统方法常见的后处理思路[1][2]。",
]

DEEP_FORMULAS = [
    "基于深度学习的人脸替换可抽象为身份特征提取与目标属性保持的联合映射过程。设源人脸图像为 I_s，目标图像或目标帧为 I_t，身份编码器 F_id 从源图中提取身份嵌入 e_s = F_id(I_s)，随后由生成或替换网络给出输出 I_hat = G(I_t, e_s)。在该表达式中，e_s 主要承担身份信息约束，而 I_t 中的姿态、表情、光照与背景结构则作为目标属性被尽量保留。因此，深度学习方法的关键不再只是局部几何对齐，而是如何在高维特征空间内同时协调身份一致性与目标场景一致性。",
    "从训练目标看，此类方法通常将总损失写为 L = lambda_id L_id + lambda_rec L_rec + lambda_per L_per + lambda_adv L_adv。其中，身份保持项可表示为 L_id = 1 - cos(E(I_hat), E(I_s))，E(.) 表示用于身份判别的嵌入网络；L_rec 用于约束内容重建或结构一致性，L_per 用于衡量高层感知差异，L_adv 则用于提升结果的真实感与分布一致性。就本文所使用的 InsightFace 方案而言，其工程优势在于能够直接调用预训练身份表征与换脸能力，在不重新训练系统模型的前提下完成较高质量的人脸身份迁移[4][6][9][10][11][12]。",
]

CASE_INTRO_OLD = "为了使测试分析不局限于统计表与文字说明，本文结合 image 目录中新增的样例图像，对传统方法与 InsightFace 的输出结果进行了重新整理，并生成两组对比图用于说明系统在不同素材条件下的表现差异。需要指出的是，当前图像材料分别来自静态图像处理、历史输出样例与摄像头运行界面，因此本节重点关注工程可用性与视觉表现特征，而不将其视为严格控制变量的数值对照实验。"
CASE_INTRO_NEW = "为了使测试分析不局限于统计表与文字说明，本文结合 image 目录中新增的样例图像，对传统方法与 InsightFace 的输出结果进行了重新整理，并生成两组对比图用于说明系统在真实照片源人脸与绘制源人脸条件下的表现差异。需要指出的是，当前图像材料主要来自静态图像处理与历史输出样例，因此本节重点关注工程可用性与视觉表现特征，而不将其视为严格控制变量的数值对照实验。"


def figure_prefix(text: str) -> str | None:
    match = re.match(r"^(图\s+\d+\.\d+)", text.strip())
    return match.group(1) if match else None


def insert_before_marker(lines: list[str], marker: str, paragraphs: list[str]) -> None:
    for idx, line in enumerate(lines):
        if line.strip() == marker:
            existing = [candidate.strip() for candidate in lines[max(0, idx - 8):idx] if candidate.strip()]
            if paragraphs[0] in existing:
                return
            block: list[str] = []
            for paragraph in paragraphs:
                block.extend([paragraph, ""])
            lines[idx:idx] = block
            return
    raise ValueError(f"Marker not found in markdown: {marker}")


def remove_figure_explanations_md(lines: list[str]) -> list[str]:
    result: list[str] = []
    previous_nonempty = ""
    for line in lines:
        stripped = line.strip()
        current_prefix = figure_prefix(stripped) if stripped else None
        previous_prefix = figure_prefix(previous_nonempty) if previous_nonempty else None
        if stripped and current_prefix and previous_prefix == current_prefix and stripped != previous_nonempty:
            continue
        result.append(line)
        if stripped:
            previous_nonempty = stripped
    return result


def update_markdown() -> None:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()

    insert_before_marker(lines, "表 2.1 传统方法与深度学习方法对比", TRADITIONAL_FORMULAS)
    insert_before_marker(lines, "表 2.2 关键技术与系统作用映射表", DEEP_FORMULAS)

    lines = remove_figure_explanations_md(lines)

    old_camera_caption = "图 6.5 摄像头模式运行界面"
    lines = [line for line in lines if line.strip() != old_camera_caption]
    lines = [line for line in lines if "图 6.5 进一步展示了系统在摄像头模式下的实际运行状态" not in line]
    lines = [
        CASE_INTRO_NEW
        if "当前图像材料分别来自静态图像处理、历史输出样例与摄像头运行界面" in line
        else line
        for line in lines
    ]

    for idx, line in enumerate(lines):
        if line.strip() == "图 5.6 关键线程与模块协作时序图":
            lines[idx] = "图 5.7 关键线程与模块协作时序图"

    for idx, line in enumerate(lines):
        if line.strip() == "图 5.5 摄像头模式处理流程图":
            next_nonempty = next((candidate.strip() for candidate in lines[idx + 1:] if candidate.strip()), "")
            if next_nonempty != "图 5.6 摄像头模式运行界面截图":
                lines[idx + 1:idx + 1] = ["", "图 5.6 摄像头模式运行界面截图"]
            break

    text = "\n".join(lines).rstrip() + "\n"
    OUTPUT_MD.write_text(text, encoding="utf-8")


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_normal_style(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.append(p_style)
    p_style.set(qn("w:val"), "Normal")


def add_paragraph_after(anchor: Paragraph, text: str, *, bold: bool = False) -> Paragraph:
    new_para = paragraph_after(anchor)
    set_normal_style(new_para)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def previous_nonempty_paragraph(paragraph: Paragraph) -> Paragraph | None:
    current = paragraph._p.getprevious()
    while current is not None:
        if current.tag != qn("w:p"):
            current = current.getprevious()
            continue
        para = Paragraph(current, paragraph._parent)
        text = (para.text or "").strip()
        if text or "w:drawing" in para._p.xml:
            return para
        current = current.getprevious()
    return None


def next_nonempty_paragraph(paragraph: Paragraph) -> Paragraph | None:
    current = paragraph._p.getnext()
    while current is not None:
        if current.tag != qn("w:p"):
            current = current.getnext()
            continue
        para = Paragraph(current, paragraph._parent)
        text = (para.text or "").strip()
        if text or "w:drawing" in para._p.xml:
            return para
        current = current.getnext()
    return None


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if (paragraph.text or "").strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_if_exists(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if (paragraph.text or "").strip() == text:
            return paragraph
    return None


def insert_formula_paragraphs(doc: Document) -> None:
    for marker, paragraphs in [
        ("表 2.1 传统方法与深度学习方法对比", TRADITIONAL_FORMULAS),
        ("表 2.2 关键技术与系统作用映射表", DEEP_FORMULAS),
    ]:
        anchor = find_paragraph(doc, marker)
        previous = previous_nonempty_paragraph(anchor)
        existing_texts = []
        probe = previous
        for _ in range(4):
            if probe is None:
                break
            text = (probe.text or "").strip()
            if text:
                existing_texts.append(text)
            probe = previous_nonempty_paragraph(probe)
        if paragraphs[0] in existing_texts:
            continue

        current_anchor = previous if previous is not None else anchor
        for paragraph_text in paragraphs:
            current_anchor = add_paragraph_after(current_anchor, paragraph_text, bold=False)


def remove_figure_explanations_docx(doc: Document) -> None:
    to_remove: list[Paragraph] = []
    previous_text = ""
    previous_para: Paragraph | None = None
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        current_prefix = figure_prefix(text)
        previous_prefix = figure_prefix(previous_text)
        if current_prefix and previous_prefix == current_prefix and text != previous_text:
            to_remove.append(paragraph)
            continue
        previous_text = text
        previous_para = paragraph
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def move_camera_figure(doc: Document) -> None:
    old_caption = find_paragraph(doc, "图 6.5 摄像头模式运行界面")
    old_image = previous_nonempty_paragraph(old_caption)
    if old_image is None or "w:drawing" not in old_image._p.xml:
        raise RuntimeError("Failed to locate original camera screenshot image paragraph.")

    anchor_caption = find_paragraph(doc, "图 5.5 摄像头模式处理流程图")
    existing_next = next_nonempty_paragraph(anchor_caption)
    if existing_next is not None and (existing_next.text or "").strip() == "图 5.6 摄像头模式运行界面截图":
        pass
    else:
        new_image_p = deepcopy(old_image._p)
        anchor_caption._p.addnext(new_image_p)
        new_image_para = Paragraph(new_image_p, anchor_caption._parent)
        set_normal_style(new_image_para)
        new_caption_para = paragraph_after(new_image_para)
        set_normal_style(new_caption_para)
        run = new_caption_para.add_run("图 5.6 摄像头模式运行界面截图")
        run.bold = True

    remove_paragraph(old_caption)
    remove_paragraph(old_image)


def renumber_thread_figure(doc: Document) -> None:
    paragraph = find_paragraph(doc, "图 5.6 关键线程与模块协作时序图")
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run("图 5.7 关键线程与模块协作时序图")
    run.bold = True


def update_docx() -> None:
    doc = Document(SOURCE_DOCX)
    insert_formula_paragraphs(doc)
    remove_figure_explanations_docx(doc)
    move_camera_figure(doc)
    renumber_thread_figure(doc)
    intro_para = next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if "当前图像材料分别来自静态图像处理、历史输出样例与摄像头运行界面" in ((paragraph.text or "").strip())
        ),
        None,
    )
    if intro_para is not None:
        for child in list(intro_para._p):
            if child.tag != qn("w:pPr"):
                intro_para._p.remove(child)
        run = intro_para.add_run(CASE_INTRO_NEW)
        run.bold = False
    leftover = next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if "图 6.5 进一步展示了系统在摄像头模式下的实际运行状态" in ((paragraph.text or "").strip())
        ),
        None,
    )
    if leftover is not None:
        remove_paragraph(leftover)
    doc.save(OUTPUT_DOCX)


def main() -> None:
    update_markdown()
    update_docx()
    print(f"output_md={OUTPUT_MD}")
    print(f"output_docx={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
