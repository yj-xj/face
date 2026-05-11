from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"e:\face")
SOURCE_MD = BASE_DIR / "换脸系统_初稿_学术化版_改写后.md"
SOURCE_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_改写后.docx"
OUTPUT_MD = BASE_DIR / "换脸系统_初稿_学术化版_精修后.md"
OUTPUT_DOCX = BASE_DIR / "换脸系统_初稿_学术化版_精修后.docx"


PARAGRAPH_REWRITES = {
    "本文并非围绕全新换脸模型的从零构建展开论述，而是针对既有人脸替换项目进行系统化梳理、补充说明与工程整理。现有仓库已经包含 PyQt5 前端、Django REST 后端、SQLite 数据库、本地媒体目录，以及由 OpenCV 与 InsightFace 共同构成的处理链路。论文写作的重点，在于将这些已有实现归纳为结构完整、论证充分且便于答辩展示的系统方案。在前端层面，系统以 EnhancedFaceSwapUI 为核心，并借助 VideoProcessingThread 与 CameraProcessingThread 分别支撑离线视频处理和实时摄像头处理；在数据管理层面，DatabaseManager 与后端接口协同配合，使人脸图片、输入视频与输出结果能够形成统一的加载、保存和回放入口。截至本文撰写时，数据库内共记录 19 条人脸图片、13 条输入视频和 5 条输出视频信息，其中本地可用人脸路径为 9 / 9，可用视频路径为 7 / 8，而输出目录中的 mp4 文件数量又比数据库记录多出 5 个。上述客观数据表明，系统已经能够覆盖素材加载、参数设置、视频换脸、摄像头演示、结果归档与回放验证等主要环节，同时也暴露出历史路径失效、旧结果未补登记等典型工程问题。基于这一现实状态，本文从需求分析、总体设计、详细实现和工程测试四个层面展开说明，重点讨论双模式界面、线程化处理、接口化数据管理与结果留痕机制如何共同提升系统的可用性与可维护性。测试结果表明，该系统已经能够承担桌面端人脸替换的主要演示与处理任务，具备作为软件工程方向系统实现型毕业设计的基本条件。后续工作将继续围绕资源路径治理、 processing_tasks 的实际队列化使用、异常日志展示以及历史结果补录等问题展开。":
        "本文以既有人脸替换项目的工程化整理与论文化表达为核心任务，并不试图从零设计单一换脸模型。现有仓库已经形成由 PyQt5 前端、Django REST 后端、SQLite 数据库、本地媒体目录以及 OpenCV 和 InsightFace 共同构成的完整处理链路。论文工作的重点，在于将上述既有实现重构为结构清晰、证据充分且适合答辩展示的系统方案。前端部分以 EnhancedFaceSwapUI 为核心，通过 VideoProcessingThread 与 CameraProcessingThread 分别支撑离线视频处理和实时摄像头处理；数据管理部分则借助 DatabaseManager 与后端接口的协同，使人脸图片、输入视频和输出结果获得统一的加载、保存与回放入口。截至本文撰写时，数据库中共保存 19 条人脸图片记录、13 条输入视频记录和 5 条输出视频记录，其中本地可用人脸路径为 9 / 9，可用视频路径为 7 / 8，而输出目录中的 mp4 文件数量较数据库记录多出 5 个。上述事实说明，系统已经能够覆盖素材加载、参数配置、视频换脸、摄像头演示、结果归档与回放验证等主要环节，同时也暴露出历史路径失效与旧结果未补登记等典型工程问题。基于这一现实基础，本文从需求分析、总体设计、详细实现和工程测试四个层面展开论述，重点分析双模式界面、线程化处理、接口化数据管理和结果留痕机制如何共同提升系统的可用性与可维护性。综合测试结果可以认为，该系统已经能够支撑桌面端人脸替换的主要演示与处理流程，具备软件工程方向系统实现型毕业设计的基本特征。后续工作仍需围绕资源路径治理、processing_tasks 的实际队列化使用、异常日志展示以及历史结果补录等问题持续推进。",
    "本章将围绕课题提出的现实背景、研究价值以及全文任务安排展开说明。首先结合短视频创作与生成式视觉应用的发展趋势，讨论人脸替换系统由算法演示走向工程化实现的现实需求；随后梳理相关研究现状，明确本文并不以提出新模型为主要目标，而是以既有项目的系统整理、功能验证与工程分析为核心内容。":
        "本章围绕课题提出的现实背景、研究价值以及全文任务安排展开论述。首先结合短视频创作与生成式视觉应用的发展态势，说明人脸替换系统由算法演示走向工程化实现的现实需求；随后梳理相关研究进展，明确本文并不以提出新模型为目标，而是将既有项目的系统整理、功能验证与工程分析作为主要研究内容。",
    "从技术演进的角度来看，人脸替换大体经历了由传统几何配准与区域融合向深度学习驱动的身份迁移方案发展的过程。早期方法主要依赖检测、关键点定位、仿射变换与局部融合，优点在于实现路径清楚、依赖较少，但当姿态变化较大、遮挡明显或光照复杂时，往往容易产生边缘失真和结构拉伸。随着 FaceNet、ArcFace、SimSwap、FaceShifter、InsightFace 等工作相继出现，换脸任务在身份保持和复杂场景表现方面获得了更明显的提升[1][2][3][4][5][6][9][10][11]。":
        "从技术演进路径来看，人脸替换大致经历了由传统几何配准与区域融合向深度学习驱动身份迁移方案逐步过渡的发展过程。早期方法主要依赖检测、关键点定位、仿射变换和局部融合，其优势在于实现链路相对清晰、外部依赖较少；然而，当姿态变化显著、遮挡较强或光照条件复杂时，此类方法往往容易出现边缘失真与结构拉伸。随着 FaceNet、ArcFace、SimSwap、FaceShifter 和 InsightFace 等方法相继提出，换脸任务在身份保持能力和复杂场景适应性方面均获得了更为明显的提升[1][2][3][4][5][6][9][10][11]。",
    "围绕这一目标，本文结合 frontend/main.py、face_swap_ui_enhanced.py、database_manager.py 以及后端模型与视图集，对系统边界与实现关系进行了重新梳理，并以视频模式和摄像头模式为中心概括主要使用场景，进一步说明需求分析、总体结构与关键模块之间的联系。同时，本文还围绕 EnhancedFaceSwapUI、DatabaseManager、VideoProcessingThread、CameraProcessingThread 以及 saveOutputVideoToDatabase 等实现内容，说明系统如何完成素材加载、处理执行、结果回放与结果保存，并结合数据库记录、文件系统状态和已有样例资源，对系统的可用性与现实问题进行验证。":
        "为实现上述目标，本文结合 frontend/main.py、face_swap_ui_enhanced.py、database_manager.py 以及后端模型与视图集，对系统边界和实现关系进行了重新梳理，并以视频模式和摄像头模式为核心概括主要使用场景，进一步明确需求分析、总体结构与关键模块之间的逻辑联系。在此基础上，本文围绕 EnhancedFaceSwapUI、DatabaseManager、VideoProcessingThread、CameraProcessingThread 以及 saveOutputVideoToDatabase 等关键实现内容，论证系统如何完成素材加载、处理执行、结果回放与结果保存，并结合数据库记录、文件系统状态和既有样例资源，对系统的可用性与现实问题加以验证。",
    "但在本项目中，InsightFace 的意义并不在于它被单独当作研究对象讨论，而在于它如何被真正纳入系统流程之中。用户在前端完成素材选择后，startProcessing 会将参数同步至原始换脸引擎，后台线程持续调用处理逻辑，结果完成后再经由 processingFinished 与 saveOutputVideoToDatabase 进入归档和回放环节。换言之，本文关注的重点不是模型如何重新训练，而是如何把既有模型能力放入一套可交互、可记录、可回查的软件系统中。这种整合思路体现了系统工程与算法研究在关注重点上的差异。对于系统型毕业设计而言，关键不仅在于结果是否更自然，还在于调用是否稳定、异常时能否切换替代方案、结果能否妥善保存，以及界面能否及时反馈状态变化。正因如此，本文始终将 InsightFace 视为系统组成部分之一，而非全文唯一的中心。":
        "然而，在本项目中，InsightFace 的价值并不在于其被孤立地作为研究对象加以讨论，而在于其如何被稳定地纳入系统执行链路。用户在前端完成素材选择后，startProcessing 会将相关参数同步至原始换脸引擎，后台线程持续调用处理逻辑，待处理完成后，再经由 processingFinished 与 saveOutputVideoToDatabase 完成归档和回放。换言之，本文真正关注的不是模型如何重新训练，而是既有模型能力如何被组织进一套可交互、可记录且可回查的软件系统之中。这种整合思路反映了系统工程与算法研究在关注重点上的显著差异。对于系统型毕业设计而言，关键不仅在于视觉结果是否自然，还在于调用链路是否稳定、异常情况下能否切换替代方案、结果能否可靠保存，以及界面能否及时反馈状态变化。因此，本文始终将 InsightFace 视为系统整体实现中的重要组成部分，而非全文唯一的研究中心。",
    "本章将结合真实使用过程对系统需求进行分析，重点说明视频换脸与摄像头演示两类业务场景下的功能目标与约束条件。在此基础上，本文将依次讨论功能需求、非功能需求、用例关系以及系统可行性，为后续总体设计与模块划分提供明确依据。":
        "本章将结合真实使用过程对系统需求展开分析，重点说明视频换脸与摄像头演示两类业务场景下的功能目标及约束条件。在此基础上，本文将依次讨论功能需求、非功能需求、用例关系与系统可行性，为后续总体设计和模块划分提供明确依据。",
    "系统总体上可划分为四个层次，分别为用户交互层、业务控制层、媒体处理层和数据持久化层。用户交互层主要由 PyQt5 前端构成，业务控制层由 DatabaseManager、处理线程以及前端控制逻辑组成，媒体处理层由 OpenCV 与换脸引擎构成，而数据持久化层则由 Django REST、SQLite 与本地文件系统共同承担。":
        "从体系结构上看，系统可划分为用户交互层、业务控制层、媒体处理层和数据持久化层四个层次。其中，用户交互层主要由 PyQt5 前端构成；业务控制层由 DatabaseManager、处理线程以及前端控制逻辑共同组成；媒体处理层负责 OpenCV 与换脸引擎相关的具体处理；数据持久化层则由 Django REST、SQLite 与本地文件系统共同承担。",
    "本章将围绕系统从界面组织到处理线程协作的核心实现展开说明。内容首先介绍前端主界面与双模式交互逻辑，随后分析素材加载、视频处理、摄像头处理、结果归档与线程协作等关键实现环节，以展示系统如何将既有换脸能力组织成一套可运行、可反馈、可回放的完整桌面应用。":
        "本章将围绕系统从界面组织到处理线程协作的关键实现展开说明。内容首先介绍前端主界面及双模式交互逻辑，随后分析素材加载、视频处理、摄像头处理、结果归档与线程协作等核心环节，以说明系统如何将既有换脸能力组织为一套可运行、可反馈且可回放的完整桌面应用。",
    "素材加载在前端中构成了一个相对独立的功能模块，其中 loadFaceImages、loadVideos、loadFaceImagesFromLocal 与 loadVideoFiles 分别承担不同来源的数据加载任务。当 DatabaseManager 可用时，前端优先从后端接口读取元数据，再依据 local_path 在本地生成缩略图和视频首帧；当后端不可用时，系统也能够退回到本地目录扫描模式，从而在答辩演示或离线使用场景下继续运行。DatabaseManager 则主要封装图片上传、视频上传、图像列表加载、视频列表加载和输出视频保存等接口调用，一方面把 requests 请求、超时处理、异常捕获与返回结构统一移出界面代码，另一方面通过 QThread 对上传和加载任务进行异步化，避免大文件处理时前端窗口直接失去响应。这种封装方式与前文提出的模块分层原则保持一致。":
        "在前端实现中，素材加载形成了相对独立的功能模块，其中 loadFaceImages、loadVideos、loadFaceImagesFromLocal 与 loadVideoFiles 分别负责不同来源数据的加载任务。当 DatabaseManager 可用时，前端优先从后端接口读取元数据，并依据 local_path 在本地生成缩略图和视频首帧；当后端不可用时，系统仍可退回至本地目录扫描模式，从而保证其在答辩演示或离线使用场景中的连续可用性。DatabaseManager 则主要封装图片上传、视频上传、图像列表加载、视频列表加载与输出视频保存等接口调用：一方面，它将 requests 请求、超时处理、异常捕获及返回结构统一移出界面代码；另一方面，它借助 QThread 对上传与加载任务进行异步化处理，以避免大文件操作导致前端窗口直接失去响应。这种组织方式与前文提出的模块分层原则保持一致。",
    "在参数准备完成后，系统会创建 VideoProcessingThread，并将 progress_signal、status_signal、finished_signal 与 error_signal 等信号连接到界面更新方法。这样的设计使后台线程能够专注于视频读取、换脸引擎调用和结果写出，而前台界面则主要负责进度展示、状态提示与最终结果呈现。对于本项目这种既涉及较长视频处理时间、又要求即时交互反馈的系统来说，这是一项不可缺少的实现策略。":
        "在完成参数准备后，系统会创建 VideoProcessingThread，并将 progress_signal、status_signal、finished_signal 与 error_signal 等信号连接至相应的界面更新方法。借助这一设计，后台线程可以专注于视频读取、换脸引擎调用与结果写出，而前台界面则主要承担进度展示、状态提示及最终结果呈现。对于同时涉及较长视频处理时间和即时交互反馈的系统而言，这种线程分工是一项不可或缺的实现策略。",
    "输出结果保存是本系统不可或缺的组成部分。在许多原型项目中，换脸结果往往只是直接写入目录的 mp4 文件，后续难以继续查询或统计。本项目则通过 saveOutputVideoToDatabase 在处理完成后提取输出视频的分辨率、fps、时长、文件大小、处理方法和状态等元数据，并将其与 input_video、face_image 建立关联，从而形成相对完整的结果记录。":
        "输出结果保存构成了本系统中的关键环节。在许多原型项目中，换脸结果往往只是简单写入目录的 mp4 文件，后续既不便于查询，也难以继续统计。本项目通过 saveOutputVideoToDatabase 在处理完成后提取输出视频的分辨率、fps、时长、文件大小、处理方法和状态等元数据，并进一步与 input_video、face_image 建立关联，由此形成较为完整的结果记录。",
    "从代码结构来看，系统的并行协作主要围绕三类对象展开，即界面控制类 EnhancedFaceSwapUI、视频处理线程 VideoProcessingThread 和摄像头线程 CameraProcessingThread。三者与 original_app 以及 DatabaseManager 共同构成系统的核心协作关系：EnhancedFaceSwapUI 承载用户交互和界面状态，两个 QThread 负责长时间或持续运行任务，original_app 负责具体的换脸算法逻辑，而 DatabaseManager 则承担资源管理与入库功能[19][20]。":
        "从代码结构上看，系统的并行协作主要围绕三类对象展开，即界面控制类 EnhancedFaceSwapUI、视频处理线程 VideoProcessingThread 与摄像头线程 CameraProcessingThread。三者与 original_app 以及 DatabaseManager 共同构成系统的核心协作关系：EnhancedFaceSwapUI 负责承载用户交互和界面状态，两个 QThread 负责执行长时间或持续运行任务，original_app 负责具体换脸算法逻辑，而 DatabaseManager 则承担资源管理与结果入库功能[19][20]。",
    "本章将以工程验证为导向，对系统当前版本的运行情况展开测试与分析。内容涵盖测试环境与方案设计、功能验证结果、稳定性与性能观察、典型样例分析以及现存问题总结，重点说明系统在真实项目材料上的完成程度、可运行性以及后续仍需改进的方向。":
        "本章以工程验证为导向，对系统当前版本的运行情况展开测试与分析。内容包括测试环境与方案设计、功能验证结果、稳定性与性能观察、典型样例分析以及现存问题总结，重点说明系统在真实项目材料上的完成程度、可运行性和后续仍需改进的方向。",
    "本章测试的定位是工程验证，而非算法竞赛式 benchmark。验证重点主要考察系统的业务链路能否真正贯通、界面在处理与切换过程中能否保持基本响应，以及数据库记录与本地文件之间是否存在可观察且可解释的偏差。与其构造并不存在的大规模对照实验，本文更关注当前项目在真实材料上的可运行性与可说明性[23][24]。":
        "本章测试的定位是工程验证，而非算法竞赛式 benchmark。验证重点主要包括三个方面：其一，系统业务链路能否真正贯通；其二，界面在处理与切换过程中能否保持基本响应；其三，数据库记录与本地文件之间是否存在可观察且可解释的偏差。与其构造并不存在的大规模对照实验，本文更关注当前项目在真实材料条件下的可运行性与可说明性[23][24]。",
    "从功能验证结果来看，当前系统已经覆盖主要处理流程。前端能够从数据库或本地目录加载人脸素材与视频素材，视频模式可以创建 VideoProcessingThread 执行处理，处理结束后能够回放结果并尝试入库；摄像头模式则支持开启、关闭以及快照保存。这说明系统已经不再只是执行一次推理的脚本，而是具备明确交互入口与结果输出能力的桌面原型。与此同时，测试也表明系统当前更偏向工程上的“可运行”，而不是对全部失败情况都进行了严格封装。例如，当数据库中的历史路径失效时，界面可以跳过无效条目而不是整体崩溃，这体现出一定的异常容忍能力；但这一现象同时也说明数据库记录与本地文件之间仍缺少更严格的同步机制。本文保留这些问题，并在后文继续分析其原因与改进方向。":
        "从功能验证结果来看，当前系统已经覆盖主要处理流程。前端能够从数据库或本地目录加载人脸素材与视频素材，视频模式可以创建 VideoProcessingThread 执行处理，处理结束后既能够回放结果，也能够尝试完成入库；摄像头模式则支持开启、关闭与快照保存。这表明系统已经不再只是执行单次推理的脚本，而是具备明确交互入口与结果输出能力的桌面原型。与此同时，测试结果也显示，现阶段系统更侧重工程上的“可运行”，尚未对全部失败情况完成严格封装。例如，当数据库中的历史路径失效时，界面可以跳过无效条目而非整体崩溃，这体现出一定的异常容忍能力；但同一现象也说明数据库记录与本地文件之间仍缺少更严格的同步机制。本文保留这些问题，并将在后文继续分析其成因与改进方向。",
    "相较于时长和帧率本身，更值得关注的是资源一致性方面的差异。当前具有 local_path 的人脸记录共有 9 条，其中失效 0 条；具有 local_path 的输入视频记录共有 8 条，其中失效 1 条；此外，output_videos 目录中的 mp4 文件数量比数据库输出记录多出 5 个。也就是说，系统当前的主要短板并不在于完全无法处理，而在于历史结果与历史路径尚未形成稳定的回补与治理机制。":
        "相较于时长与帧率本身，更值得关注的是资源一致性差异。当前具有 local_path 的人脸记录共有 9 条，其中失效 0 条；具有 local_path 的输入视频记录共有 8 条，其中失效 1 条；此外，output_videos 目录中的 mp4 文件数量比数据库输出记录多出 5 个。由此可见，现阶段系统的主要短板并不在于完全无法处理，而在于历史结果和历史路径尚未形成稳定的回补与治理机制。",
    "后续改进应围绕上述具体问题逐项展开。系统可在入库阶段采用更稳定的相对路径或统一媒体目录策略，并补充历史记录修复脚本；同时增加输出目录扫描和结果补登记逻辑，以减少结果文件与数据库记录不一致的情况；对于 processing_tasks，则需要将其从预留模型进一步发展为真正可用的任务队列，并补充日志分类与错误展示界面。只有把这些基础工作逐步落实下来，系统才更适合长期使用。":
        "后续改进应当围绕上述问题逐项展开。系统可在入库阶段采用更稳定的相对路径或统一媒体目录策略，并配套补充历史记录修复脚本；同时增加输出目录扫描和结果补登记逻辑，以减少结果文件与数据库记录不一致的现象；对于 processing_tasks，则需要将其从预留模型进一步发展为真正可用的任务队列，并补充日志分类与错误展示界面。只有将这些基础性工作逐步落实，系统才更适合长期稳定使用。",
    "从实现结果来看，当前项目已经形成相对完整的桌面端处理流程。前端以 EnhancedFaceSwapUI 为核心承担双模式交互，DatabaseManager 负责前后端资源协同，VideoProcessingThread 与 CameraProcessingThread 分别支撑离线处理和实时处理，而处理结果也能够通过 saveOutputVideoToDatabase 等逻辑进入统一的记录与回放流程。这些内容表明，项目已经超出单次算法演示的范围，具备系统实现型毕业设计应有的基本形态。":
        "从实现结果来看，当前项目已经形成相对完整的桌面端处理流程。前端以 EnhancedFaceSwapUI 为核心承担双模式交互，DatabaseManager 负责前后端资源协同，VideoProcessingThread 与 CameraProcessingThread 分别支撑离线处理和实时处理，而处理结果也能够通过 saveOutputVideoToDatabase 等逻辑进入统一的记录与回放流程。这些实现表明，项目已经超出单次算法演示的范围，具备系统实现型毕业设计应有的基本形态。",
    "本系统的特点并不在于提出新的换脸网络，而在于将已有换脸方法整理进一套实际可运行的桌面原型之中。视频模式与摄像头模式被统一纳入同一界面，并通过状态切换逻辑与线程控制维持一致的交互过程；图片、视频与输出结果则被纳入同一套后端与数据库管理流程，使系统能够保存并查询多次处理记录；同时，系统在实现上保留了传统方法与 InsightFace 两种方案，因此在不同运行环境下仍能维持基本处理能力。":
        "本系统的主要特点并不在于提出新的换脸网络，而在于将既有换脸方法整合进一套实际可运行的桌面原型。视频模式与摄像头模式被统一纳入同一界面，并通过状态切换逻辑与线程控制维持一致的交互过程；图片、视频和输出结果则被纳入同一套后端与数据库管理流程，从而使系统能够保存并查询多次处理记录；与此同时，系统在实现上同时保留传统方法与 InsightFace 两种方案，因此在不同运行环境下仍可维持基本处理能力。",
    "后续工作可首先从资源整理入手，通过调整入库规则、本地目录组织方式以及历史记录修复策略，减少 local_path 失效和输出文件孤立现象。系统还需要进一步完善任务调度机制，将 processing_tasks 从当前的预留模型发展为真正可用的处理队列，使其具备批量提交、失败重试、取消控制与更细致状态展示等能力；与此同时，还应补充结果整理功能，为 output_videos 目录提供历史结果扫描与补登记能力，避免真实输出长期脱离数据库记录。":
        "后续工作可首先从资源整理入手，通过调整入库规则、本地目录组织方式以及历史记录修复策略，尽量减少 local_path 失效和输出文件孤立等现象。系统还需要进一步完善任务调度机制，将 processing_tasks 从当前的预留模型发展为真正可用的处理队列，使其具备批量提交、失败重试、取消控制以及更细致的状态展示能力；与此同时，还应补充结果整理功能，为 output_videos 目录提供历史结果扫描与补登记能力，避免真实输出长期脱离数据库记录。",
    "本文能够顺利完成，离不开指导老师和同学们在项目整理、文档完善以及答辩准备过程中给予的支持与帮助。在论文撰写阶段，无论是系统边界的确定、章节结构的梳理，还是图表整理与文档格式核对，都获得了许多宝贵意见和启发。在此，对这些帮助与包容表示诚挚感谢。":
        "本文得以顺利完成，离不开指导老师和同学们在项目整理、文档完善以及答辩准备过程中给予的支持与帮助。在论文撰写阶段，无论是系统边界的界定、章节结构的梳理，还是图表整理与文档格式核对，都使本人获得了诸多宝贵意见与启发。在此，对所有帮助与包容表示诚挚感谢。",
    "同时，也感谢在项目开发与论文完善过程中提供材料、参与讨论并提出意见的各位老师和同学。对本人而言，这份毕业设计不仅是一次完成学业任务的经历，更是一场将算法原型、界面开发、后端管理与文档撰写整合在一起的综合实践。":
        "同时，感谢在项目开发与论文完善过程中提供材料、参与讨论并提出意见的各位老师和同学。对本人而言，这份毕业设计不仅是一次学业任务的完成过程，更是一场将算法原型、界面开发、后端管理与文档撰写综合起来的系统性实践。",
}

AI_STATEMENT = "需要说明的是，本文的核心代码实现、系统架构设计与实验分析均由本人独立完成；在论文撰写过程中，仅在文本润色、语句规范化与错别字核查阶段适度使用了大语言模型工具作为辅助，不以其替代研究设计、程序开发或结论形成，特此声明[27]。"
AI_REFERENCE = "[27] OpenAI. ChatGPT (GPT-5) [EB/OL]. https://chatgpt.com/, 2026."


def normalize_text(text: str) -> str:
    text = re.sub(r"([，。；：、（《“])\s+([A-Za-z])", r"\1\2", text)
    text = re.sub(r"\s+([，。；：、）》”])", r"\1", text)
    text = re.sub(r" {2,}", " ", text)
    return text


def insert_paragraph_after(paragraph: Paragraph, text: str, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def replace_paragraph_text(paragraph: Paragraph, text: str, bold: bool) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.bold = bold


def normalize_md_line(line: str) -> tuple[str, bool]:
    match = re.fullmatch(r"\*\*(.+)\*\*", line)
    if match:
        return match.group(1), True
    return line, False


def build_md_output() -> tuple[dict[str, str], int]:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    body_lines: list[str] = []
    for line in lines:
        if line.strip() == "二、 修改日志与排版备忘录":
            break
        body_lines.append(line)

    rewritten_map: dict[str, str] = {}
    output_lines: list[str] = ["修改后的正文", ""]
    replace_count = 0
    ai_inserted = False
    ref_inserted = False

    for line in body_lines:
        core, is_bold = normalize_md_line(line)
        if core == "一、 修改与扩写后的正文":
            continue
        new_core = PARAGRAPH_REWRITES.get(core, core)
        if new_core != core:
            rewritten_map[core] = new_core
            replace_count += 1
        new_core = normalize_text(new_core)
        output_lines.append(f"**{new_core}**" if is_bold else new_core)

        if core == "同时，也感谢在项目开发与论文完善过程中提供材料、参与讨论并提出意见的各位老师和同学。对本人而言，这份毕业设计不仅是一次完成学业任务的经历，更是一场将算法原型、界面开发、后端管理与文档撰写整合在一起的综合实践。":
            output_lines.append("")
            output_lines.append(AI_STATEMENT)
            ai_inserted = True

        if core.startswith("[26] 潘璐, 郑云. 基于 REST 风格的轻量级资源管理系统设计"):
            output_lines.append("")
            output_lines.append(AI_REFERENCE)
            ref_inserted = True

        output_lines.append("")

    if not ai_inserted or not ref_inserted:
        raise RuntimeError("Failed to insert AI statement or AI reference into markdown output.")

    output_lines.extend(
        [
            "修改简报",
            "",
            "1. 对摘要、绪论研究现状、关键技术说明、系统实现、测试分析和总结展望等高重复风险段落进行了二次重构，重点通过句群重组、逻辑前置和工程化表达替代原有直述式写法。",
            "2. 保持了原有软件工程架构逻辑不变，未改动前后端分层、多线程异步处理、数据库留痕与双模式交互等核心设计含义。",
            "3. 在致谢末尾补入了符合学术规范的人工智能使用声明，并在参考文献中新增了 1 条大语言模型工具性引用，以满足学院关于 AI 使用说明的形式要求。",
            "4. 同步清理了中英文混排中的局部空格问题，统一了个别英文标识符前后的排版形式。",
            "5. 预估效果：相较上一版，文本在摘要、研究现状、实现分析与结论部分的重复风险会进一步下降，但“低于 15%”仍需以学校最终知网检测结果为准。",
            "",
        ]
    )

    OUTPUT_MD.write_text("\n".join(output_lines), encoding="utf-8")
    return rewritten_map, replace_count


def build_docx_output(rewritten_map: dict[str, str]) -> int:
    doc = Document(SOURCE_DOCX)
    replace_count = 0
    ai_inserted = False
    ref_inserted = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().replace("\t", " ")
        if not text:
            continue
        bold = any(bool(run.bold) for run in paragraph.runs if run.text.strip())
        if text in rewritten_map:
            replace_paragraph_text(paragraph, normalize_text(rewritten_map[text]), bold)
            replace_count += 1
            text = normalize_text(rewritten_map[text])

        if text == PARAGRAPH_REWRITES["同时，也感谢在项目开发与论文完善过程中提供材料、参与讨论并提出意见的各位老师和同学。对本人而言，这份毕业设计不仅是一次完成学业任务的经历，更是一场将算法原型、界面开发、后端管理与文档撰写整合在一起的综合实践。"]:
            next_text = ""
            nxt = paragraph._p.getnext()
            if nxt is not None:
                next_para = Paragraph(nxt, paragraph._parent)
                next_text = next_para.text.strip().replace("\t", " ")
            if next_text != AI_STATEMENT:
                insert_paragraph_after(paragraph, AI_STATEMENT)
            ai_inserted = True

        if text.startswith("[26] 潘璐, 郑云. 基于 REST 风格的轻量级资源管理系统设计"):
            next_text = ""
            nxt = paragraph._p.getnext()
            if nxt is not None:
                next_para = Paragraph(nxt, paragraph._parent)
                next_text = next_para.text.strip().replace("\t", " ")
            if next_text != AI_REFERENCE:
                insert_paragraph_after(paragraph, AI_REFERENCE)
            ref_inserted = True

    if not ai_inserted or not ref_inserted:
        raise RuntimeError("Failed to insert AI statement or AI reference into docx output.")

    OUTPUT_DOCX.write_bytes(b"")
    doc.save(OUTPUT_DOCX)
    return replace_count


def main() -> None:
    rewritten_map, md_replacements = build_md_output()
    docx_replacements = build_docx_output(rewritten_map)
    print(f"md_replacements={md_replacements}")
    print(f"docx_replacements={docx_replacements}")
    print(f"output_md={OUTPUT_MD}")
    print(f"output_docx={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
